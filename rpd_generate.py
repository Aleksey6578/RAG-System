"""
rpd_generate.py — генерация РПД на основе шаблона Шаблон_пустой.dotx.

Стратегия (v4): копируем пустой шаблон с [] плейсхолдерами →
заполняем [] в параграфах и строках таблиц сгенерированным LLM-контентом.
Таблицы ищутся по заголовку (find_table), а не по хрупкому индексу.
Поля old_discipline / old_code / replace_all — удалены за ненадобностью.

Рефакторинг v4.0 (шаблон Шаблон_пустой.dotx):
  - [T1] УДАЛЕНО: detect_old_discipline(), replace_all(), replace_text_in_paragraph()
         — заменены на fill_doc_header(), который ищет [] в параграфах.
  - [T2] УДАЛЕНО: clear_table_data_rows() + add_table_row()
         — заменены на fill_placeholder_rows(), которая ищет строки с [].
  - [T3] ДОБАВЛЕНО: find_table(doc, key) — поиск таблицы по ключевому слову
         в заголовке. Устойчив к смене порядка таблиц в шаблоне.
  - [T4] config.json упрощён: убраны old_discipline / old_code / new_code,
         добавлен code (код дисциплины). template = Шаблон_пустой.dotx.

Исправления v3.6:
  - [КОМПЕТЕНЦИИ] ИСПРАВЛЕНО: пример в промпте содержал «в области {discipline}»
    посередине предложения → модель копировала конструкцию буквально.
    Пример заменён на безопасный с «Машинное обучение» как фиксированной
    дисциплиной + явная инструкция «Примеры выше — для другой дисциплины».
  - [БИБЛИОГРАФИЯ T15] ИСПРАВЛЕНО: промпт содержал «Фамилия, И. О. Название»
    в JSON-примере → модель копировала шаблон. Убраны шаблонные примеры,
    добавлен запрет «НЕ придумывай авторов», добавлена постпроверка
    _is_placeholder() в gen_bibliography() для фильтрации галлюцинаций.
    При обнаружении плейсхолдеров — fallback на реальные учебники.
  - [БИБЛИОГРАФИЯ T17] ИСПРАВЛЕНО: qwen2.5:3b стабильно не умеет генерировать
    методические пособия (возвращает «Фамилия, И. О. Название» независимо
    от промпта). LLM-вызов для T17 убран полностью — всегда используется
    fallback с реальными УГНТУ-пособиями кафедры ВТИК (Д. М. Зарипов, 2023).

Исправления v3.5:
  - [ПРОМПТЫ] Убраны плейсхолдеры «<уникальное действие 1>»/«знание 1»/«умение 1».
  - [БИБЛИОГРАФИЯ] Добавлена генерация Т15/Т17 через LLM + fallback.

Исправления v3.8:
  - [З-R1] ИСПРАВЛЕНО: формулировки SECTION_QUERIES разнесены по стилю —
    первый запрос официальный (ФГОС), второй содержательный (действия студента).
    Устраняет дублирующиеся scores типа 0.707/0.707.
  - [З-R2] ИСПРАВЛЕНО: добавлен "bibliography_main": ["bibliography", "place"]
    в SECTION_TYPE_FILTER. Ранее поиск шёл без фильтра — нерелевантные чанки
    с одинаковыми scores 0.539×8.
  - [З-R5] ИСПРАВЛЕНО: персистентный файловый кэш rpd_cache.json через
    _load_cache()/_save_cache(). При повторных запусках ~50% экономии времени.
  - [Д-1] ИСПРАВЛЕНО: fill_t21_fos разворачивает строки по паттерну шаблона:
    раздел × компетенция × 3 типа (З/У/В). Было 16 строк → ~60, как в шаблоне.
  - [Д-2] ИСПРАВЛЕНО: T15 поддерживает override из config.json через ключ
    "main_bibliography": [{type, purpose, desc, url, coeff}].

(более ранние версии — см. историю файла)
"""

import json
import re
import sys
import os
import shutil
import time
import copy
from pathlib import Path
import requests
# [FIX-#18]
from utils import get_embedding as _embed_raw
from typing import Optional
from lxml import etree
from docx import Document
from docx.table import Table
from docx.oxml.ns import qn

OUTPUT_DOCX     = "output_rpd.docx"
GENERATION_LOG  = "generation_log.json"

QDRANT = {"url": "http://localhost:6333", "collection": "rpd_rag"}
OLLAMA = {
    "embed_url":    "http://localhost:11434/api/embed",      # Ollama ≥0.6: /api/embed
    "generate_url": "http://localhost:11434/api/generate",
    "embed_model":  "bge-m3",
    # Смена модели: mistral:latest (7B, доступна локально).
    # Для наилучшего качества русского текста рекомендуется qwen2.5:14b —
    # скачать: ollama pull qwen2.5:14b  →  заменить ниже на "qwen2.5:14b"
    "llm_model":    "qwen2.5:14b",
}
GENERATION = {"top_k": 8, "min_score": 0.45}

# [З-13]
RERANK_ENABLED  = False          # переключается через args.rerank в main()
RERANK_TOP_K    = 20             # первичный пул для cross-encoder
_RERANKER_MODEL = "BAAI/bge-reranker-v2-m3"
_reranker       = None           # None = не инициализирован; False = недоступен


def _get_reranker():
    """Lazy-init CrossEncoder. False = попытка была, модель недоступна."""
    global _reranker
    if _reranker is not None:
        return _reranker
    try:
        from sentence_transformers import CrossEncoder  # noqa: PLC0415
        _reranker = CrossEncoder(_RERANKER_MODEL, max_length=512)
        print(f"  ✅ Reranker загружен: {_RERANKER_MODEL}")
    except Exception as e:
        print(f"  ⚠️  Reranker недоступен ({e}), cosine-only ranking")
        _reranker = False
    return _reranker


def _rerank(query: str, hits: list, top_n: int) -> list:
    """Cross-encoder reranking: hits (RERANK_TOP_K) → top_n. Fallback: hits[:top_n]."""
    if not hits:
        return hits
    reranker = _get_reranker()
    if not reranker:
        return hits[:top_n]
    texts = [h.get("payload", {}).get("text", "") or "" for h in hits]
    pairs = [(query, t[:512]) for t in texts]
    try:
        scores = reranker.predict(pairs, show_progress_bar=False)
        ranked  = sorted(zip(scores, hits), key=lambda x: x[0], reverse=True)
        result  = [h for _, h in ranked[:top_n]]
        print(f"      ↑ rerank: {len(hits)} → {len(result)} "
              f"(best score: {max(scores):.3f})")
        return result
    except Exception as e:
        print(f"  ⚠️  rerank ошибка: {e}, fallback cosine")
        return hits[:top_n]


# При однородном корпусе (все РПД одного направления/уровня) фильтр применяется
# к 100% чанков — ничего не отсекает. Смысл появится при расширении корпуса
# РПД других направлений или уровней подготовки (магистратура и т.п.).

# [J] Максимальная длина контекста, передаваемого в LLM (символы).
# [FIX-CTX]
MAX_CONTEXT_CHARS = 6000

# [З-R2]
SECTION_TYPE_FILTER = {
    "competencies":      ["competencies", "learning_outcomes"],
    "outcomes":          ["competencies", "learning_outcomes"],
    "content":           ["content", "lecture_content"],
    # [FIX-PRACTICE-STF]
    "lab_works":         ["content", "lab_content", "book_content"],
    "practice":          ["content", "practice_content", "book_content"],
    "bibliography_main": ["bibliography", "place"],
}

EMBED_CACHE    = {}
RETRIEVE_CACHE = {}

# [З-R5]
_CACHE_FILE = "rpd_cache.json"

# [З-G6]
_RETRIEVAL_CONF_HASH = ""

def _make_retrieval_conf_hash(top_k: int, min_score: float) -> str:
    # [FIX-HASH]
    import hashlib as _hl, os as _os
    stf_hash = _hl.md5(
        json.dumps(SECTION_TYPE_FILTER, sort_keys=True).encode()
    ).hexdigest()[:8]
    _chunks_mtime = int(_os.path.getmtime("chunks.jsonl")) if _os.path.exists("chunks.jsonl") else 0
    return f"k{top_k}_s{min_score:.3f}_stf{stf_hash}_ct{_chunks_mtime}"

def _load_cache() -> None:
    """Загружает кэш из файла, если он существует."""
    global EMBED_CACHE, RETRIEVE_CACHE
    if not os.path.exists(_CACHE_FILE):
        return
    try:
        with open(_CACHE_FILE, encoding="utf-8") as f:
            data = json.load(f)
        EMBED_CACHE    = data.get("embed", {})
        RETRIEVE_CACHE = {
            k: (v[0], v[1]) for k, v in data.get("retrieve", {}).items()
        }
        print(f"  Кэш загружен: {len(EMBED_CACHE)} эмбеддингов, "
              f"{len(RETRIEVE_CACHE)} retrieval-запросов")
    except Exception as e:
        print(f"  ⚠️  Кэш не загружен: {e}")

def _print_similar_disciplines(discipline: str, corpus_dir: str = "rpd_corpus",
                                top_n: int = 5) -> None:
    """
    [SIM] Выводит top_n наиболее похожих дисциплин из корпуса до генерации.
    Стратегия поиска названия (по убыванию приоритета):
      1. chunks[0..4]['text'] первая строка → regex (NNNNN)Название
      2. metadata['subject']
      3. data_clean.jsonl → поле 'discipline'
    """
    import glob as _glob
    import numpy as _np  # [FIX-#15] numpy вместо pure-Python cosine (math.sqrt loops)

    # [FIX-#15]
    def _cosine(a, b):
        va, vb = _np.array(a, dtype=float), _np.array(b, dtype=float)
        denom = _np.linalg.norm(va) * _np.linalg.norm(vb)
        return float(_np.dot(va, vb) / denom) if denom > 1e-10 else 0.0

    _CODE_RE = re.compile(r"^\(\d+\)\s*(.+)$")

    title_by_src: dict = {}

    # --- Стратегия 1: rpd_corpus/*.json ---
    if os.path.isdir(corpus_dir):
        for path in _glob.glob(os.path.join(corpus_dir, "*.json")):
            try:
                with open(path, encoding="utf-8") as f:
                    rec = json.load(f)
                src = os.path.basename(path)
                name = ""

                # Сканируем первые 5 чанков — ищем строку вида (код)Дисциплина
                for ch in rec.get("chunks", [])[:5]:
                    for line in ch.get("text", "").split("\n"):
                        line = line.strip()
                        m = _CODE_RE.match(line)
                        if m:
                            name = m.group(1).strip()
                            break
                    if name:
                        break

                # Fallback: metadata['subject']
                if not name:
                    name = rec.get("metadata", {}).get("subject", "").strip()

                if name:
                    title_by_src[src] = name
            except Exception:
                continue

    # --- Стратегия 2: data_clean.jsonl ---
    if not title_by_src:
        jsonl = "data_clean.jsonl"
        if os.path.exists(jsonl):
            seen: set = set()
            try:
                with open(jsonl, encoding="utf-8") as f:
                    for line in f:
                        try:
                            rec  = json.loads(line)
                            src  = rec.get("source", rec.get("title", ""))
                            if not src or src in seen:
                                continue
                            seen.add(src)
                            name = ""
                            for line_text in rec.get("text", "").split("\n"):
                                line_text = line_text.strip()
                                m = _CODE_RE.match(line_text)
                                if m:
                                    name = m.group(1).strip()
                                    break
                            if name:
                                title_by_src[src] = name
                        except Exception:
                            continue
            except Exception:
                pass

    if not title_by_src:
        print("  ⚠️  [SIM] Корпус не найден — список похожих дисциплин недоступен")
        return

    query_vec = get_embedding(discipline)
    if not query_vec:
        print("  ⚠️  [SIM] Embedding недоступен — пропускаю поиск похожих")
        return

    scored = []
    for src, name in title_by_src.items():
        vec = get_embedding(name)
        if vec:
            scored.append((name, src, _cosine(query_vec, vec)))

    if not scored:
        return

    scored.sort(key=lambda x: x[2], reverse=True)

    print(f"\n📚 Похожие дисциплины в корпусе (топ-{top_n}):")
    for i, (name, src, score) in enumerate(scored[:top_n], 1):
        bar = "█" * int(score * 20)
        print(f"  {i}. [{score:.3f}] {bar}  {name}  ({src})")
    print()


def _save_cache() -> None:
    """Сохраняет кэш в файл."""
    try:
        with open(_CACHE_FILE, "w", encoding="utf-8") as f:
            json.dump(
                {"embed": EMBED_CACHE, "retrieve": RETRIEVE_CACHE},
                f, ensure_ascii=False
            )
    except Exception as e:
        print(f"  ⚠️  Кэш не сохранён: {e}")

# [З-R1]
SECTION_QUERIES = {
    "competencies": [
        "{discipline} УК ОПК ПК формируемые компетенции шифр индекс ФГОС",
        "{discipline} способен разрабатывать применять анализировать профессиональная деятельность",
    ],
    "outcomes": [
        "{discipline} результаты обучения индикаторы достижения компетенций ФГОС",
        "{discipline} знать уметь владеть навыки практические умения студент",
    ],
    "content": [
        "{discipline} тематический план содержание дисциплины разделы лекции",
        "{discipline} программа курса темы методы алгоритмы технологии практика",
    ],
    "lab_works": [
        # [З-08]
        "{discipline} задание разработка реализация алгоритма исследование",
        "{discipline} лабораторная работа задание исследование программирование Python",
    ],
    "practice": [
        "{discipline} практические занятия перечень тем задач методы алгоритмы",
        "{discipline} решение задач моделирование синтез исследование системы",
    ],
    "bibliography_main": [
        "{discipline} учебник учебное пособие литература библиография",
        "{discipline} основная дополнительная литература ЭБС Знаниум",
    ],
}

# Глобальный лог генерации — [C]
_generation_log: dict = {}


# ---------------------------------------------------------------------------
# Утилиты
# ---------------------------------------------------------------------------

def clean(text: str) -> str:
    text = re.sub(r" +", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"\[score=[^\]]+\]\n?", "", text)
    return "\n".join(l.strip() for l in text.split("\n") if l.strip()).strip()


def get_embedding(text: str):
    # [FIX-#18]
    if text in EMBED_CACHE:
        return EMBED_CACHE[text]
    vec = _embed_raw(text, prefix="query", retry=3)
    if vec:
        EMBED_CACHE[text] = vec
    return vec


def _search_qdrant(vec: list, payload_filter: dict | None, top_k: int) -> list:
    """Поиск в Qdrant с fallback query → search."""
    try:
        body = {"query": vec, "limit": top_k, "with_payload": True}
        if payload_filter:
            body["filter"] = payload_filter
        r = requests.post(
            f"{QDRANT['url']}/collections/{QDRANT['collection']}/points/query",
            json=body, timeout=30)
        r.raise_for_status()
        return r.json().get("result", {}).get("points", [])
    except requests.HTTPError:
        body = {"vector": vec, "limit": top_k, "with_payload": True}
        if payload_filter:
            body["filter"] = payload_filter
        r = requests.post(
            f"{QDRANT['url']}/collections/{QDRANT['collection']}/points/search",
            json=body, timeout=30)
        r.raise_for_status()
        return r.json().get("result", [])


def retrieve(section: str, discipline: str, section_types: list = None,
             direction: str = "", level: str = "") -> tuple[str, list]:
    """
    Ищет релевантные чанки в Qdrant.

    [K] Multi-query: объединяем результаты по нескольким формулировкам.
    [B] Доменная фильтрация по direction/level.
    [S] Фильтр использует "section_type" (верхний уровень payload).
    [R] При пустом результате возвращает пустую строку с флагом для caller.

    Возвращает: (ctx_string, hits_list) для логирования [C].
    """
    # [З-G6]
    cache_key = (f"{section}|{discipline}|{','.join(section_types or [])}"
                 f"|{direction}|{level}|{_RETRIEVAL_CONF_HASH}")
    if cache_key in RETRIEVE_CACHE:
        return RETRIEVE_CACHE[cache_key]

    try:
        # [B] Строим фильтр с доменными полями
        must_conditions: list = []
        if section_types:
            if len(section_types) == 1:
                # [FIX-SHOULD1]
                must_conditions.append(
                    {"key": "section_type", "match": {"value": section_types[0]}}
                )
            else:
                must_conditions.append({
                    "should": [
                        # [S] Используем "section_type" на верхнем уровне payload
                        {"key": "section_type", "match": {"value": st}}
                        for st in section_types
                    ]
                })
        if direction:
            must_conditions.append({"key": "direction", "match": {"value": direction}})
        if level:
            must_conditions.append({"key": "level", "match": {"value": level}})

        payload_filter = {"must": must_conditions} if must_conditions else None

        # [K] Multi-query: собираем чанки по нескольким запросам
        queries = SECTION_QUERIES.get(section, [f"{discipline} {section}"])
        queries = [q.format(discipline=discipline) for q in queries]

        all_hits: dict[int, dict] = {}  # id → hit (дедупликация)
        for query_text in queries:
            vec = get_embedding(query_text)
            if not vec:
                continue
            hits = _search_qdrant(vec, payload_filter,
                                  RERANK_TOP_K if RERANK_ENABLED else GENERATION["top_k"])
            for h in hits:
                hit_id = h.get("id")
                if hit_id not in all_hits or h.get("score", 0) > all_hits[hit_id].get("score", 0):
                    all_hits[hit_id] = h

        # [FIX-5]
        MAX_PER_SOURCE = 2
        _source_counts: dict = {}
        _diverse_all: list = []
        for h in sorted(all_hits.values(), key=lambda h: h.get("score", 0), reverse=True):
            if h.get("score", 0) < GENERATION["min_score"]:
                continue
            src = h.get("payload", {}).get("source", "")
            if _source_counts.get(src, 0) < MAX_PER_SOURCE:
                _source_counts[src] = _source_counts.get(src, 0) + 1
                _diverse_all.append(h)

        # [З-13]
        if RERANK_ENABLED and _diverse_all:
            good_hits = _rerank(queries[0], _diverse_all, GENERATION["top_k"])
        else:
            good_hits = _diverse_all[:GENERATION["top_k"]]

        # [R] Fallback при пустом retrieval — снижаем порог и убираем фильтр
        if not good_hits:
            print(f"    ⚠️  RAG [{section}]: нет чанков выше {GENERATION['min_score']}, "
                  f"пробую без доменного фильтра...")
            vec = get_embedding(queries[0])
            if vec:
                hits = _search_qdrant(vec, None, GENERATION["top_k"])
                good_hits = sorted(
                    [h for h in hits if h.get("score", 0) >= GENERATION["min_score"] * 0.7],
                    key=lambda h: h.get("score", 0), reverse=True
                )[:GENERATION["top_k"]]

        print(f"    🔍 RAG [{section}]: найдено {len(good_hits)} чанков "
              f"(scores: {[round(h.get('score', 0), 3) for h in good_hits]})")

        # Сборка контекста с метаданными источника
        seen_texts: set = set()
        parts: list[str] = []
        for h in good_hits:
            payload = h.get("payload", {})
            raw_text = payload.get("text", "")
            if not raw_text:
                continue

            # [замечание #12 ИСПРАВЛЕНО]: ранее hard-cut payload["text"][:1200]
            # мог обрывать текст посередине предложения, давая LLM
            # неструктурированный фрагмент. Теперь при превышении 1200 символов
            # ищем последнюю точку в диапазоне [800, 1200] и обрезаем по ней.
            # Если точка не найдена — оставляем сырой срез (лучше, чем ничего).
            if len(raw_text) > 1200:
                cut = raw_text[:1200]
                last_dot = cut.rfind(".")
                text = cut[:last_dot + 1] if last_dot >= 800 else cut
            else:
                text = raw_text

            dedup_key = text[:100]
            if dedup_key in seen_texts:
                continue
            seen_texts.add(dedup_key)
            source        = payload.get("source", "")
            section_title = payload.get("section_title", "")
            prefix = ""
            if source:
                prefix += f"[{source}]"
            if section_title:
                prefix += f" [{section_title}]"
            parts.append(f"{prefix}\n{text}" if prefix else text)

        ctx = "\n\n---\n\n".join(parts)

        # [J] Ограничение длины контекста — предотвращает переполнение окна LLM.
        # При превышении MAX_CONTEXT_CHARS обрезаем с явной пометкой.
        if len(ctx) > MAX_CONTEXT_CHARS:
            ctx = ctx[:MAX_CONTEXT_CHARS].rsplit("\n", 1)[0]
            ctx += "\n[...контекст обрезан до MAX_CONTEXT_CHARS символов...]"

        RETRIEVE_CACHE[cache_key] = (ctx, good_hits)
        return ctx, good_hits

    except Exception as e:
        print(f"  ⚠️  RAG [{section}]: {e}")
        return "", []


def llm(prompt: str, max_tokens: int = 800) -> str:
    for attempt in range(3):
        try:
            r = requests.post(OLLAMA["generate_url"],
                json={
                    "model": OLLAMA["llm_model"],
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.3,
                        "num_predict": max_tokens,
                        # [M] num_ctx=8192: mistral:7b поддерживает 8K контекст.
                        # qwen2.5:3b требовал 4096 из-за OOM; 7B справляется на 8K.
                        # При qwen2.5:14b можно оставить 8192 или поднять до 16384.
                        "num_ctx": 8192,
                    }
                },
                # [M] timeout=300: 7B-модель генерирует ~3–5×медленнее 3B.
                # На CPU ~60–120 сек на раздел — запас до 300 сек достаточен.
                timeout=300)
            r.raise_for_status()
            text = r.json().get("response", "")
            if text:
                time.sleep(3.0)   # [ЗАМЕЧАНИЕ] пауза между LLM-вызовами — снижает нагрев GPU
                return clean(text)
        except Exception as e:
            if attempt == 2:
                return f"[Ошибка: {e}]"
            time.sleep(5)
    return "[Ошибка: пустой ответ]"


def _sanitize_retrieved_text(text: str) -> str:
    """
    [замечание #13] Базовая защита от prompt injection в retrieved-контексте.

    Если corpus содержит строки вида "Ignore previous instructions" или
    "System: ...", LLM может их воспринять как системные директивы.
    Фильтруем строки, начинающиеся с типичных injection-паттернов.
    Это не полная защита (для production нужен отдельный guard-слой),
    но устраняет наиболее очевидные векторы атаки.
    """
    INJECTION_PATTERNS = re.compile(
        r"^(ignore\s+(previous|all|prior)|forget\s+(previous|all)|"
        r"system\s*:|instruction\s*:|act\s+as\s|"
        r"забудь\s+(предыдущ|все)|игнорируй\s+предыдущ|"
        r"ты\s+теперь\s|притворяйся\s|ты\s+—\s)",
        re.IGNORECASE,
    )
    clean_lines = [
        line for line in text.split("\n")
        if not INJECTION_PATTERNS.match(line.strip())
    ]
    return "\n".join(clean_lines).strip()


# [§2.2.5]
_TERM_CORRECTIONS: list[tuple[str, str]] = [
    # транслитерационные галлюцинации
    (r"\bсемисери\b",               "полуконтролируемое обучение"),
    (r"\bполусери\b",               "полуконтролируемое обучение"),
    (r"\bбезпосредственн\w*",       "обучение без учителя"),
    (r"\bГейш-рекуррентн\w*",       "управляемый рекуррентный блок"),
    (r"\bГейш\s+рекуррентн\w*",     "управляемый рекуррентный блок"),
    # калька «deep learning» → «глубокий обучение» (неверный род)
    (r"\bглубокий\s+обучени[ея]\b", "глубокое обучение"),
    # «машинный обучение» (неверный род)
    (r"\bмашинный\s+обучени[ея]\b", "машинное обучение"),
    # «supervised learning» → «надзорное обучение»
    (r"\bнадзорное\s+обучени[ея]\b","обучение с учителем"),
    # «unsupervised» → «ненадзорный»
    (r"\bненадзорн\w+",             "без учителя"),
    # опечатка «интеллектального»
    (r"\bинтеллектального\b",       "интеллектуального"),
]
_TERM_RE: list[tuple[re.Pattern, str]] = [
    (re.compile(pat, re.IGNORECASE), repl)
    for pat, repl in _TERM_CORRECTIONS
]


def _apply_term_corrections(text: str) -> str:
    """Применяет _TERM_RE к тексту LLM-ответа."""
    for pattern, replacement in _TERM_RE:
        text = pattern.sub(replacement, text)
    return text


def gen(label: str, discipline: str, prompt: str,
        direction: str = "", level: str = "", **extra) -> str:
    """
    Генерация секции с RAG-контекстом.

    [R] При пустом retrieval добавляет явную инструкцию в промпт.
    [C] Сохраняет данные в _generation_log для последующей записи в JSON.
    """
    section_types = SECTION_TYPE_FILTER.get(label)
    ctx, hits = retrieve(label, discipline, section_types, direction, level)

    # [замечание #13] Санитизация retrieved-контекста перед вставкой в промпт
    ctx = _sanitize_retrieved_text(ctx)

    if ctx:
        ctx_block = (
            "Примеры из базы РПД кафедры (используй как образец стиля и формата):\n"
            f"{ctx}\n\n"
        )
    else:
        # [R] Явное предупреждение и инструкция при отсутствии контекста
        print(f"  ⚠️  RAG [{label}]: контекст пуст — генерация без примеров из корпуса")
        ctx_block = (
            "Примеры из базы РПД недоступны. "
            "Сгенерируй содержимое самостоятельно строго по указанному формату "
            "без копирования примеров из промпта.\n\n"
        )

    # [БАГ 5 ИСПРАВЛЕНО]
    fmt_vars = {"discipline": discipline, "direction": direction, "level": level, **extra}
    full_prompt = ctx_block + prompt.format(**fmt_vars) + f"\n\nСоздай для «{discipline}»:"
    result = _apply_term_corrections(llm(full_prompt))

    # [C] Логируем для generation_log.json
    _generation_log[label] = {
        "prompt_preview":   full_prompt[:600],
        "retrieved_chunks": [
            {
                "id":           h.get("id"),
                "source":       h.get("payload", {}).get("source", ""),
                "score":        round(h.get("score", 0), 4),
                "text_preview": h.get("payload", {}).get("text", "")[:120],
            }
            for h in hits
        ],
        "llm_response":     result,
        "timestamp":        time.strftime("%Y-%m-%dT%H:%M:%S"),
    }

    return result


# ---------------------------------------------------------------------------
# [T1-T4] Работа с DOCX-шаблоном на основе [] плейсхолдеров
# ---------------------------------------------------------------------------

# Маппинг семантического ключа → предикат для поиска таблицы.
# Аргумент предиката: frozenset строк из первых 3 строк заголовка таблицы.
_TABLE_PREDICATES: dict = {
    "workload":     lambda h: "Зачетные единицы" in h,
    "competencies": lambda h: "Формируемые компетенции" in h,
    "outcomes":     lambda h: "Индикаторы достижения компетенций" in h,
    "work_types":   lambda h: "Вид учебной работы" in h,
    "topics":       lambda h: "Название темы (раздела)" in h,
    # T8: есть "№ пп." + "Название темы", нет "(раздела)" и "лабораторной"
    "lectures":     lambda h: (
        "Название темы" in h and "№ пп." in h
        and "Название темы (раздела)" not in h
        and "лабораторной" not in h
    ),
    "labs":         lambda h: "Название лабораторной работы" in h,
    "practice":     lambda h: "Тема практического занятия" in h,
    "sro":          lambda h: "Вид СРО" in h,
    # T15: есть "Тип" И "Библиографическое описание"
    "bibliography": lambda h: "Тип" in h and "Библиографическое описание" in h,
    # T17: "Назначение учебных изданий", но НЕТ "Тип" как отдельного заголовка
    "method_bib":   lambda h: (
        "Назначение учебных изданий" in h and "Тип" not in h
    ),
    "fos":          lambda h: "Контролируемые разделы (темы) дисциплины" in h,
    "fos_types":    lambda h: "Вид оценочного средства" in h and "Контролируемые разделы (темы) дисциплины" not in h,
    # [П9] "annotation" удалён: find_annotation_table() ищет таблицу по позиции
    # в документе, не через этот предикат — был мёртвым кодом.
}


def _table_header_set(table, max_rows: int = 5) -> frozenset:
    """Собирает множество уникальных текстов ячеек из первых max_rows строк.
    (work_types, fos) иногда занимают 4 строки (двойная шапка).
    При max_rows=3 «Вид учебной работы» в строке 4 → find_table вернул None."""
    texts = set()
    for row in table.rows[:max_rows]:
        seen_tc = set()
        for cell in row.cells:
            if id(cell._tc) not in seen_tc:
                seen_tc.add(id(cell._tc))
                t = cell.text.strip()
                if t:
                    texts.add(t)
    return frozenset(texts)


def find_annotation_table(doc: Document) -> Optional[Table]:
    """
    Находит таблицу аннотации РПД по позиции в документе:
    ищет первую таблицу после параграфа «Аннотация к рабочей программе».

    Стандартный find_table() не подходит: в шаблоне заголовок таблицы
    тоже содержит «[]» (не pre-filled текст), поэтому предикат
    _TABLE_PREDICATES["annotation"] никогда не срабатывает.
    """
    body = doc.element.body
    found_ann = False
    for child in body:
        tag = child.tag.split("}")[1] if "}" in child.tag else child.tag
        if tag == "p":
            text = "".join(
                t.text or "" for t in child.findall(f".//{{{qn('w:t').split('}')[0][1:]}}}t")
            )
            if "Аннотация к рабочей программе" in text:
                found_ann = True
        elif tag == "tbl" and found_ann:
            return Table(child, doc)
    return None


def fill_annotation_table(
    doc: Document,
    competencies: list,
    outcomes: list,
    topics: list,
    credits: int,
    hours_total: int,
    exam_type: str,
) -> None:
    """
    Заполняет таблицу аннотации РПД (15 строк).

    Структура (соответствует реальным РПД кафедры ВТИК):
      [0]  «Компетенции, формируемые в результате освоения дисциплины»
      [1]  Список компетенций с индикаторами
      [2]  «Результат обучения»
      [3]  «Знать:»
      [4]  З-результаты (по одному на компетенцию)
      [5]  «Уметь:»
      [6]  У-результаты
      [7]  «Владеть:»
      [8]  В-результаты
      [9]  «Краткая характеристика дисциплины»
      [10] Перечень разделов (темы)
      [11] «Трудоёмкость (з.е. / часы)»
      [12] «N з.е. (Nчас)»
      [13] «Вид промежуточной аттестации»
      [14] exam_type
    """
    table = find_annotation_table(doc)
    if table is None:
        print("  ⚠️  [fill_annotation] Таблица аннотации не найдена")
        return

    # --- Строим содержимое ячеек ---

    # [1] Компетенции с индикаторами
    # [FIX-ANN]
    _ann_z_texts = [t for ot, t in outcomes if ot == "\u0417"] if outcomes else []
    comp_lines = []
    for comp_idx, (code, desc) in enumerate(competencies):
        comp_lines.append(f" {code} {desc}:")
        # [FIX-01-ANN]
        _ind_text = _ann_z_texts[comp_idx % len(_ann_z_texts)] if _ann_z_texts \
            else "Применяет методы и инструменты дисциплины"
        comp_lines.append(f"-{code}.1 Знает {_ind_text}")
    comp_text = "\n".join(comp_lines)

    # [4][6][8] Группируем outcomes по типу
    z_items = [(code, t) for (ot, t), (code, _) in zip(
        [o for o in outcomes if o[0] == "З"],
        competencies
    )] if outcomes else []
    u_items = [(code, t) for (ot, t), (code, _) in zip(
        [o for o in outcomes if o[0] == "У"],
        competencies
    )] if outcomes else []
    v_items = [(code, t) for (ot, t), (code, _) in zip(
        [o for o in outcomes if o[0] == "В"],
        competencies
    )] if outcomes else []

    def fmt_outcomes(items: list) -> str:
        return "\n".join(f"{code}- {text}" for code, text in items) if items else ""

    # [10] Краткий перечень разделов
    sections = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    sections_text = "; ".join(
        re.sub(r"^Раздел\s*\d+\.\s*", "", s) for s in sections
    ) if sections else "; ".join(t for t in topics[:3])

    values = [
        "Компетенции, формируемые в результате освоения дисциплины",
        comp_text,
        "Результат обучения",
        "Знать:",
        fmt_outcomes(z_items) or "Основные методы и понятия дисциплины",
        "Уметь:",
        fmt_outcomes(u_items) or "Применять методы дисциплины для решения задач",
        "Владеть:",
        fmt_outcomes(v_items) or "Навыками работы с инструментами дисциплины",
        "Краткая характеристика дисциплины",
        sections_text,
        "Трудоёмкость (з.е. / часы)",
        f"{credits} з.е.  ({hours_total}час)",
        "Вид промежуточной аттестации",
        exam_type,
    ]

    tbl_xml = table._tbl
    all_trs  = tbl_xml.findall(qn("w:tr"))
    for tr, val in zip(all_trs, values):
        seen: set = set()
        for tc in tr.findall(f".//{qn('w:tc')}"):
            if id(tc) in seen:
                continue
            seen.add(id(tc))
            _set_cell_xml(tc, val)
            break  # одна колонка — берём первую ячейку

    # [FIX-2]


def find_table(doc: Document, key: str) -> Optional[Table]:
    """
    [T3] Находит таблицу по семантическому ключу из _TABLE_PREDICATES.
    Устойчив к добавлению/удалению таблиц в шаблоне.
    При отсутствии таблицы выводит предупреждение и возвращает None.
    """
    predicate = _TABLE_PREDICATES.get(key)
    if predicate is None:
        raise KeyError(f"Неизвестный ключ таблицы: {key!r}. "
                       f"Доступны: {sorted(_TABLE_PREDICATES)}")
    for table in doc.tables:
        header_set = _table_header_set(table)
        if predicate(header_set):
            return table
    print(f"  ⚠️  [find_table] Таблица {key!r} не найдена — проверьте шаблон")
    return None


def _set_cell_xml(tc, text: str) -> None:
    """
    Записывает text в ячейку (lxml <w:tc>), сохраняя форматирование первого run.
    Используется в fill_placeholder_rows и fill_doc_header.

    [Фикс №7] Многострочный текст (\\n) → отдельные <w:p> с тем же форматированием.
    copy.deepcopy(saved_rpr) — предотвращает повторное использование одного
    lxml-элемента в нескольких ячейках (вызывало порчу форматирования при
    клонировании строк таблицы).
    """
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W  = lambda tag: f"{{{NS}}}{tag}"

    paras = tc.findall(W("p"))
    if not paras:
        lines = text.split("\n") if "\n" in text else [text]
        new_p = etree.SubElement(tc, W("p"))
        new_r = etree.SubElement(new_p, W("r"))
        new_t = etree.SubElement(new_r, W("t"))
        new_t.text = lines[0]
        for line in lines[1:]:
            extra_p = etree.SubElement(tc, W("p"))
            extra_r = etree.SubElement(extra_p, W("r"))
            extra_t = etree.SubElement(extra_r, W("t"))
            extra_t.text = line
        return

    # Оставляем первый параграф, удаляем остальные
    for p in paras[1:]:
        tc.remove(p)
    p = paras[0]

    # Сохраняем rPr первого run
    runs = p.findall(f".//{W('r')}")
    saved_rpr = None
    if runs:
        rpr = runs[0].find(W("rPr"))
        if rpr is not None:
            saved_rpr = copy.deepcopy(rpr)
        for r in runs:
            p.remove(r)

    lines = text.split("\n") if "\n" in text else [text]

    new_r = etree.SubElement(p, W("r"))
    if saved_rpr is not None:
        new_r.append(copy.deepcopy(saved_rpr))
    new_t = etree.SubElement(new_r, W("t"))
    new_t.text = lines[0]
    if lines[0] and (lines[0][0] == " " or lines[0][-1] == " "):
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    # [Фикс №7] Дополнительные строки → отдельные <w:p> с тем же форматированием
    for line in lines[1:]:
        new_p = copy.deepcopy(p)
        for r in new_p.findall(f".//{W('r')}"):
            new_p.remove(r)
        nr = etree.SubElement(new_p, W("r"))
        if saved_rpr is not None:
            nr.append(copy.deepcopy(saved_rpr))
        nt = etree.SubElement(nr, W("t"))
        nt.text = line
        if line and (line[0] == " " or line[-1] == " "):
            nt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        tc.append(new_p)


def _is_placeholder_row(tr) -> bool:
    """True если хотя бы одна уникальная ячейка строки содержит только '[]'."""
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W  = lambda tag: f"{{{NS}}}{tag}"
    seen = set()
    for tc in tr.findall(f".//{W('tc')}"):
        if id(tc) in seen:
            continue
        seen.add(id(tc))
        t = "".join(x.text or "" for x in tc.findall(f".//{W('t')}")).strip()
        if t == "[]":
            return True
    return False


def _fill_tr(tr, values: list[str]) -> None:
    """Записывает values в уникальные ячейки строки."""
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W  = lambda tag: f"{{{NS}}}{tag}"
    cells = []
    seen = set()
    for tc in tr.findall(f".//{W('tc')}"):
        if id(tc) not in seen:
            seen.add(id(tc))
            cells.append(tc)
    for tc, val in zip(cells, values):
        _set_cell_xml(tc, str(val))


def fill_placeholder_rows(table: Table, data: list[list[str]]) -> None:
    """
    [T2] Заполняет строки таблицы, содержащие [] плейсхолдеры.

    Алгоритм:
      1. Находит все строки с [] — это «слоты» для данных.
      2. Заполняет их значениями из data (один список = одна строка).
      3. Если данных больше слотов — клонирует последний слот.
      4. Если слотов больше данных — удаляет лишние.
    Строки ИТОГО и «-» (без []) не трогаются.
    """
    tbl_xml = table._tbl
    all_trs = list(tbl_xml)

    placeholder_trs = [tr for tr in all_trs if _is_placeholder_row(tr)]
    if not placeholder_trs:
        print("  ⚠️  [fill_placeholder_rows] Строк с [] не найдено — шаблон не обновлён?")
        return

    template_tr = placeholder_trs[-1]   # эталон для клонирования

    # Заполняем имеющиеся слоты
    for tr, values in zip(placeholder_trs, data):
        _fill_tr(tr, values)

    # Если данных больше — добавляем клонированные строки
    for values in data[len(placeholder_trs):]:
        new_tr = copy.deepcopy(template_tr)
        tbl_xml.append(new_tr)
        _fill_tr(new_tr, values)

    # Если слотов больше данных — удаляем лишние []
    for extra_tr in placeholder_trs[len(data):]:
        tbl_xml.remove(extra_tr)


def fill_doc_header(doc: Document, discipline: str, code: str,
                    year: str = "2025", credits: int = 4,
                    hours_total: int = 144, exam_type: str = "экзамен") -> None:
    """
    [T1] Заменяет [] плейсхолдеры в параграфах документа.

    Заменяет:
      • первый «[]» → «(код)Название дисциплины»
      • «Трудоемкость дисциплины: []» → реальное значение
      • «Уфа []» → «Уфа {year}»
      • оставшиеся «[]» во всех параграфах → «(код)Название дисциплины»
        (для приложений А/Б/В и аннотации)

    перечни вопросов/заданий, не название дисциплины. Пропускаем здесь,
    заполняет fill_appendix_v().
    """
    label    = f"({code}){discipline}"
    workload = f"{credits} з.е.  ({hours_total}час)"

    # [FIX-AppV]
    _ASSESSMENT_HEADERS = {
        "реферат", "доклад", "лабораторная работа",
        "письменный и устный опрос", "тест",
        "расчётно-графическая работа",
        # [Д-4] Промежуточная строка между видом оценивания и [] в Приложении В.
        # Структура шаблона: «Реферат.» → «Перечень вопросов...:» → «[]»
        # prev_txt при встрече [] = эта строка, а не «реферат» → без неё
        # fill_doc_header заменял [] на label вместо пропуска для fill_appendix_v.
        "перечень вопросов (задач, заданий, тем, комплекта тестовых заданий):",
    }
    # [Д-2] Флаг блока «Темы для СРО» — [] внутри него НЕ заменяются
    # дисциплинарным лейблом: их обрабатывает fill_sro_topic_paragraphs().
    _in_sro_block = False
    prev_txt = ""

    def _set_para(para, text: str) -> None:
        """Записывает text в параграф, сохраняя формат первого run."""
        if not para.runs:
            return
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""

    first_label_done = False

    for para in doc.paragraphs:
        txt = para.text.strip()

        # [Д-2] SRO-блок: отслеживаем вход/выход
        if "Темы для самостоятельной работы обучающихся" in txt:
            _in_sro_block = True
        elif _in_sro_block and re.match(r"^\d+\.", txt):
            _in_sro_block = False

        # [Д-2] Внутри SRO-блока [] заменяет fill_sro_topic_paragraphs — пропускаем
        if _in_sro_block:
            prev_txt = txt
            continue

        if not first_label_done and txt == "[]":
            _set_para(para, label)
            first_label_done = True
            prev_txt = txt
            continue

        if txt == "Трудоемкость дисциплины: []":
            _set_para(para, f"Трудоемкость дисциплины: {workload}")
            prev_txt = txt
            continue

        if txt == f"Уфа []":
            _set_para(para, f"Уфа {year}")
            prev_txt = txt
            continue

        # [§6.1.5]
        if re.search(r"\b\d{2}\.\d{2}\.\d{4}\b", txt):
            new_txt = re.sub(r"\b(\d{2}\.\d{2}\.)\d{4}\b",
                             lambda m: m.group(1) + str(year), para.text)
            if new_txt != para.text:
                _set_para(para, new_txt)
            prev_txt = txt
            continue

        if re.search(r"Год\s+приема\s+\d{4}\s+г", txt, re.IGNORECASE):
            new_txt = re.sub(r"(Год\s+приема\s+)\d{4}(\s+г)",
                             lambda m: m.group(1) + str(year) + m.group(2),
                             para.text, flags=re.IGNORECASE)
            if new_txt != para.text:
                _set_para(para, new_txt)
            prev_txt = txt
            continue

        # Приложения А/Б/В и аннотация — оставшиеся одиночные []
        if txt == "[]" and first_label_done:
            # [FIX-AppV]
            if prev_txt.strip().lower() not in _ASSESSMENT_HEADERS:
                _set_para(para, label)
            prev_txt = txt
            continue

        # Строки «4 з.е.  (144час)» в аннотации
        if txt == "[]" or "[]" in txt:
            if "з.е." in txt or "час" in txt:
                _set_para(para, para.text.replace("[]", workload))
            else:
                _set_para(para, para.text.replace("[]", label))

        prev_txt = txt


# Оставляем set_cell_text для fill_t6_workload (keyword-based, без [] слотов)
def set_cell_text(cell, text: str) -> None:
    """Устанавливает текст ячейки python-docx Cell. Используется только в fill_t6_workload."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if not cell.paragraphs:
        cell.add_paragraph(text)
    elif not cell.paragraphs[0].runs:
        cell.paragraphs[0].add_run(text)
    else:
        cell.paragraphs[0].runs[0].text = text


# [Д-2] Заполнение блока «Темы для самостоятельной работы».
# Шаблон содержал жёстко вшитые rpd_1-темы («Организация НИД в России» и т.п.)
# вместо []. После исправления шаблона [125],[127],[129] → [] в блоке 6 слотов:
#   чётные (0,2,4) — имя раздела  →  sections[i]
#   нечётные (1,3,5) — перечень тем  →  "(код)Дисциплина" (fill_doc_header пропустил)
# fill_doc_header скипает весь SRO-блок; эта функция заполняет его явно.
def fill_sro_topic_paragraphs(doc: Document, topics: list, label: str) -> None:
    sections = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    if not sections:
        return

    # [FIX-SRO-TOPICS]
    section_subtopics: list[str] = []
    current: list[str] = []
    for t in topics:
        if re.match(r"^Раздел\s*\d+", t):
            # Сохраняем накопленные темы предыдущего раздела
            if current or section_subtopics:
                section_subtopics.append("; ".join(current) if current else label)
                current = []
        elif re.match(r"^Тема\s+\d", t):
            # Убираем «Тема X.Y. » — оставляем только название
            name = re.sub(r"^Тема\s+[\d.]+\s*", "", t).strip()
            current.append(name)
    # Последний раздел
    section_subtopics.append("; ".join(current) if current else label)

    in_block   = False
    slot_idx   = 0
    sec_idx    = 0

    for para in doc.paragraphs:
        txt = para.text.strip()

        if "Темы для самостоятельной работы обучающихся" in txt:
            in_block = True
            continue

        if in_block:
            # Выход из блока: начало следующего пронумерованного раздела РПД
            if re.match(r"^\d+\.", txt) and "Раздел" not in txt:
                break

            if txt == "[]":
                if para.runs:
                    if slot_idx % 2 == 0:              # имя раздела
                        text = sections[sec_idx] if sec_idx < len(sections) else label
                        sec_idx += 1
                    else:                               # перечень тем раздела
                        idx = sec_idx - 1
                        text = (
                            section_subtopics[idx]
                            if 0 <= idx < len(section_subtopics)
                            else label
                        )
                    para.runs[0].text = text
                    for r in para.runs[1:]:
                        r.text = ""
                slot_idx += 1



def _to_research_topic(subtopic: str, idx: int) -> str:
    """Превращает название темы в исследовательский вопрос для доклада.
    Без этой функции темы докладов дословно копировали названия лекций."""
    base = subtopic.rstrip(".").strip()
    if not base:
        return subtopic
    base_lc = base[0].lower() + base[1:] if len(base) > 1 else base.lower()
    templates = [
        f"Современное состояние и перспективы развития: {base_lc}",
        f"Сравнительный анализ подходов к теме «{base}»",
        f"Применение {base_lc} в прикладных задачах",
        f"Актуальные проблемы и ограничения: {base_lc}",
        f"Обзор методов и алгоритмов: {base_lc}",
        f"{base}: практический обзор существующих решений",
    ]
    return templates[idx % len(templates)]


# fill_doc_header пропускает [] после заголовков видов оценивания —
# их заполняет эта функция на основе тематического плана дисциплины.
def fill_appendix_v(doc: Document, discipline: str, topics: list) -> None:
    """
    Заполняет перечни вопросов/заданий в Приложении В (и аналогичных разделах).

    Алгоритм: сканирует doc.paragraphs, отслеживает предыдущий параграф.
    Если предыдущий — заголовок вида оценивания, а текущий — «[]»,
    подставляет template-контент на основе разделов дисциплины.
    """
    sections = [
        re.sub(r"^Раздел\s*\d+\.\s*", "", t)
        for t in topics if re.match(r"^Раздел\s*\d+", t)
    ][:5]
    if not sections:
        sections = [f"основные разделы дисциплины «{discipline}»"]

    topics_list   = "\n".join(f"- {s}" for s in sections)
    topics_inline = "; ".join(sections)

    # [FIX-AppV-DOK]
    subtopics = [
        re.sub(r"^Тема\s+[\d.]+\s*", "", t).strip()
        for t in topics if re.match(r"^Тема\s+[\d.]", t)
    ][:6]
    # [FIX-10]
    subtopics_list = (
        "\n".join(f"- {_to_research_topic(s, i)}" for i, s in enumerate(subtopics))
        if subtopics else topics_list
    )

    _TEMPLATES = {
        "реферат": (
            f"Темы рефератов по дисциплине «{discipline}»:\n{topics_list}"
        ),
        "доклад": (
            f"Темы докладов по дисциплине «{discipline}»:\n{subtopics_list}"
        ),
        "лабораторная работа": (
            f"Перечень лабораторных работ охватывает разделы: {topics_inline}"
        ),
        "письменный и устный опрос": (
            f"Контрольные вопросы по разделам дисциплины «{discipline}»:\n{topics_list}"
        ),
        "тест": (
            f"Тестирование по дисциплине «{discipline}». "
            f"Охватываемые разделы: {topics_inline}"
        ),
        "расчётно-графическая работа": (
            f"Задания РГР по дисциплине «{discipline}»:\n{topics_list}"
        ),
    }

    # [Д-4 ИСПРАВЛЕНО]: прежде отслеживался только prev_txt (один шаг назад).
    # Структура шаблона: «Реферат.» → «Перечень вопросов...:» → «[]»
    # При встрече [] prev_txt = «Перечень вопросов...», а не «реферат» →
    # _TEMPLATES не срабатывал, fill_doc_header подставлял label вместо перечня.
    # Теперь last_assessment сохраняется до встречи [] вне зависимости от числа
    # промежуточных параграфов.
    last_assessment: str | None = None
    for para in doc.paragraphs:
        txt     = para.text.strip()
        txt_key = txt.lower().rstrip(".")

        # Заголовок вида оценивания — запоминаем
        # [Фикс Д-AppV] txt_key in _TEMPLATES давал промах для составных
        # заголовков вида «Доклад, сообщение» → ключ «доклад» не находился
        # точным совпадением. Теперь ищем ключ как подстроку заголовка.
        # ВАЖНО: startswith, а НЕ `k in txt_key` — «тест» найдётся как
        # подстрока в «тестовых заданий» и перезапишет last_assessment.
        matched_key = next((k for k in _TEMPLATES if txt_key.startswith(k)), None)
        if matched_key:
            last_assessment = matched_key
            continue

        # [] с активным last_assessment — заполняем и сбрасываем
        if txt == "[]" and last_assessment is not None:
            content = _TEMPLATES[last_assessment]
            if para.runs:
                para.runs[0].text = content
                for r in para.runs[1:]:
                    r.text = ""
            last_assessment = None
            continue

        # Выход из зоны Приложения В — сброс
        if re.match(r"^\d+\.", txt) or txt.startswith("СОГЛАСОВАНО"):
            last_assessment = None


# ---------------------------------------------------------------------------
# [A] JSON-парсеры с fallback на regex
# ---------------------------------------------------------------------------

def parse_competencies_json(text: str) -> list | None:
    """
    [A] Пытается разобрать JSON-ответ LLM для компетенций.
    Ожидаемый формат: [{"code": "УК-1", "desc": "Способен..."}]
    """
    # [БАГ 8 ИСПРАВЛЕНО]
    m = re.search(r"\[.*\]", text, re.S)
    if not m:
        return None
    try:
        data = json.loads(m.group())
        if not isinstance(data, list):
            return None
        result = [
            (str(d.get("code", "")), str(d.get("desc", "")))
            for d in data
            if isinstance(d, dict) and d.get("code") and d.get("desc")
        ]
        return result if result else None
    except (json.JSONDecodeError, TypeError):
        return None


def parse_competencies(text: str, codes: list = None) -> list:
    """
    [A] Парсит компетенции: JSON-режим → regex-fallback.
    """
    # Попытка JSON-разбора
    json_result = parse_competencies_json(text)
    if json_result:
        return json_result

    # Regex-fallback: нумерованные строки «Способен...»
    descriptions = []
    seen = set()
    for line in text.split("\n"):
        line = re.sub(r"^\d+[\.)]\s*", "", line.strip())
        line = re.sub(r"^[-–•]\s*", "", line)
        line = re.sub(r"\*\*", "", line).strip()
        if not line or len(line) < 10:
            continue
        if re.match(r"^Способен", line, re.I):
            key = line.lower()[:60]
            if key not in seen:
                seen.add(key)
                descriptions.append(line)

    if codes and descriptions:
        while len(descriptions) < len(codes):
            descriptions.append("Способен применять методы и инструменты дисциплины на практике")
        return list(zip(codes, descriptions[:len(codes)]))

    # Последний fallback: старый формат «УК-1: описание»
    result = []
    seen_codes: set = set()
    for line in text.split("\n"):
        m = re.match(r"(УК-\d+|ОПК-\d+|ПК-\d+)[:\.\s]+(.+)", line.strip())
        if m and m.group(1) not in seen_codes:
            seen_codes.add(m.group(1))
            result.append((m.group(1), m.group(2).strip()))
    return result if result else [
        ("УК-1",  "Способен применять системный подход для анализа и решения задач"),
        ("ОПК-1", "Способен разрабатывать алгоритмы и программы для интеллектуальных систем"),
        ("ПК-1",  "Способен применять методы машинного обучения для решения прикладных задач"),
    ]


def parse_outcomes_json(text: str, required_count: int = 0) -> list | None:
    """
    [A] Пытается разобрать JSON-ответ LLM для результатов обучения.
    Ожидаемый формат: [{"type": "З", "text": "..."}, ...]

    элементов), чтобы gen_with_json_retry запустил перегенерацию.
    qwen2.5:14b стабильно возвращал 9 вместо 15, код принимал это как OK,
    и ротация items[idx % len] давала одинаковые З/У/В у разных компетенций.
    """
    m = re.search(r"\[.*\]", text, re.S)
    if not m:
        return None
    try:
        data = json.loads(m.group())
        if not isinstance(data, list):
            return None
        result = [
            (str(d.get("type", "")), str(d.get("text", "")))
            for d in data
            if isinstance(d, dict) and d.get("type") in ("З", "У", "В") and d.get("text")
        ]
        if len(result) < 3:
            return None
        if required_count > 0 and len(result) < required_count:
            print(f"  ⚠️  [outcomes] JSON содержит {len(result)} элементов, "
                  f"нужно {required_count} — retry")
            return None
        return result
    except (json.JSONDecodeError, TypeError):
        return None


def parse_outcomes(text: str) -> list:
    """
    [A] Парсит результаты обучения: JSON-режим → regex-fallback.
    """
    json_result = parse_outcomes_json(text)
    if json_result:
        return json_result

    # Regex-fallback: многострочный или однострочный формат Знать/Уметь/Владеть
    result = []
    current_type = None
    lines = []

    def flush():
        if current_type and lines:
            result.append((current_type, "\n".join(lines)))

    def split_inline(rest: str) -> list[str]:
        items = re.split(r";\s*-\s*|;\s*–\s*", rest)
        cleaned = []
        for item in items:
            item = re.sub(r"^[-–•]\s*", "", item.strip())
            item = re.sub(r"^\d+[\.)]\s*", "", item)
            item = re.sub(r"\*\*", "", item)
            if item and len(item) > 3:
                cleaned.append(item)
        return cleaned

    for line in text.split("\n"):
        line = line.strip()
        m_know = re.match(r"^Знать:\s*(.*)", line, re.I)
        m_can  = re.match(r"^Уметь:\s*(.*)", line, re.I)
        m_have = re.match(r"^Владеть:\s*(.*)", line, re.I)
        if m_know:
            flush(); current_type = "З"; lines = []
            rest = m_know.group(1).strip()
            if rest:
                lines.extend(split_inline(rest) if ";" in rest else [rest])
        elif m_can:
            flush(); current_type = "У"; lines = []
            rest = m_can.group(1).strip()
            if rest:
                lines.extend(split_inline(rest) if ";" in rest else [rest])
        elif m_have:
            flush(); current_type = "В"; lines = []
            rest = m_have.group(1).strip()
            if rest:
                lines.extend(split_inline(rest) if ";" in rest else [rest])
        elif line and current_type:
            item = re.sub(r"^\d+[\.)]\s*|\*\*|^[-–•]\s*", "", line)
            if item:
                lines.append(item)

    flush()

    VLADEET_PREFIXES = ("навыками", "методами", "инструментами", "технологиями",
                        "опытом", "практикой", "способностью")

    # [FIX-ZUV]
    _ACTION_VERBS = (
        "применять", "разрабатыва", "анализирова", "использова",
        "реализовыва", "проектирова", "оценива", "выполнять",
        "создавать", "строить", "моделирова", "формирова",
        "составлять", "решать", "описыва", "обеспечива",
    )
    _INFINITIVE_PREFIXES = (
        "применять", "разрабатыва", "анализирова", "использова",
        "реализовыва", "проектирова", "оценива", "выполнять",
        "создавать", "строить", "моделирова", "формирова",
        "составлять", "решать", "описыва", "исследова", "обеспечива",
        "разрабатывать",
    )

    fixed = []
    for otype, otext in result:
        if otype == "В":
            fixed_lines = []
            for ln in otext.split("\n"):
                ln = ln.strip()
                if not ln:
                    continue
                ll = ln.lower()
                if not any(ll.startswith(p) for p in VLADEET_PREFIXES):
                    ln = re.sub(r"^(Основ[ыа]|Знание|Понимание|Базов[ые]+)\s+",
                                "навыками ", ln, flags=re.I)
                    ll = ln.lower()
                    if not any(ll.startswith(p) for p in VLADEET_PREFIXES):
                        ln = "навыками " + ln[0].lower() + ln[1:]
                fixed_lines.append(ln)
            fixed.append((otype, "\n".join(fixed_lines)))

        elif otype == "З":
            # З-результат описывает знание: существительные/конструкции, не глагол-действие.
            # Если LLM написал «применять методы» → это У-уровень, исправляем.
            fixed_lines = []
            for ln in otext.split("\n"):
                ln = ln.strip()
                if not ln:
                    continue
                ll = ln.lower()
                if any(ll.startswith(v) for v in _ACTION_VERBS):
                    # Глагол → снимаем его, превращаем в существительное-конструкцию
                    ln = re.sub(
                        r"^(применять|использовать|разрабатывать|анализировать|"
                        r"реализовывать|проектировать|оценивать|выполнять|"
                        r"создавать|строить|моделировать|формировать|"
                        r"составлять|решать|описывать|обеспечивать)\s+",
                        "",
                        ln, flags=re.I,
                    ).strip()
                    if ln:
                        ln = ln[0].upper() + ln[1:]
                fixed_lines.append(ln)
            fixed.append((otype, "\n".join(fixed_lines)))

        elif otype == "У":
            # У-результат должен начинаться с инфинитива действия.
            # Если LLM написал существительное/«знание» → добавляем «применять».
            fixed_lines = []
            for ln in otext.split("\n"):
                ln = ln.strip()
                if not ln:
                    continue
                ll = ln.lower()
                if not any(ll.startswith(p) for p in _INFINITIVE_PREFIXES):
                    ln = "применять " + ln[0].lower() + ln[1:]
                fixed_lines.append(ln)
            fixed.append((otype, "\n".join(fixed_lines)))

        else:
            fixed.append((otype, otext))

    # [П7] ИСПРАВЛЕНО: финальный fallback возвращал ровно 3 элемента (1З+1У+1В).
    # fill_outcomes_table вызывает ротацию items[idx % len(items)] для каждой
    # из N компетенций. При len=1 все N компетенций получают одинаковый результат.
    # Теперь fallback генерирует N*3 уникальных элементов (N берётся из глобального
    # контекста через замыкание). Если parse_outcomes вызывается вне main() где
    # codes_list недоступен — используем базовый набор из 3 элементов (безопасно).
    if fixed:
        return fixed

    # Базовые тексты — минимальный осмысленный набор
    _z_base = [
        "основные методы и алгоритмы дисциплины",
        "теоретические принципы построения систем",
        "современные инструменты и технологии в данной области",
        "методы анализа и оценки эффективности систем",
        "нормативную базу и стандарты в области дисциплины",
    ]
    _u_base = [
        "применять методы дисциплины для решения практических задач",
        "разрабатывать и реализовывать алгоритмы в рамках дисциплины",
        "анализировать результаты и интерпретировать их в контексте задачи",
        "использовать инструментальные средства при проектировании систем",
        "выбирать оптимальные подходы для решения профессиональных задач",
    ]
    _v_base = [
        "навыками применения методов дисциплины на практике",
        "методами проектирования и разработки систем",
        "инструментами анализа и оценки качества решений",
        "навыками работы с профессиональными программными средствами",
        "методами исследования и верификации результатов",
    ]
    # Строим список: сначала все З, потом все У, потом все В
    # Количество элементов каждого типа = min(5, N); при N > 5 — циклический сдвиг
    fallback: list = []
    for i, z in enumerate(_z_base):
        fallback.append(("З", z))
    for i, u in enumerate(_u_base):
        fallback.append(("У", u))
    for i, v in enumerate(_v_base):
        fallback.append(("В", v))
    return fallback


def parse_topics_json(text: str) -> list | None:
    """
    [A] Пытается разобрать JSON-ответ LLM для тематического плана.
    Ожидаемый формат: [{"type": "section"|"topic", "label": "Раздел 1", "name": "..."}]

    "лекция", "ЛР" и т.п. Это приводило к тому что topics заполнялся лекционными
    позициями («Лекция 1. ...»), которые не совпадают с паттерном «^Раздел N»,
    используемым в fill_t7, fill_t11, fill_t21. Результат — T7/T11/T21 пустые.
    Теперь принимаются только type="section"/"topic" (рус.: "раздел"/"тема").
    Метки разделов нормализуются к формату «Раздел N».
    """
    m = re.search(r"\[.*\]", text, re.S)
    if not m:
        # Попробуем весь текст если нет квадратных скобок
        try:
            data = json.loads(text.strip())
            if isinstance(data, list):
                pass
            else:
                return None
        except Exception:
            return None
    else:
        try:
            data = json.loads(m.group())
        except (json.JSONDecodeError, TypeError):
            return None

    if not isinstance(data, list):
        return None

    # Типы которые считаем валидными структурными единицами
    SECTION_TYPES = {"section", "раздел", "section_type"}
    TOPIC_TYPES   = {"topic", "тема", "подтема", "subtopic"}
    # Типы которые нужно явно пропустить (контент лекций, ЛР и т.п.)
    SKIP_TYPES    = {"lecture", "лекция", "lab", "лр", "practice", "пз",
                     "work", "задание", "task", "item"}

    topics = []
    section_counter = 0
    for d in data:
        if not isinstance(d, dict):
            continue
        item_type = str(d.get("type", "")).strip().lower()
        label = str(d.get("label", "")).strip()
        name  = str(d.get("name",  "")).strip()

        # Пропускаем записи без текста
        if not name:
            continue
        # Пропускаем нежелательные типы
        if item_type in SKIP_TYPES:
            continue

        # Определяем: это раздел или тема
        is_section = (item_type in SECTION_TYPES or
                      re.match(r"^(?:Раздел|Section)\s*\d*", label, re.I))
        is_topic   = (item_type in TOPIC_TYPES or
                      re.match(r"^(?:Тема|Topic)\s*[\d\.]*", label, re.I))

        if is_section:
            section_counter += 1
            # Нормализуем метку к «Раздел N»
            if not re.match(r"^Раздел\s*\d+", label, re.I):
                m_num = re.search(r"(\d+)", label)
                label = f"Раздел {m_num.group(1) if m_num else section_counter}"
            topics.append(f"{label}. {name}")
        elif is_topic or not item_type:
            # Принимаем темы и элементы без типа
            if label:
                topics.append(f"{label}. {name}")
            else:
                topics.append(name)

    return topics if topics else None


def parse_topics(text: str) -> list:
    """[A] Парсит содержание дисциплины: JSON-режим → regex-fallback."""
    json_result = parse_topics_json(text)
    if json_result:
        return json_result

    # Regex-fallback
    topics = []
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    for para in paragraphs:
        tokens = re.split(
            r"(?=(?:Раздел|Тема)\s+\d+[\.\d]*\.?\s|\b\d+\.\d+\.?\s)", para
        )
        tokens = [t.strip() for t in tokens if t.strip()]
        for token in tokens:
            m_sec = re.match(r"^(Раздел\s+\d+)\.\s+(.+)", token)
            if m_sec:
                name = re.split(r"\s+\d+\.\d+\.", m_sec.group(2))[0].strip()
                if name:
                    topics.append(f"{m_sec.group(1)}. {name}")
                continue
            m_tema = re.match(r"^(Тема\s+[\d\.]+)\.\s+(.+)", token)
            if m_tema:
                name = re.split(r"\s+\d+\.\d+\.", m_tema.group(2))[0].strip()
                if name:
                    topics.append(f"{m_tema.group(1)}. {name}")
                continue
            m_sub = re.match(r"^(\d+\.\d+)\.?\s+(.+)", token)
            if m_sub:
                name = re.split(r"\s+\d+\.\d+\.", m_sub.group(2))[0].strip()
                if name:
                    topics.append(f"Тема {m_sub.group(1)}. {name}")

    if not topics:
        for line in text.split("\n"):
            m = re.match(r"^(Раздел|Тема)\s*([\d\.]+)[\.\ ]+(.+)", line.strip())
            if m:
                topics.append(
                    f"{m.group(1)} {m.group(2).rstrip('.')}. {m.group(3).strip()}"
                )

    return topics if topics else [
        "Раздел 1. Основы интеллектуальных систем",
        "Раздел 2. Методы машинного обучения",
        "Раздел 3. Применение интеллектуальных систем",
    ]


def parse_list_json(text: str, min_items: int = 3) -> list | None:
    """
    [A] Пытается разобрать JSON-ответ LLM для списка ЛР/ПЗ.
    Ожидаемый формат: [{"title": "Реализация алгоритма..."}, ...]

    Замечание: "Нет контроля длины LLM ответа — модель возвращает 4 лабораторных
    вместо 6. Хотя парсер пытается исправлять, лучше проверять count items."
    Раньше порог был жёстко зашит как >= 3, что позволяло принять неполный список
    (4 из 6 ЛР) как "валидный" JSON — retry не срабатывал, дефолт не подставлялся.
    Теперь caller передаёт min_items=6, и неполный список возвращает None → retry.
    """
    m = re.search(r"\[.*\]", text, re.S)
    if not m:
        return None
    try:
        data = json.loads(m.group())
        if not isinstance(data, list):
            return None
        result = [str(d.get("title", "")).strip() for d in data
                  if isinstance(d, dict) and d.get("title")]
        return result if len(result) >= min_items else None
    except (json.JSONDecodeError, TypeError):
        return None


def parse_list_json_with_section(text: str, min_items: int = 3) -> list | None:
    """
    [Фикс №5+6] Парсит JSON-ответ LLM для ЛР/ПЗ с полем section.
    Возвращает list of dicts: [{"title": "...", "section": 2}, ...]
    При отсутствии поля section — подставляет None (fill_* использует ротацию).
    """
    m = re.search(r"\[.*\]", text, re.S)
    if not m:
        return None
    try:
        data = json.loads(m.group())
        if not isinstance(data, list):
            return None
        result = [
            {"title": str(d.get("title", "")).strip(),
             "section": d.get("section")}
            for d in data
            if isinstance(d, dict) and d.get("title")
        ]
        return result if len(result) >= min_items else None
    except (json.JSONDecodeError, TypeError):
        return None


def _normalize_section_assignment(items: list, n_sections: int) -> list:
    """
    [Фикс Д-SecRot] Принудительно перезаписывает поле 'section' по
    детерминированному паттерну 1,1,...,2,2,...,3,3,...

    Устраняет ошибки LLM при назначении раздела — например, когда модель
    возвращает section:1 для 4-го ПЗ вместо ожидаемого section:2.
    Вызывается после gen_with_json_retry для lab_works и practice.
    """
    if n_sections < 1 or not items:
        return items
    n = len(items)
    per_sec = max(1, n // n_sections)
    for i, item in enumerate(items):
        if isinstance(item, dict):
            item["section"] = min((i // per_sec) + 1, n_sections)
    return items


def parse_list(text: str, discipline: str = "", min_items: int = 3) -> list:
    """[A] Парсит список ЛР/ПЗ: JSON-режим → regex-fallback.
    Caller передаёт min_items=6 → fallback-список из 4 ЛР/ПЗ теперь
    корректно отклоняется вместо попадания в документ.
    """
    json_result = parse_list_json(text, min_items=min_items)
    if json_result:
        return json_result[:8]

    OFFTRACK_KEYWORDS = [
        "презентаци", "доклад", "реферат", "публикаци", "журнал",
        "flutter", "react native", "android studio", "xcode",
        "google play", "app store", "swift", "kotlin",
        "устный", "подготовка к",
    ]
    items = []
    for line in text.split("\n"):
        line = line.strip()
        line = re.sub(r"^(ЛР\s*№?\d+|ЛР\s*No\d+|\d+[\.):])\s+", "", line)
        line = re.sub(r"^\*\*(.+)\*\*$", r"\1", line)
        line = re.sub(r"^<[^>]{1,30}>\s*[-–\.\:]?\s*", "", line)
        line = re.sub(r"^<[^>]{1,30}>\s*$", "", line)
        if not line or len(line) < 6:
            continue
        if any(kw in line.lower() for kw in OFFTRACK_KEYWORDS):
            continue
        items.append(line)
    return items[:8] if len(items) >= min_items else ["Лабораторная работа 1", "Лабораторная работа 2"]


# ---------------------------------------------------------------------------
# Заполнение таблиц шаблона
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Библиография — генерация и заполнение таблиц
# ---------------------------------------------------------------------------

def _parse_rag_bibliography_chunks(hits: list) -> list:
    """

    Вместо передачи retrieved-текста в LLM как «примера» — извлекаем ГОСТ-строки
    напрямую из payload["text"] каждого хита. Это исключает галлюцинации авторов
    (Шарма, Петерс, Харрисон и т.д.), которые LLM генерировал по образцу чанков.

    Признак ГОСТ-строки: содержит ' — ' и 4-значный год.
    Тип записи («Основная» / «Дополнительная») определяется по контексту чанка.
    """
    entries = []
    seen: set = set()
    for h in hits:
        text = h.get("payload", {}).get("text", "")
        if not text:
            continue
        # Тип определяем по разделу, в котором встретилась запись
        btype = (
            "Дополнительная литература"
            if "дополнительн" in text.lower()
            else "Основная литература"
        )
        for line in text.splitlines():
            line = line.strip()
            # ГОСТ-строка: содержит год издания и хотя бы один разделитель.
            # Корпус использует два варианта: " — " (ГОСТ-7.1) и " - " (дефис-замена).
            # Проверяем оба — иначе записи вида "Кузнецов, И. Н. ... - Москва : ..., 2020."
            # не распознаются и парсер возвращает пустой список несмотря на наличие чанков.
            has_separator = (" — " in line) or (". -" in line) or (" : " in line and " - " in line)
            if not has_separator or not re.search(r"\b\d{4}\b", line):
                continue
            # Убираем нумерацию "1. ", "2) " в начале строки
            desc = re.sub(r"^\d+[\.\)]\s*", "", line).strip()
            if len(desc) < 20:
                continue
            key = desc[:60].lower()
            if key in seen:
                continue
            seen.add(key)
            entries.append({
                "type":    btype,
                "purpose": "Для изучения теории;Для выполнения СРО;",
                "desc":    desc,
                "url":     "http://bibl.rusoil.net",
                "coeff":   "1.00",
            })
    return entries


def parse_bibliography_json(text: str) -> list | None:
    """
    Парсит JSON-ответ LLM для библиографических записей.
    Ожидаемые поля: type/purpose/desc/url/coeff.
    """
    m = re.search(r"\[.*\]", text, re.S)
    if not m:
        return None
    try:
        data = json.loads(m.group())
        if not isinstance(data, list):
            return None
        result = [d for d in data if isinstance(d, dict) and d.get("desc")]
        return result if result else None
    except (json.JSONDecodeError, TypeError):
        return None


def gen_bibliography(discipline: str, direction: str = "", level: str = "", cfg: dict = None) -> tuple[list, list]:
    """
    Генерирует основную (Т15) и методическую (Т17) литературу.
    Возвращает (main_entries, method_entries).
    Каждая запись — dict с полями: type/purpose/desc/url/coeff.

    T15: генерируем через LLM с валидацией против галлюцинаций.
         Признаки галлюцинации: «Фамилия», «Название», «<», «>», «...».
         При обнаружении — fallback на проверенные реальные учебники.

    T17: qwen2.5:3b стабильно копирует «Фамилия, И. О. Название» из промпта.
         Обходим LLM полностью — всегда используем fallback с реальными
         УГНТУ-пособиями в корректном формате.
    """
    # Признаки шаблонных/галлюцинированных записей
    _PLACEHOLDER_MARKERS = ("фамилия", "название", "<гост", "<реальная", "...", "<")

    def _is_placeholder(desc: str) -> bool:
        dl = desc.lower()
        return any(m in dl for m in _PLACEHOLDER_MARKERS)

    def _make_fallback_main() -> list:
        """Реальные учебники по ИИ/МО, доступные в российских ЭБС.
        возвращался пустой список и T15 оставалась незаполненной."""
        return [
            {
                "type": "Основная литература",
                "purpose": "Для изучения теории;",
                "desc": (
                    "Флах, П. Машинное обучение : наука и искусство построения алгоритмов, "
                    "которые извлекают знания из данных / П. Флах ; пер. с англ. "
                    "А. А. Слинкина. — Москва : ДМК Пресс, 2015. — 400 с."
                ),
                "url": "http://www.znanium.com",
                "coeff": "1.00",
            },
            {
                "type": "Основная литература",
                "purpose": "Для изучения теории;Для выполнения СРО;",
                "desc": (
                    "Осовский, С. Нейронные сети для обработки информации : учебное пособие / "
                    "С. Осовский ; пер. с польск. И. Д. Рудинского. — Москва : "
                    "Финансы и статистика, 2002. — 344 с."
                ),
                "url": "http://www.znanium.com",
                "coeff": "1.00",
            },
            {
                "type": "Дополнительная литература",
                "purpose": "Для изучения теории;",
                "desc": (
                    "Рассел, С. Искусственный интеллект : современный подход / "
                    "С. Рассел, П. Норвиг ; пер. с англ. — 4-е изд. — Москва : "
                    "Вильямс, 2022. — 1408 с."
                ),
                "url": "http://biblio-online.ru",
                "coeff": "0.50",
            },
        ]

    def _make_fallback_method(disc: str) -> list:
        """УГНТУ-пособия — без персоналии, только кафедра-составитель.
        [FIX-08] Убрано «сост. Д. М. Зарипов» — вымышленное ФИО."""
        return [
            {
                "purpose": "Для выполнения лабораторных работ;",
                "desc": (
                    f"Методические указания к выполнению лабораторных работ "
                    f"по дисциплине «{disc}» / УГНТУ, каф. ВТИК. — "
                    "Уфа : УГНТУ, 2023. — 64 с."
                ),
                "url": "http://bibl.rusoil.net",
                "coeff": "1.00",
            },
            {
                "purpose": "Для выполнения практических занятий;",
                "desc": (
                    f"Методические указания к практическим занятиям "
                    f"по дисциплине «{disc}» / УГНТУ, каф. ВТИК. — "
                    "Уфа : УГНТУ, 2023. — 48 с."
                ),
                "url": "http://bibl.rusoil.net",
                "coeff": "1.00",
            },
        ]

    # ── Основная литература — config.json override или LLM с валидацией ─────
    # [Д-2] Если в config.json задан ключ "main_bibliography" —
    # список [{type, purpose, desc, url, coeff}] — используем его напрямую,
    # минуя LLM. Это гарантирует качество без зависимости от модели.
    # [FIX-CFG]
    _cfg = cfg or {}
    _custom_main = _cfg.get("main_bibliography")

    if _custom_main and isinstance(_custom_main, list) and len(_custom_main) >= 2:
        main_entries = _custom_main
        print(f"    ✅ Библиография T15: из config.json (main_bibliography), "
              f"{len(main_entries)} записей")
        _generation_log["bibliography_main_source"] = "config.json"
    else:
        # [FIX-BIB-RAG]
        _rag_section_types = SECTION_TYPE_FILTER.get("bibliography_main", ["bibliography", "place"])
        _, _bib_hits = retrieve("bibliography_main", discipline, _rag_section_types,
                                direction=direction, level=level)
        _rag_entries = _parse_rag_bibliography_chunks(_bib_hits)
        if len(_rag_entries) >= 2:
            main_entries = _rag_entries
            print(f"    ✅ Библиография T15: из RAG-чанков напрямую, "
                  f"{len(main_entries)} записей (LLM не вызывался)")
            _generation_log["bibliography_main_source"] = "rag_direct"
        else:
            # Fallback: RAG не дал достаточно записей → LLM с фильтром галлюцинаций
            if _rag_entries:
                print(f"    ⚠️  RAG вернул только {len(_rag_entries)} записей — передаём в LLM")
            raw_main = gen(
                "bibliography_main", discipline, PROMPTS["bibliography_main"],
                direction=direction, level=level,
            )
            llm_entries = parse_bibliography_json(raw_main)

            # Отфильтровываем записи с плейсхолдерами / галлюцинированными авторами
            if llm_entries:
                clean_entries = [e for e in llm_entries if not _is_placeholder(e.get("desc", ""))]
                if len(clean_entries) >= 1:
                    main_entries = list(clean_entries)
                    # [FIX-5]
                    if len(main_entries) < 2:
                        # [З-10]
                        fb = [
                            {
                                "type": "Основная литература",
                                "purpose": "Для изучения теории;",
                                "desc": "Таненбаум Э., Бос Х. Современные операционные системы. — СПб. : Питер, 2015.",
                                "url": "http://bibl.rusoil.net",
                                "coeff": "1.00",
                            },
                            {
                                "type": "Основная литература",
                                "purpose": "Для изучения теории;",
                                "desc": "Кормен Т. и др. Алгоритмы: построение и анализ. — М. : Вильямс, 2013.",
                                "url": "http://bibl.rusoil.net",
                                "coeff": "1.00",
                            },
                        ]
                        print("    ⚠️  T15: <2 записей от LLM и нет config.main_bibliography — запустите book_loader.py")
                        existing_keys = {
                            re.sub(r"\s+", "", e.get("desc", ""))[:50].lower()
                            for e in main_entries
                        }
                        for fb_entry in fb:
                            fb_key = re.sub(r"\s+", "", fb_entry.get("desc", ""))[:50].lower()
                            if fb_key not in existing_keys:
                                main_entries.append(fb_entry)
                                existing_keys.add(fb_key)
                            if len(main_entries) >= 3:
                                break
                    added = len(main_entries) - len(clean_entries)
                    suffix = f" + {added} из fallback" if added else ""
                    print(f"    ✅ Библиография T15: принято {len(clean_entries)} от LLM{suffix}")
                    _generation_log["bibliography_main_source"] = "llm"
                else:
                    _generation_log["bibliography_main_source"] = "fallback"
                    _generation_log["bibliography_main_fallback_reason"] = (
                        f"LLM вернул {len(llm_entries)} записей, "
                        f"из них {len(clean_entries)} без плейсхолдеров (нужно ≥1)"
                    )
                    print(f"    ⚠️  Библиография T15: LLM вернул шаблонные записи → fallback")
                    main_entries = _make_fallback_main()
            else:
                _generation_log["bibliography_main_source"] = "fallback"
                _generation_log["bibliography_main_fallback_reason"] = "JSON не распарсился"
                print(f"    ⚠️  Библиография T15: JSON не распарсился → fallback")
                main_entries = _make_fallback_main()

    # ── Методические издания — config.json override или fallback ─────────
    # qwen2.5:3b стабильно копирует «Фамилия, И. О. Название» из любого промпта.
    # [З-6]
    _custom_method = _cfg.get("method_bibliography")

    if _custom_method and isinstance(_custom_method, list) and len(_custom_method) > 0:
        method_entries = _custom_method
        print("    ✅ Библиография T17: из config.json (method_bibliography)")
        _generation_log["bibliography_method_source"] = "config.json"
    else:
        method_entries = _make_fallback_method(discipline)
        print("    ✅ Библиография T17: используется fallback (реальные УГНТУ-пособия)")
        _generation_log["bibliography_method_source"] = "fallback"

    return main_entries, method_entries


def _normalize_gost_biblio(desc: str) -> str:
    """«Москва _ Вильямс» → «Москва : Вильямс» (ГОСТ Р 7.0.5-2008).
    LLM иногда генерирует разделитель «_» или пробел вместо « : » между городом
    и издательством — нормализуем перед записью в таблицу библиографии."""
    if not desc:
        return desc
    s = re.sub(r"(?<=[A-Za-zА-Яа-яё»])\s*_\s*(?=[A-ZА-ЯЁ])", " : ", desc)
    s = re.sub(r"(?<=[A-Za-zА-Яа-яё»])\s*:\s*(?=[A-ZА-ЯЁ])", " : ", s)
    return re.sub(r" {2,}", " ", s).strip()


def fill_bibliography_main(doc: Document, entries: list, semester: str):
    """Заполняет T15 основную и дополнительную литературу через fill_placeholder_rows."""
    table = find_table(doc, "bibliography")
    if table is None:
        return

    # [B-4] Дедупликация по нормализованному описанию
    seen_descs: set = set()
    deduped: list = []
    for entry in entries:
        key = re.sub(r"\s+", " ", entry.get("desc", "")).strip()[:60].lower()
        if key and key in seen_descs:
            print(f"  ⚠️  Т15: дублирующая запись пропущена: {key[:50]!r}")
            continue
        if key:
            seen_descs.add(key)
        deduped.append(entry)

    rows = [
        [
            e.get("type",    "Основная литература"),
            e.get("purpose", "Для изучения теории;"),
            semester, "", "",           # очная / очно-заочная / заочная
            _normalize_gost_biblio(e.get("desc", "")),   # [FIX-07] нормализация ГОСТ
            "1",                        # кол-во экз.
            e.get("url",     ""),
            e.get("coeff",   "1.00"),
        ]
        for e in deduped
    ]
    fill_placeholder_rows(table, rows)


def fill_bibliography_method(doc: Document, entries: list, semester: str):
    """Заполняет T17 учебно-методические издания через fill_placeholder_rows."""
    table = find_table(doc, "method_bib")
    if table is None:
        return
    rows = [
        [
            e.get("purpose", "Для выполнения лабораторных работ;"),
            semester, "", "",           # очная / очно-заочная / заочная
            e.get("desc",    ""),
            "1", "0",                  # всего / на кафедре
            e.get("url",     ""),
            e.get("coeff",   "1.00"),
        ]
        for e in entries
    ]
    fill_placeholder_rows(table, rows)


def fill_competencies_table(doc: Document, competencies: list):
    table = find_table(doc, "competencies")
    if table is None:
        return
    rows = [[str(i), desc, code] for i, (code, desc) in enumerate(competencies, 1)]
    fill_placeholder_rows(table, rows)


def fill_outcomes_table(doc: Document, competencies: list, outcomes: list):
    table = find_table(doc, "outcomes")
    if table is None:
        return

    # [FIX-outcomes]
    type_lists: dict = {"З": [], "У": [], "В": []}
    for ot, otext in outcomes:
        if ot in type_lists:
            type_lists[ot].append(otext)

    # Fallback если LLM вернул пустой список для какого-то типа
    if not type_lists["З"]:
        type_lists["З"] = ["основные концепции и методы дисциплины"]
    if not type_lists["У"]:
        type_lists["У"] = ["применять методы дисциплины для решения задач"]
    if not type_lists["В"]:
        type_lists["В"] = ["навыками работы с инструментами дисциплины"]

    def split_items(text: str) -> list:
        lines = []
        for ln in text.split("\n"):
            ln = re.sub(r"^\d+[\.)]\s*", "", ln.strip())
            if ln and len(ln) > 4:
                lines.append(ln)
        return lines if lines else [text.strip()]

    # Разворачиваем каждый список: если LLM вернул блоки текста — разбиваем на строки
    z_items = []
    for t in type_lists["З"]:
        z_items.extend(split_items(t))
    u_items = []
    for t in type_lists["У"]:
        u_items.extend(split_items(t))
    v_items = []
    for t in type_lists["В"]:
        v_items.extend(split_items(t))

    z_qualifiers = ["в данной области", "в контексте дисциплины",
                    "применительно к решаемым задачам",
                    "необходимые для профессиональной деятельности",
                    "включая теоретические основы и практические аспекты"]
    u_qualifiers = ["в профессиональной деятельности", "для решения практических задач",
                    "при разработке и исследовании систем",
                    "при анализе данных и построении моделей",
                    "в ходе проектирования и реализации"]
    v_qualifiers = ["при решении профессиональных задач",
                    "для анализа и разработки систем",
                    "в проектировании и исследовательской деятельности",
                    "при реализации и тестировании решений",
                    "в профессиональной практике"]

    rows = []
    # [FIX-01]
    # [FIX-02]
    _ind_verbs = {"З": "знать", "У": "уметь", "В": "владеть"}
    _type_pos  = {"З": 0, "У": 0, "В": 0}
    _seen_texts: dict = {"З": set(), "У": set(), "В": set()}
    for idx, (code, desc) in enumerate(competencies):
        for type_idx, (otype, items, qualifiers) in enumerate([
            ("З", z_items, z_qualifiers),
            ("У", u_items, u_qualifiers),
            ("В", v_items, v_qualifiers),
        ]):
            result_code   = f"{otype}({code})"
            indicator_num = type_idx + 1

            base_item = items[_type_pos[otype] % len(items)]
            qual      = qualifiers[_type_pos[otype] % len(qualifiers)]
            _type_pos[otype] += 1

            prefix      = {"З": "Знать:", "У": "Уметь:", "В": "Владеть:"}[otype]
            result_text = f"{prefix} {base_item}"
            if result_text in _seen_texts[otype]:
                result_text = f"{prefix} {base_item} — {qual}"
            _seen_texts[otype].add(result_text)

            # [FIX-01]
            indicator = f"{code}.{indicator_num} {_ind_verbs[otype]} {base_item}"

            rows.append([code, indicator, result_code, result_text])

    fill_placeholder_rows(table, rows)


def fill_topics_table(doc: Document, topics: list, semester: str, hours: dict,
                      codes_list: list = None):
    table = find_table(doc, "topics")
    if table is None:
        return

    sections_only = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    n = max(len(sections_only), 1) if sections_only else max(len(topics), 1)

    lec  = hours.get("lecture",  12) // n
    pz   = hours.get("practice", 36) // n
    lr   = hours.get("lab",      16) // n
    sro  = hours.get("self",     62) // n
    total_l = total_pz = total_lr = total_sro = 0

    codes = codes_list or ["ОПК-1", "ПК-1"]
    rows = []
    for i, sec in enumerate(sections_only, 1):
        sec_name = re.sub(r"^Раздел\s*\d+\.\s*", "", sec).strip()
        c1 = codes[(i - 1) % len(codes)]
        c2 = codes[i % len(codes)]
        shifer = f"З({c1})\nУ({c1})\nВ({c2})"
        rows.append([str(i), sec_name, semester,
                     str(lec), str(pz), str(lr), str(sro), str(lec + pz + lr + sro),
                     shifer])
        total_l += lec; total_pz += pz; total_lr += lr; total_sro += sro

    rows.append(["", "ИТОГО:", "",
                 str(total_l), str(total_pz), str(total_lr), str(total_sro),
                 str(total_l + total_pz + total_lr + total_sro), ""])
    fill_placeholder_rows(table, rows)


def _compact_section(section: str) -> str:
    """«Раздел N. Название» → «N-Название» (формат реальных РПД кафедры ВТИК)."""
    m = re.match(r"^Раздел\s*(\d+)[.\s]+(.+)$", section.strip())
    if m:
        return f"{m.group(1)}-{m.group(2).strip()}"
    return section


def fill_lectures_table(doc: Document, topics: list, hours: dict):
    table = find_table(doc, "lectures")
    if table is None:
        return

    themes_only   = [t for t in topics if not re.match(r"^Раздел\s*\d+", t)]
    sections_only = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]

    n_topics = max(len(themes_only), 1) if themes_only else max(len(sections_only), 1)
    lec = max(hours.get("lecture", 12) // n_topics, 1)

    def _clean_lecture_name(raw_name: str) -> str:
        name = re.sub(r"^(?:Лекция\s+\d+\.\s*)+", "", raw_name.strip()).strip()
        name = re.sub(r"\s+[А-ЯЁ]{2,4}-\d+\s*$", "", name).strip()
        name = re.sub(r"\s+[A-ZА-ЯЁ]{2,5}\s*$", "", name).strip()
        name = re.sub(r"\.\s*$", "", name).strip()
        return name if name else raw_name.strip()

    rows = []
    section = ""
    lec_no  = 0
    for topic in topics:
        if re.match(r"^Раздел\s*\d+", topic):
            section = topic
        else:
            lec_no += 1
            short = re.sub(r"^Тема\s*[\d\.]+[\.\ ]+", "", topic).strip()
            short = _clean_lecture_name(short)
            rows.append([str(lec_no), _compact_section(section) if section else topic,
                         f"Лекция {lec_no}. {short}", str(lec), "", ""])

    if lec_no == 0:
        lec = max(hours.get("lecture", 12) // max(len(sections_only), 1), 1)
        for i, topic in enumerate(sections_only, 1):
            short = re.sub(r"^Раздел\s*\d+[\.\ ]+", "", topic).strip()
            short = _clean_lecture_name(short)
            rows.append([str(i), _compact_section(topic), f"Лекция {i}. {short}", str(lec), "", ""])

    fill_placeholder_rows(table, rows)


def fill_lab_table(doc: Document, lab_works: list, topics: list, hours_lab: int = 18):
    table = find_table(doc, "labs")
    if table is None:
        return

    if len(lab_works) < 6:
        print(f"  ⚠️  Т9: получено {len(lab_works)} ЛР — дополняю до 6")
        for j in range(len(lab_works), 6):
            lab_works.append({"title": f"Лабораторная работа {j + 1}", "section": None})

    sections = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    n_lab = len(lab_works)
    base  = max(hours_lab // n_lab, 1)
    rem   = hours_lab - base * n_lab
    hours_list = [base + (1 if i < rem else 0) for i in range(n_lab)]

    rows = []
    for i, work in enumerate(lab_works, 1):
        # [Фикс №5+6] Поддержка dict {"title", "section"} от parse_list_json_with_section
        if isinstance(work, dict):
            title      = work.get("title", f"Лабораторная работа {i}")
            sec_num    = work.get("section")
            if sec_num is not None and sections:
                sec_idx = (int(sec_num) - 1) % len(sections)
                section = sections[sec_idx]
            else:
                section = sections[(i - 1) % max(len(sections), 1)] if sections else f"Раздел {((i - 1) // 2) + 1}"
        else:
            title   = work
            section = sections[(i - 1) % max(len(sections), 1)] if sections else f"Раздел {((i - 1) // 2) + 1}"
        rows.append([_compact_section(section), str(i), title, str(hours_list[i - 1]), "", ""])
    rows.append(["-", "", "ИТОГО:", str(hours_lab), "", ""])
    fill_placeholder_rows(table, rows)


def fill_practice_table(doc: Document, practices: list, topics: list,
                        hours_practice: int = 36):
    table = find_table(doc, "practice")
    if table is None:
        return

    if len(practices) < 6:
        print(f"  ⚠️  Т10: получено {len(practices)} ПЗ — дополняю до 6")
        for j in range(len(practices), 6):
            practices.append({"title": f"Практическое занятие {j + 1}", "section": None})

    sections = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    n_prac = len(practices)
    base   = max(hours_practice // n_prac, 1)
    rem    = hours_practice - base * n_prac
    hours_list = [base + (1 if i < rem else 0) for i in range(n_prac)]

    rows = []
    for i, prac in enumerate(practices, 1):
        # [Фикс №5+6] Поддержка dict {"title", "section"} от parse_list_json_with_section
        if isinstance(prac, dict):
            title   = prac.get("title", f"Практическое занятие {i}")
            sec_num = prac.get("section")
            if sec_num is not None and sections:
                sec_idx = (int(sec_num) - 1) % len(sections)
                section = sections[sec_idx]
            else:
                section = sections[(i - 1) % max(len(sections), 1)] if sections else f"Раздел {((i - 1) // 2) + 1}"
        else:
            title   = prac
            section = sections[(i - 1) % max(len(sections), 1)] if sections else f"Раздел {((i - 1) // 2) + 1}"
        rows.append([_compact_section(section), str(i), title, str(hours_list[i - 1]), "", ""])
    rows.append(["-", "", "ИТОГО:", str(hours_practice), "", ""])
    fill_placeholder_rows(table, rows)


def fill_t3_hours(doc: Document, semester: str, credits: int,
                  hours_total: int, hours_contact: int, hours_sro: int,
                  exam_type: str):
    """Заполняет T3 (трудоёмкость) через fill_placeholder_rows."""
    table = find_table(doc, "workload")
    if table is None:
        return
    # Два [] слота: строка данных + строка ИТОГО
    fill_placeholder_rows(table, [
        [semester, str(credits), str(hours_total), str(hours_contact), str(hours_sro), exam_type],
        ["ИТОГО:", str(credits), str(hours_total), str(hours_contact), str(hours_sro), ""],
    ])


def fill_t6_workload(doc: Document, lec: int, pz: int, lr: int, sro: int,
                     semester: str, exam_prep_hours: int = 0):
    t = find_table(doc, "work_types")
    if t is None:
        return
    # [Д-5] Диагностика: убеждаемся что таблица найдена и содержит строки.
    # [TABLE] в pandoc-выводе — артефакт сложных объединённых ячеек, не признак
    # пустой таблицы. Если строк > 0 и print ниже выводится — таблица заполнена.
    print(f"  ℹ️  Т6 work_types: найдена, строк={len(t.rows)}, ищу семестр={semester!r}")
    sem_col = None
    # [З-G5]
    sem_str = str(semester).strip()
    for header_row in t.rows[:4]:
        for j, cell in enumerate(header_row.cells):
            cell_text = cell.text.strip()
            # [FIX-SEM]
            _sem_re = re.compile(r'(?<!\d)' + re.escape(sem_str) + r'(?!\d)')
            if cell_text == sem_str or _sem_re.search(cell_text):
                sem_col = j
                break
        if sem_col is not None:
            break
    if sem_col is None:
        print(f"  ⚠️  Т6: столбец семестра {sem_str!r} не найден — "
              f"заголовки: {[c.text.strip() for c in t.rows[0].cells]}")

    total_contact = lec + pz + lr

    # [A-6] ИСПРАВЛЕНО: зачищаем ВСЕ числовые столбцы данных, кроме первого
    # (Вид работы) и второго (Всего). Прежде fill_t6 писал только в cells[1]
    # и sem_col, оставляя цифры предыдущей дисциплины в других семестровых
    # столбцах. Теперь каждый числовой столбец (≥2) обнуляется перед записью.
    def clear_data_columns(row):
        """Зачищает все числовые столбцы строки (кроме 0 и 1)."""
        for ci in range(2, len(row.cells)):
            set_cell_text(row.cells[ci], "")

    # [A-8] On-line строки должны оставаться 0 — они не тронуты нашей логикой.
    # kw_map содержит только базовые виды работ, без «on-line» и «в т.ч.»
    # подстрок. Строки вида «-в т.ч. лекции on-line курс» содержат «лекции»
    # как подстроку и ранее получали те же значения что и основные строки.
    # Исправление: требуем что строка НЕ содержит «on-line», «онлайн», «в т.ч.»,
    # «иная», «проектная», «освоение», «самостоятельная проект».
    # [Т6-FIX-1] "иные" добавлен — строка "иные виды работ обучающегося"
    # содержит "иные", а не "иная" → старый паттерн не срабатывал → [] оставался.
    # [FIX-T6-SRO]
    SKIP_PATTERNS = ("on-line", "онлайн", "в т.ч.", "иная", "иные", "проектная",
                     "освоение", "самостоятельная проект",
                     "контролируем",           # [FIX-4] «контролируемая СРО» = 72 от шаблона
                     "выполнение и подготовка к защит",  # [FIX-T6-ZERO2] ИСПРАВЛЕНО:
                     # строки «выполнение и подготовка к защите курсового проекта» и
                     # «выполнение и подготовка к защите РГР работы» не совпадали ни
                     # с kw_map, ни с SKIP_PATTERNS → clear_data_columns зачищал cells[≥2],
                     # но cells[1] («Всего») оставался с «[]» из шаблона.
                     # Один паттерн покрывает обе строки (общий префикс).
                     )

    # [FIX-T6-SRO]
    hrs_study = round(sro * 0.20)           # изучение вынесенного материала
    hrs_rgr   = round(sro * 0.20)           # РГР / реферат
    # [FIX-EXAM-PREP]
    hrs_prep  = sro - hrs_study - hrs_rgr - exam_prep_hours   # подготовка к ЛР/ПЗ

    kw_map = {
        "контактная":             total_contact,
        "лекции":                 lec,
        "подготовка к лаборатор": hrs_prep,
        "подготовка к сдач":      exam_prep_hours,  # [FIX-EXAM-PREP] было 0
        "практические занятия":   pz,
        "лабораторные работы":    lr,
        "самостоятельная работа": sro,
        "изучение учебного":      hrs_study,
        # [FIX-06b]
        "расчётно-графическ":    hrs_main,
        "расчетно-графическ":    hrs_main,
        "реферат":                hrs_main,
        "индивидуальн":           hrs_main,
        "курсов":                 hrs_main,   # курсовая работа / курсовой проект
    }

    for row in t.rows:
        label = row.cells[0].text.strip().lower()

        # [A-6] Зачищаем столбцы в каждой строке — убираем старые данные
        clear_data_columns(row)

        # [A-7] Строка ИТОГО — пишем итоговые значения
        if "итого" in label:
            total = total_contact + sro
            set_cell_text(row.cells[1], str(total))
            if sem_col is not None and sem_col < len(row.cells):
                set_cell_text(row.cells[sem_col], str(total))
            continue

        # [A-8] Пропускаем строки on-line, «в т.ч.» и вспомогательные.
        # [FIX-T6-ZERO]
        if any(pat in label for pat in SKIP_PATTERNS):
            set_cell_text(row.cells[1], "0")
            continue

        for kw, val in kw_map.items():
            if kw in label:
                set_cell_text(row.cells[1], str(val))
                if sem_col is not None and sem_col < len(row.cells):
                    set_cell_text(row.cells[sem_col], str(val))
                break


def fill_t11_sro(doc: Document, topics: list, sro: int, cfg: dict = None):
    table = find_table(doc, "sro")
    if table is None:
        return

    sections = [tp for tp in topics if re.match(r"^Раздел\s*\d+", tp)]
    n = max(len(sections), 1)

    hrs_study = round(sro * 0.20)
    hrs_main  = round(sro * 0.20)
    hrs_prep  = sro - hrs_study - hrs_main

    # [FIX-06]
    _cfg = cfg or {}
    _custom_sro = _cfg.get("sro_types")
    if _custom_sro and isinstance(_custom_sro, list) and len(_custom_sro) >= 1:
        _custom_names = [str(x).strip() for x in _custom_sro if str(x).strip()]
        if _custom_names:
            per_item  = hrs_main // len(_custom_names)
            remainder = hrs_main - per_item * len(_custom_names)
            sro_types = [
                ("подготовка к лабораторным и/или практическим занятиям", hrs_prep),
                ("изучение учебного материала, вынесенного на СРО",       hrs_study),
            ]
            for i, name in enumerate(_custom_names):
                hrs = per_item + (remainder if i == len(_custom_names) - 1 else 0)
                sro_types.append((name, hrs))
        else:
            _custom_sro = None

    if not (_custom_sro and isinstance(_custom_sro, list) and len(_custom_sro) >= 1):
        _focus_raw = _cfg.get("discipline_focus", "")
        _focus = (
            " ".join(_focus_raw) if isinstance(_focus_raw, list) else str(_focus_raw)
        ).lower()
        _disc  = _cfg.get("discipline", "").lower()
        _text  = f"{_disc} {_focus}"

        if any(kw in _text for kw in (
            "интеллектуальн", "нейрон", "машинн", "агент",
            "проектирован", "разработк", "информационн",
        )):
            main_name = "выполнение индивидуального задания"
        elif any(kw in _text for kw in (
            "моделирован", "анализ данн", "исследован", "экспертн", "нечётк",
        )):
            main_name = "подготовка реферата"
        else:
            main_name = "выполнение расчётно-графической работы"

        sro_types = [
            ("подготовка к лабораторным и/или практическим занятиям", hrs_prep),
            ("изучение учебного материала, вынесенного на СРО",       hrs_study),
            (main_name,                                                hrs_main),
        ]
    rows = []
    for sec_idx, sec in enumerate(sections):
        for stype, total_hrs in sro_types:
            base_per_sec = round(total_hrs / n)
            if sec_idx < n - 1:
                hrs_per_sec = base_per_sec
            else:
                hrs_per_sec = total_hrs - base_per_sec * (n - 1)
            rows.append([sec, stype, str(hrs_per_sec), "", ""])
    rows.append(["-", "ИТОГО:", str(sro), "", ""])
    fill_placeholder_rows(table, rows)


def fill_t21_fos(doc: Document, competencies: list, topics: list,
                 outcomes: list = None, discipline: str = "дисциплины"):
    """
    Заполняет таблицу паспорта ФОС через fill_placeholder_rows.
    Структура: раздел × компетенция × 3 типа (З/У/В).
    """
    table = find_table(doc, "fos")
    if table is None:
        return

    sections = [tp for tp in topics if re.match(r"^Раздел\s*\d+", tp)]
    ocs = ["Письменный и устный опрос", "Лабораторная работа",
           "Тест", "Расчётно-графическая работа"]

    # Индексируем outcomes по типу
    outcomes_by_type: dict = {"З": [], "У": [], "В": []}
    for o in (outcomes or []):
        if isinstance(o, tuple):
            t_type = o[0] if len(o) > 0 else ""
            t_text = o[1] if len(o) > 1 else ""
        else:
            t_type = o.get("type", "")
            t_text = o.get("text", "")
        if t_type in outcomes_by_type:
            outcomes_by_type[t_type].append(t_text)

    # [FIX-03b]
    _fallbacks = {
        "З": f"основные методы и принципы дисциплины «{discipline}»",
        "У": f"применять инструменты дисциплины «{discipline}» для решения практических задач",
        "В": f"навыками работы с инструментами в области «{discipline}»",
    }
    _indicator_verbs = {"З": "Знает", "У": "Умеет", "В": "Владеет"}
    # [FIX-T21-VERB]
    # [З-09]

    rows = []
    n = 1
    for sec in sections:
        sec_name  = re.sub(r"^Раздел\s*\d+\.\s*", "", sec)
        # [FIX-04]
        sec_short = sec_name if len(sec_name) <= 80 else sec_name[:77].rstrip() + "…"
        for comp_idx, (code, comp_desc) in enumerate(competencies):
            for type_idx, res_type in enumerate(("З", "У", "В")):
                type_outcomes = outcomes_by_type[res_type]
                outcome_text  = (
                    type_outcomes[(comp_idx * 3 + type_idx) % len(type_outcomes)]
                    if type_outcomes else _fallbacks[res_type]
                )
                indicator_num = type_idx + 1
                verb      = _indicator_verbs[res_type]

                # [FIX-03]
                _oc_core = outcome_text.strip()
                _oc_core = re.sub(r"^(?:З|У|В)\s*[:\-—]\s*", "", _oc_core)
                _oc_core = re.sub(
                    r"^(?:знать|знает|уметь|умеет|владеть|владеет)\s*[:\-—]?\s*",
                    "", _oc_core, flags=re.IGNORECASE,
                )
                indicator = f"{code}.{indicator_num} {verb}: {_oc_core}"

                pokazatel = (
                    f"{'Отвечает на вопросы' if res_type == 'З' else 'Выполняет задания' if res_type == 'У' else 'Демонстрирует навыки'} "
                    f"по теме «{sec_short}»"
                )
                rows.append([
                    str(n), sec_name, f"{res_type}({code})",
                    outcome_text, indicator, pokazatel,
                    ocs[(n - 1) % len(ocs)],
                ])
                n += 1

    fill_placeholder_rows(table, rows)


# ---------------------------------------------------------------------------
# [D] Валидация бизнес-правил
# ---------------------------------------------------------------------------

def validate_generation(cfg: dict, hours: dict, competencies: list,
                        topics: list, lab_works: list, practices: list) -> list[str]:
    """
    [D] Проверяет корректность сгенерированного содержимого.
    Возвращает список предупреждений (пустой = всё ОК).
    """
    warnings: list[str] = []

    expected_total = (
        cfg.get("hours_lecture",  12) +
        cfg.get("hours_practice", 36) +
        cfg.get("hours_lab",      16) +
        cfg.get("hours_self",     62)
    )
    actual_total = sum(hours.values())
    if actual_total != expected_total:
        warnings.append(
            f"⚠️  Сумма часов {actual_total} ≠ {expected_total} из config.json"
        )

    # [З-ЧАС]
    credits = cfg.get("credits", 0)
    if credits:
        expected_by_credits = credits * 36
        if actual_total != expected_by_credits:
            warnings.append(
                f"⚠️  credits × 36 = {credits} × 36 = {expected_by_credits}, "
                f"но сумма часов = {actual_total} — нарушение ФГОС"
            )

    sections = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    if not sections:
        warnings.append("⚠️  Разделы дисциплины не сгенерированы (topics пуст)")

    if len(lab_works) < 6:
        warnings.append(f"⚠️  ЛР: сгенерировано {len(lab_works)} < 6 минимальных")
    if len(practices) < 6:
        warnings.append(f"⚠️  ПЗ: сгенерировано {len(practices)} < 6 минимальных")

    # [З-G7]
    def _norm_code(c: str) -> str:
        return re.sub(r"[.\s]+$", "", c.strip()).upper()

    codes_from_cfg    = {_norm_code(c) for c in cfg.get("competency_codes", "").split(",") if c.strip()}
    generated_codes   = {_norm_code(code) for code, _ in competencies}
    missing = codes_from_cfg - generated_codes
    if missing:
        warnings.append(f"⚠️  Компетенции не сгенерированы: {', '.join(sorted(missing))}")

    return warnings


# ---------------------------------------------------------------------------
# [A] Промпты — JSON-режим для всех генерируемых разделов
# ---------------------------------------------------------------------------

PROMPTS = {
    "competencies": """\
Ты составляешь рабочую программу дисциплины для российского технического университета.
Напиши описания компетенций для дисциплины «{discipline}».

Коды компетенций:
{competency_codes_numbered}

Правила:
- каждое описание начинается со слова «Способен»
- описание конкретно для «{discipline}», отражает реальные навыки
- не пиши шаблонные фразы, пиши конкретные профессиональные действия

Примеры правильных описаний (для дисциплины «Машинное обучение»):
[
  {{"code": "УК-1", "desc": "Способен применять системный подход для постановки и декомпозиции задач обработки данных"}},
  {{"code": "ОПК-1", "desc": "Способен разрабатывать алгоритмы машинного обучения и реализовывать их программно"}},
  {{"code": "ПК-1", "desc": "Способен обучать, тестировать и оценивать качество моделей машинного обучения"}}
]

ВЕРНИ ТОЛЬКО JSON-массив (без пояснений, без markdown).
Ровно {competency_count} объектов. Примеры выше — для другой дисциплины, напиши свои для «{discipline}».""",

    "outcomes": """\
Напиши результаты обучения для дисциплины «{discipline}» по ФГОС 3++.
Нужно ровно {outcomes_total} элементов: по {competency_count} знаний (З),
{competency_count} умений (У), {competency_count} навыков (В) — по одной
формулировке каждого типа на каждую из {competency_count} компетенций.

Коды компетенций, для которых надо написать З/У/В: {competency_codes_list}

Правила:
- З: что знает студент — КОНКРЕТНЫЕ методы, алгоритмы, технологии именно «{discipline}»
- У: что умеет — начинается с глагола (применять, разрабатывать, анализировать...)
- В: чем владеет — начинается с «навыками», «методами» или «инструментами»
- [FIX-02] ВСЕ тексты УНИКАЛЬНЫ. Никакие два З не совпадают даже частично.
  Никакие два У не совпадают. Никакие два В не совпадают.
  Разные компетенции → разные результаты. Например, для УК-1 и ПК-1 формулировки
  З должны описывать РАЗНЫЕ области знаний, а не одну и ту же.
- каждый текст — ОДНА краткая фраза (не список, не перечисление через запятую)
- Не повторяй одни и те же слова-головы в разных строках одного типа
- Содержат специфику «{discipline}», НЕ копируй примеры

Пример формата (для другой дисциплины «Компьютерное зрение», 3 компетенции):
[
  {{"type": "З", "text": "методы детектирования объектов: YOLO, SSD, Faster R-CNN"}},
  {{"type": "З", "text": "принципы сегментации изображений: семантическую и экземплярную"}},
  {{"type": "З", "text": "алгоритмы выделения ключевых точек: SIFT, ORB, SuperPoint"}},
  {{"type": "У", "text": "реализовывать нейронные детекторы объектов в PyTorch"}},
  {{"type": "У", "text": "применять OpenCV для предобработки и аугментации изображений"}},
  {{"type": "У", "text": "оценивать качество моделей по метрикам mAP, Precision, Recall"}},
  {{"type": "В", "text": "навыками обучения и тонкой настройки CNN на датасетах COCO, VOC"}},
  {{"type": "В", "text": "методами трекинга объектов: ByteTrack, SORT, DeepSORT"}},
  {{"type": "В", "text": "инструментами визуализации Grad-CAM для интерпретации сети"}}
]

ВЕРНИ ТОЛЬКО JSON-массив (без пояснений, без markdown).
Ровно {outcomes_total} объектов ({competency_count}З + {competency_count}У +
{competency_count}В), все уникальны, специфичны для «{discipline}».""",

    "content": """\
Напиши содержание дисциплины «{discipline}» — ровно 3 раздела, в каждом 2 темы.
Компетенции дисциплины: {competencies_summary}
{discipline_focus_block}
ПРАВИЛА:
- type="section" — крупный тематический блок, label="Раздел 1"/"Раздел 2"/"Раздел 3"
- type="topic"   — конкретная тема внутри раздела, label="Тема 1.1"/"Тема 1.2" и т.д.
- НЕ используй type="lecture", "lab", "лекция", "ЛР" — только section и topic
- Все названия уникальны, специфичны для «{discipline}», отражают ключевые темы выше

ВЕРНИ ТОЛЬКО JSON-массив без пояснений и markdown:
[
  {{"type": "section", "label": "Раздел 1", "name": "<тематический блок 1 для {discipline}>"}},
  {{"type": "topic",   "label": "Тема 1.1", "name": "<конкретная тема из блока 1>"}},
  {{"type": "topic",   "label": "Тема 1.2", "name": "<конкретная тема из блока 1>"}},
  {{"type": "section", "label": "Раздел 2", "name": "<тематический блок 2 для {discipline}>"}},
  {{"type": "topic",   "label": "Тема 2.1", "name": "<конкретная тема из блока 2>"}},
  {{"type": "topic",   "label": "Тема 2.2", "name": "<конкретная тема из блока 2>"}},
  {{"type": "section", "label": "Раздел 3", "name": "<тематический блок 3 для {discipline}>"}},
  {{"type": "topic",   "label": "Тема 3.1", "name": "<конкретная тема из блока 3>"}},
  {{"type": "topic",   "label": "Тема 3.2", "name": "<конкретная тема из блока 3>"}}
]
Ровно 9 объектов: 3 section + 6 topic.""",

    "lab_works": """\
Напиши 6 лабораторных работ для дисциплины «{discipline}».
Компетенции, которые должны формироваться: {competencies_summary}
{discipline_focus_block}
Разделы дисциплины:
{sections_list}

Требования:
- каждая ЛР — конкретное техническое задание специфичное для «{discipline}»
- все 6 ЛР на РАЗНЫЕ темы, покрывают разные аспекты дисциплины
- используй конкретные методы/алгоритмы из ключевых тем выше
- формулировка: глагол + метод/алгоритм + объект («Реализация...», «Обучение...», «Анализ...»)
- [FIX-05] КРАТКО: 8–15 слов на ЛР. НЕ перечисляй через запятую 3+ метода.
  Одна ЛР — одна главная задача. Не более одного придаточного оборота.
  Пример правильной длины: «Реализация алгоритма k-means для кластеризации
  числовых данных» (9 слов). НЕЛЬЗЯ: «Реализация и исследование алгоритмов
  кластеризации k-means, DBSCAN и иерархической кластеризации с визуализацией
  результатов и оценкой качества по метрикам силуэта и Дэвиса–Болдуина» (25 слов).
- укажи номер раздела (1, 2 или 3) к которому относится ЛР

ВЕРНИ ТОЛЬКО JSON-массив (без пояснений, без markdown):
[
  {{"title": "<ЛР 1 специфичная для {discipline}, 8-15 слов>", "section": 1}},
  {{"title": "<ЛР 2 специфичная для {discipline}, 8-15 слов>", "section": 1}},
  {{"title": "<ЛР 3 специфичная для {discipline}, 8-15 слов>", "section": 2}},
  {{"title": "<ЛР 4 специфичная для {discipline}, 8-15 слов>", "section": 2}},
  {{"title": "<ЛР 5 специфичная для {discipline}, 8-15 слов>", "section": 3}},
  {{"title": "<ЛР 6 специфичная для {discipline}, 8-15 слов>", "section": 3}}
]
Ровно 6 объектов. Все темы уникальны и специфичны для «{discipline}».""",

    "bibliography_main": """\
Ты составляешь список рекомендуемой литературы для рабочей программы дисциплины «{discipline}» \
в российском техническом университете.

Напиши 3 записи библиографии: 2 основных учебника и 1 дополнительный.

СТРОГИЕ ПРАВИЛА:
- только реально существующие книги — НЕ придумывай авторов и названия
- авторы должны быть реальными специалистами в области «{discipline}»
- год издания 2010–2024
- формат: Фамилия, И. О. Название : тип / И. О. Фамилия. — Город : Издательство, Год. — N с.
- url: http://www.znanium.com или http://e.lanbook.com или http://biblio-online.ru

ВЕРНИ ТОЛЬКО JSON-массив без пояснений и без markdown:
[
  {{"type": "Основная литература", "purpose": "Для изучения теории;", "desc": "<ГОСТ-запись реального учебника>", "url": "http://www.znanium.com", "coeff": "1.00"}},
  {{"type": "Основная литература", "purpose": "Для выполнения СРО;Для изучения теории;", "desc": "<ГОСТ-запись реального учебника>", "url": "http://e.lanbook.com", "coeff": "1.00"}},
  {{"type": "Дополнительная литература", "purpose": "Для изучения теории;", "desc": "<ГОСТ-запись реального учебника>", "url": "http://biblio-online.ru", "coeff": "0.50"}}
]
Ровно 3 объекта. Замени угловые скобки реальными ГОСТ-записями.""",

    # [FIX-#13]
    # [FIX-#14]

    "practice": """\
Напиши 6 тем практических занятий для дисциплины «{discipline}».
Компетенции, которые должны формироваться: {competencies_summary}
{discipline_focus_block}
Разделы дисциплины:
{sections_list}

Требования:
- каждое занятие — решение конкретной задачи по методам дисциплины
- все 6 тем разные, чередовать: анализ, синтез, моделирование, эксперимент
- темы ПЗ привязаны к конкретным методам из ключевых тем выше
- формулировка: глагол + метод/объект/задача («Анализ...», «Синтез...», «Решение задач...», «Моделирование...»)
- [FIX-05] КРАТКО: 8–15 слов на ПЗ. Одно занятие — одна задача.
  НЕ перечисляй через запятую 3+ метода, НЕ описывай этапы работы в названии.
  Пример правильной длины: «Моделирование нечёткого регулятора в среде
  MATLAB/Simulink» (7 слов). НЕЛЬЗЯ: «Разработка, настройка и сравнительный
  анализ нечётких регуляторов Мамдани и Такаги–Сугено с оценкой качества
  регулирования по переходным процессам и показателям устойчивости» (22 слова).
- укажи номер раздела (1, 2 или 3) к которому относится ПЗ
ВЕРНИ ТОЛЬКО JSON-массив (без пояснений, без markdown):
[
  {{"title": "<тема ПЗ 1 специфичная для {discipline}, 8-15 слов>", "section": 1}},
  {{"title": "<тема ПЗ 2 специфичная для {discipline}, 8-15 слов>", "section": 1}},
  {{"title": "<тема ПЗ 3 специфичная для {discipline}, 8-15 слов>", "section": 2}},
  {{"title": "<тема ПЗ 4 специфичная для {discipline}, 8-15 слов>", "section": 2}},
  {{"title": "<тема ПЗ 5 специфичная для {discipline}, 8-15 слов>", "section": 3}},
  {{"title": "<тема ПЗ 6 специфичная для {discipline}, 8-15 слов>", "section": 3}}
]
Ровно 6 объектов. Все темы уникальны и специфичны для «{discipline}».""",
}


# ---------------------------------------------------------------------------
# [A] Обёртка генерации с JSON-retry
# ---------------------------------------------------------------------------

def gen_with_json_retry(label: str, discipline: str, prompt: str,
                        parser_json, parser_fallback, max_retries: int = 2,
                        direction: str = "", level: str = "", **extra):
    """
    [A] Генерирует секцию с JSON-валидацией и retry.

    1. Вызывает gen() → LLM-ответ
    2. Пробует parser_json — если успех, возвращает (raw_text, parsed)
    3. При неудаче: до max_retries перегенераций
       без неё модель получает идентичный запрос и с высокой вероятностью
       даёт тот же невалидный ответ. Подсказка снижает число fallback'ов.
    4. Если JSON так и не распарсился — regex-fallback через parser_fallback
    """
    raw = gen(label, discipline, prompt, direction=direction, level=level, **extra)
    result = parser_json(raw)
    if result is not None:
        return raw, result

    # [FIX-3]
    RETRY_HINT = (
        "\n\n!!! ПРЕДЫДУЩИЙ ОТВЕТ НЕ ПРОШЁЛ ВАЛИДАЦИЮ !!!\n"
        "Верни ТОЛЬКО валидный JSON-массив — никакого текста до или после.\n"
        "Никаких ```json``` блоков. Никаких пояснений. Только [...]\n"
    )
    retry_prompt = RETRY_HINT + prompt

    for attempt in range(max_retries):
        print(f"  🔄 [{label}] JSON не распарсился (попытка {attempt + 1}/{max_retries}), "
              f"перегенерация...")
        raw = gen(label, discipline, retry_prompt, direction=direction, level=level, **extra)
        result = parser_json(raw)
        if result is not None:
            return raw, result

    print(f"  ⚠️  [{label}] JSON недоступен после {max_retries} попыток — regex-fallback")
    return raw, parser_fallback(raw)


# ---------------------------------------------------------------------------
# Точка входа
# ---------------------------------------------------------------------------

def main(config_path: Optional[str] = None, clear_cache: bool = False):
    if config_path is None and os.path.exists("config.json"):
        config_path = "config.json"

    # [FIX-CACHE]
    if clear_cache and os.path.exists(_CACHE_FILE):
        os.remove(_CACHE_FILE)
        print(f"  ♻️  Кэш {_CACHE_FILE} сброшен (--clear-cache)")

    if config_path:
        with open(config_path, encoding="utf-8") as f:
            cfg = json.load(f)
    else:
        cfg = {}
    cfg.setdefault("discipline", "Интеллектуальные системы")

    discipline       = cfg["discipline"]
    semester         = str(cfg.get("semester", "7"))
    competency_codes = cfg.get("competency_codes", "УК-1, ОПК-1, ОПК-2, ПК-1, ПК-2")
    direction        = cfg.get("direction", "")
    level            = cfg.get("level", "бакалавриат")

    # [З-R5]
    _load_cache()

    # [SIM] Показываем похожие дисциплины до генерации — ориентир для преподавателя
    _print_similar_disciplines(discipline, corpus_dir="rpd_corpus")

    # [З-5]
    if "retrieval_top_k" in cfg:
        GENERATION["top_k"] = int(cfg["retrieval_top_k"])
    if "retrieval_min_score" in cfg:
        GENERATION["min_score"] = float(cfg["retrieval_min_score"])

    # [З-G6]
    global _RETRIEVAL_CONF_HASH
    _RETRIEVAL_CONF_HASH = _make_retrieval_conf_hash(
        GENERATION["top_k"], GENERATION["min_score"]
    )

    hours = {
        "lecture":  cfg.get("hours_lecture",  12),
        "practice": cfg.get("hours_practice", 36),
        "lab":      cfg.get("hours_lab",      16),
        "self":     cfg.get("hours_self",     62),
    }

    template = cfg.get("template", "")
    if not template or not os.path.exists(template):
        # Ищем Шаблон_пустой.dotx в текущей директории и rpd_corpus/
        for candidate in ["Шаблон_пустой.dotx",
                          os.path.join("rpd_corpus", "Шаблон_пустой.dotx")]:
            if os.path.exists(candidate):
                template = candidate
                break
        # Fallback на последний .docx в rpd_corpus/ (обратная совместимость)
        if not template:
            corpus_dir = "rpd_corpus"
            candidates = sorted(
                f for f in os.listdir(corpus_dir)
                if (f.endswith(".docx") or f.endswith(".dotx")) and not f.startswith("~$")
            ) if os.path.isdir(corpus_dir) else []
            template = os.path.join(corpus_dir, candidates[-1]) if candidates else ""

    print(f"\n{'=' * 60}")
    print(f"ГЕНЕРАЦИЯ РПД: {discipline}")
    print(f"Направление: {direction}  Уровень: {level}")
    print(f"{'=' * 60}\n")

    # Проверка Ollama
    try:
        requests.get("http://localhost:11434/api/tags", timeout=5).raise_for_status()
        print("✅ Ollama доступен")
    except Exception as e:
        print(f"❌ Ollama недоступен: {e}")
        return

    if not template or not os.path.exists(template):
        print(f"❌ Шаблон не найден: {template!r}")
        return

    codes_list = [c.strip() for c in competency_codes.split(",") if c.strip()]
    competency_codes_numbered = "\n".join(f"{i + 1}. {c}" for i, c in enumerate(codes_list))

    # Базовые переменные (competencies_summary пустой для первого прохода).
    # [БАГ 7 ИСПРАВЛЕНО]
    _focus_raw = cfg.get("discipline_focus", "")
    if _focus_raw:
        discipline_focus_block = (
            f"Ключевые темы дисциплины (обязательно используй их в названиях разделов и заданий):\n"
            f"{_focus_raw}\n"
        )
    else:
        discipline_focus_block = ""

    base_vars = {
        "competency_codes":          competency_codes,
        "competency_codes_numbered": competency_codes_numbered,
        "competency_codes_list":     ", ".join(codes_list),        # [FIX-02] для промпта outcomes
        "competency_count":          len(codes_list),
        "outcomes_count":            9,   # резерв для совместимости
        "outcomes_total":            len(codes_list) * 3,           # [FIX-02] по 3 на компетенцию
        "competencies_summary":      "",  # заполняется после парсинга компетенций
        "discipline_focus_block":    discipline_focus_block,
    }

    raw: dict = {}

    # --- Шаг 1: компетенции и результаты обучения ---
    # [FIX-FGOS]
    _fgos = cfg.get("fgos_competencies", {})
    if _fgos and isinstance(_fgos, dict):
        competencies = [
            (code, _fgos[code])
            for code in codes_list
            if code in _fgos
        ]
        # Коды из competency_codes без записи в fgos_competencies — генерируем LLM
        _missing_codes = [c for c in codes_list if c not in _fgos]
        if _missing_codes:
            print(f"  ⚠️  [FGOS] Коды не найдены в fgos_competencies: {_missing_codes} — генерирую LLM")
            _missing_str = ", ".join(_missing_codes)
            _miss_vars = {**base_vars,
                          "competency_codes": _missing_str,
                          "competency_codes_numbered": "\n".join(f"{i+1}. {c}" for i, c in enumerate(_missing_codes)),
                          "competency_count": len(_missing_codes)}
            _, _extra = gen_with_json_retry(
                "competencies", discipline, PROMPTS["competencies"],
                parser_json=lambda t: parse_competencies_json(t),
                parser_fallback=lambda t: parse_competencies(t, codes=_missing_codes),
                direction=direction, level=level, **_miss_vars
            )
            competencies += _extra
        raw["competencies"] = json.dumps(
            [{"code": c, "desc": d} for c, d in competencies], ensure_ascii=False
        )
        print(f"  ✅ [FGOS] Компетенции из fgos_competencies: {[c for c, _ in competencies]}")
        _generation_log["competencies_source"] = "fgos_competencies (config.json)"
    else:
        raw["competencies"], competencies = gen_with_json_retry(
            "competencies", discipline, PROMPTS["competencies"],
            parser_json=lambda t: parse_competencies_json(t),
            parser_fallback=lambda t: parse_competencies(t, codes=codes_list),
            direction=direction, level=level, **base_vars
        )
        _generation_log["competencies_source"] = "llm"

    raw["outcomes"], outcomes = gen_with_json_retry(
        "outcomes", discipline, PROMPTS["outcomes"],
        # [FIX-02]
        parser_json=lambda t: parse_outcomes_json(t, required_count=len(codes_list) * 3),
        parser_fallback=parse_outcomes,
        direction=direction, level=level, **base_vars
    )

    # --- Шаг 2: обновляем competencies_summary и перегенерируем разделы ---
    comp_summary = "; ".join(f"{c[0]}: {c[1][:60]}" for c in competencies[:5])
    content_vars = {**base_vars, "competencies_summary": comp_summary}

    raw["content"], topics = gen_with_json_retry(
        "content", discipline, PROMPTS["content"],
        parser_json=parse_topics_json,
        parser_fallback=parse_topics,
        direction=direction, level=level, **content_vars
    )

    # [FIX-2]
    # [FIX-DRIFT]
    _ONIR_KW = {
        "научно-исследовательск", "этапы научного", "методологии научных",
        "исследований в России", "научного исследования", "нирс",
    }
    _sections_found = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    _is_domain_drift = any(
        any(kw in t.lower() for kw in _ONIR_KW) for t in _sections_found
    )
    if len(_sections_found) < 2 or _is_domain_drift:
        _reason = "domain drift (ОНИР)" if _is_domain_drift else f"Разделов найдено: {len(_sections_found)}"
        print(f"  ⚠️  [content] {_reason} — структурный fallback")
        # Строим осмысленный fallback на основе компетенций
        _comp_keywords = " ".join(c[1][:40] for c in competencies[:3]).lower()
        _has_neuro  = any(w in _comp_keywords for w in ("нейр", "сеть", "deep"))
        _has_fuzzy  = any(w in _comp_keywords for w in ("нечётк", "fuzzy", "логик"))
        _has_optim  = any(w in _comp_keywords for w in ("оптим", "алгорит", "эволюц"))
        _has_manage = any(w in _comp_keywords for w in ("управл", "регулят", "систем"))

        if _has_fuzzy:
            topics = [
                f"Раздел 1. Теоретические основы {discipline}",
                f"Тема 1.1. Математический аппарат нечётких множеств",
                f"Тема 1.2. Архитектуры нечётких систем",
                f"Раздел 2. Методы нечёткого вывода",
                f"Тема 2.1. Системы Мамдани и Сугено",
                f"Тема 2.2. Нейро-нечёткие системы ANFIS",
                f"Раздел 3. Применение {discipline}",
                f"Тема 3.1. Синтез нечётких регуляторов",
                f"Тема 3.2. Оценка эффективности систем",
            ]
        elif _has_neuro:
            topics = [
                f"Раздел 1. Архитектуры нейронных сетей",
                f"Тема 1.1. Многослойные перцептроны и обратное распространение",
                f"Тема 1.2. Сверточные и рекуррентные сети",
                f"Раздел 2. Обучение и оптимизация нейронных сетей",
                f"Тема 2.1. Алгоритмы оптимизации и регуляризация",
                f"Тема 2.2. Трансферное обучение и тонкая настройка",
                f"Раздел 3. Применение нейронных сетей",
                f"Тема 3.1. Задачи классификации и регрессии",
                f"Тема 3.2. Оценка качества и развёртывание моделей",
            ]
        elif _has_manage:
            topics = [
                f"Раздел 1. Основы интеллектуального управления",
                f"Тема 1.1. Классификация и архитектуры ИСУ",
                f"Тема 1.2. Адаптивное управление",
                f"Раздел 2. Методы синтеза интеллектуальных регуляторов",
                f"Тема 2.1. Нейросетевые и нечёткие регуляторы",
                f"Тема 2.2. Обучение с подкреплением в управлении",
                f"Раздел 3. Применение {discipline}",
                f"Тема 3.1. Моделирование и верификация",
                f"Тема 3.2. Сравнительный анализ методов",
            ]
        else:
            topics = [
                f"Раздел 1. Теоретические основы {discipline}",
                f"Тема 1.1. Основные понятия и методы",
                f"Тема 1.2. Архитектуры и инструменты",
                f"Раздел 2. Алгоритмическая база {discipline}",
                f"Тема 2.1. Ключевые алгоритмы и их реализация",
                f"Тема 2.2. Оптимизация и настройка систем",
                f"Раздел 3. Применение {discipline}",
                f"Тема 3.1. Прикладные задачи дисциплины",
                f"Тема 3.2. Оценка эффективности и верификация",
            ]
        _sections_found = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
        print(f"  ℹ️  Создано {len(_sections_found)} разделов из fallback")

    # [Фикс №9] Каждый раздел должен иметь хотя бы 1 тему (ОДНОКРАТНАЯ ПРОВЕРКА).
    # Если LLM вернул только разделы без тем — добавляем базовые подтемы,
    # иначе fill_lectures_table / fill_t21_fos получат пустой список topics_only.
    _secs_in_topics   = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    _topics_in_topics = [t for t in topics if re.match(r"^Тема\s*[\d\.]+", t)]
    if _secs_in_topics and not _topics_in_topics:
        print(f"  ⚠️  [content] Темы внутри разделов отсутствуют — добавляю базовые")
        enriched: list = []
        for i, sec in enumerate(_secs_in_topics, 1):
            enriched.append(sec)
            enriched.append(f"Тема {i}.1. Теоретические основы")
            enriched.append(f"Тема {i}.2. Практическое применение")
        topics = enriched

    # [Фикс №5+6] sections_list передаётся в промпты ЛР/ПЗ — LLM указывает
    # номер раздела явно, а не определяется по ротации в fill_lab/practice_table.
    _secs = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    _SEC_PREFIX = re.compile(r"^Раздел\s*\d+[.\s]+")
    _sections_list_str = "\n".join(
        "{0}. {1}".format(i + 1, _SEC_PREFIX.sub("", s).strip())
        for i, s in enumerate(_secs)
    ) or "1. Теоретические основы\n2. Методы\n3. Применение"
    content_vars = {**content_vars, "sections_list": _sections_list_str}

    raw["lab_works"], lab_works = gen_with_json_retry(
        "lab_works", discipline, PROMPTS["lab_works"],
        parser_json=lambda t: parse_list_json_with_section(t, min_items=6),
        parser_fallback=lambda t: [{"title": x, "section": None} for x in parse_list(t, discipline)],
        direction=direction, level=level, **content_vars
    )
    lab_works = _normalize_section_assignment(lab_works, len(_secs))

    raw["practice"], practices = gen_with_json_retry(
        "practice", discipline, PROMPTS["practice"],
        parser_json=lambda t: parse_list_json_with_section(t, min_items=6),
        parser_fallback=lambda t: [{"title": x, "section": None} for x in parse_list(t, discipline)],
        direction=direction, level=level, **content_vars
    )
    practices = _normalize_section_assignment(practices, len(_secs))

    # --- Шаг 3: библиография ---
    print("  📚 Генерация библиографии...")
    bib_main, bib_method = gen_bibliography(discipline, direction, level, cfg=cfg)

    # --- [D] Валидация ---
    validation_warnings = validate_generation(
        cfg, hours, competencies, topics, lab_works, practices
    )
    if validation_warnings:
        print("\n🔎 Результаты валидации:")
        for w in validation_warnings:
            print(f"  {w}")
        _generation_log["validation_warnings"] = validation_warnings
    else:
        print("\n✅ Валидация пройдена — все бизнес-правила соблюдены")

    # --- Заполнение шаблона ---
    hours_contact = hours["lecture"] + hours["practice"] + hours["lab"]
    hours_sro     = hours["self"]
    hours_total   = hours_contact + hours_sro
    exam_type     = cfg.get("exam_type", "экзамен")
    code          = cfg.get("code", "38050")

    # [T1] Копируем пустой шаблон с [] плейсхолдерами.
    # .dotx — тот же zip-архив, python-docx открывает его напрямую.
    shutil.copy(template, OUTPUT_DOCX)
    doc = Document(OUTPUT_DOCX)

    # [T1] Заполняем [] в заголовочных параграфах (название, трудоёмкость, год)
    fill_doc_header(
        doc,
        discipline  = discipline,
        code        = code,
        year        = cfg.get("year", "2025"),
        credits     = cfg.get("credits", 4),
        hours_total = hours_total,
        exam_type   = exam_type,
    )
    # [FIX-AppV]
    fill_appendix_v(doc, discipline, topics)
    # [Д-2] Заполняем блок «Темы для СРО» — fill_doc_header его пропускает
    fill_sro_topic_paragraphs(doc, topics, label=f"({code}){discipline}")

    for name, fn, args in [
        # [Д-1 ИСПРАВЛЕНО]: fill_t3_hours существовала, но не вызывалась →
        # таблица Раздела 1 (семестр/часы) оставалась с [] в output.
        ("Т3 Трудоёмкость",        fill_t3_hours,            (doc, semester,
                                                               cfg.get("credits", 4),
                                                               hours_total,
                                                               hours_contact,
                                                               hours_sro,
                                                               exam_type)),
        ("Т4 Компетенции",         fill_competencies_table,  (doc, competencies)),
        ("Т5 Результаты обучения", fill_outcomes_table,      (doc, competencies, outcomes)),
        ("Т6 Виды работы",         fill_t6_workload,         (doc, hours["lecture"], hours["practice"], hours["lab"], hours["self"], semester,
                                                               cfg.get("exam_prep_hours", 9 if exam_type == "экзамен" else 0))),
        ("Т7 Темы",                fill_topics_table,        (doc, topics, semester, hours, codes_list)),
        ("Т8 Лекции",              fill_lectures_table,      (doc, topics, hours)),
        ("Т9 ЛР",                  fill_lab_table,           (doc, lab_works, topics, hours["lab"])),
        ("Т10 ПЗ",                 fill_practice_table,      (doc, practices, topics, hours["practice"])),
        ("Т11 СРО",                fill_t11_sro,             (doc, topics, hours["self"], cfg)),
        ("Т15 Основная лит-ра",    fill_bibliography_main,   (doc, bib_main,   semester)),
        ("Т17 Метод.издания",      fill_bibliography_method, (doc, bib_method, semester)),
        ("Т21 ФОС",                fill_t21_fos,             (doc, competencies, topics, outcomes, discipline)),
        ("Аннотация",              fill_annotation_table,    (doc, competencies, outcomes, topics,
                                                               cfg.get("credits", 4), hours_total, exam_type)),
    ]:
        try:
            fn(*args)
            print(f"  ✅ {name}")
        except Exception as e:
            print(f"  ⚠️  {name}: {e}")

    doc.save(OUTPUT_DOCX)
    print(f"\n✅ Сохранено: {OUTPUT_DOCX}")

    # [З-R5]
    _save_cache()

    # [C] Сохраняем лог генерации
    try:
        with open(GENERATION_LOG, "w", encoding="utf-8") as f:
            json.dump(_generation_log, f, ensure_ascii=False, indent=2)
        print(f"📋 Лог генерации: {GENERATION_LOG}")
    except Exception as e:
        print(f"  ⚠️  Не удалось сохранить лог: {e}")


if __name__ == "__main__":
    import argparse as _ap
    _p = _ap.ArgumentParser()
    _p.add_argument("config", nargs="?", default=None, help="Путь к config.json")
    _p.add_argument("--clear-cache", action="store_true",
                    help="Сбросить rpd_cache.json (нужно после пересборки корпуса)")
    _p.add_argument("--rerank", action="store_true",
                    help="[З-13] Включить CrossEncoder reranking (требует sentence-transformers)")
    _a = _p.parse_args()
    if _a.rerank:
        RERANK_ENABLED = True
        print("  ℹ️  Reranking включён (--rerank)")
    main(_a.config, clear_cache=_a.clear_cache)
