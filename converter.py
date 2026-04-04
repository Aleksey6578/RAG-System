"""
converter.py — конвертация DOCX-файлов РПД в JSON-блоки.

Исправления v3:
  - [E] Структурированные таблицы: {"headers": [...], "rows": [...]}
  - [1] БАГ flush_buffer level. [2] БАГ split_into_chunks.
  - [3] Разбивка длинных таблиц. [4] Глобальный try/except.

Исправления v3.1:
  - [5] document_id (MD5) в каждом блоке.

Исправления v3.2:
  - [6] ИСПРАВЛЕНО: нормализация section_level в int.
    Замечание: "section_level = null — нестабильная иерархия. Лучше level=1/2/3."
    Теперь section_level — всегда int: 0 = неизвестен, 1–6 = уровень заголовка.
    _normalize_level("2.1") → 2, None → 0, raw int из стиля → без изменений.

  - [7] ИСПРАВЛЕНО: document_metadata вынесена из тела чанков.
    Замечание: "document_metadata внутри первого chunk — ошибка архитектуры.
    Правильная структура: {document: {...}, chunks: [...]}."
    process_document() теперь возвращает dict:
      {
        "document_id": "<md5>",
        "metadata":    { ...заголовок, автор, дата... },
        "chunks":      [ ...блоки без document_metadata... ]
      }
    prepare_texts.py обновлён для чтения обоих форматов (dict + list).

Исправления v3.3:
  - [8] ИСПРАВЛЕНО: несоответствие section_type между converter.py и chunking.py.
    Замечание: ключевые слова "литератур", "библиограф", "учебно-методич" были
    сопоставлены с типом "place" в SECTION_TYPE_MAP, тогда как classify_section()
    в chunking.py возвращал для тех же заголовков "bibliography".
    Поскольку build_metadata() отдаёт приоритет block_stype из converter, разделы
    с литературой оседали в Qdrant под тегом "place" и становились невидимы для
    retrieval-запросов с фильтром section_type = "bibliography".
    Исправление: разделены в отдельную запись с типом "bibliography", остальные
    ("ресурс", "библиотек", "программн", "информационн") → "place" как прежде.
"""

from docx import Document
from docx.document import Document as DocumentClass
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
import hashlib
import json
import re
from pathlib import Path
from typing import List, Dict, Optional, Tuple

RPD_CORPUS         = "rpd_corpus"
RPD_JSON           = "rpd_json"
MAX_CHUNK_WORDS    = 180
MIN_CHUNK_WORDS    = 20
MAX_HEADING_LENGTH = 300
MAX_TABLE_ROWS     = 30
SECTION_REGEX      = r"^(\d{1,2}(\.\d{1,2}){0,3})(?:[.\s]+)(.+)$"
KEY_HEADERS        = [
    "Цели дисциплины", "Формируемые компетенции",
    "Результаты обучения", "Содержание дисциплины",
]
SECTION_RE = re.compile(SECTION_REGEX)

SECTION_TYPE_MAP = [
    (["цел", "задач"],                                              "goals"),
    (["компетенц"],                                                 "competencies"),
    (["доступн", "инвалид", "огранич", "здоровь", "овз"],         "accessibility"),
    (["результат обучен", "индикатор"],                            "learning_outcomes"),
    (["содержани", "тематическ"],                                   "content"),
    (["лабораторн", "практическ", "семинар", "самостоятельн"],     "assessment"),
    (["трудоёмк", "трудоемк", "объём", "объем", "часов", "учебн", "нагрузк"], "hours"),
    (["фонд оценочн", "оценочн", "контрол", "аттестац",
      "промежуточн", "текущ"],                                      "assessment"),
    # [8] ИСПРАВЛЕНО: "литератур", "библиограф", "учебно-методич" выделены
    # в отдельный тип "bibliography" — согласовано с classify_section() в chunking.py.
    # Раньше они попадали в "place", что делало разделы литературы невидимыми
    # для retrieval-фильтров с section_type = "bibliography".
    # [З-C1] ИСПРАВЛЕНО: добавлены "учебной литератур", "обеспеченност", "сведени" —
    # заголовок «СВЕДЕНИЯ об обеспеченности дисциплины учебной литературой» не
    # содержал прежних ключевых слов → попадал в "other" → 0 чанков bibliography
    # в Qdrant → retrieval литературы всегда возвращал пустой контекст → fallback.
    (["литератур", "библиограф", "учебно-методич",
      "учебной литератур", "обеспеченност", "сведени"],             "bibliography"),
    (["ресурс", "библиотек", "программн", "информационн"],         "place"),
]


# ---------------------------------------------------------------------------
# Вспомогательные функции
# ---------------------------------------------------------------------------

def generate_doc_id(doc_path: Path) -> str:
    return hashlib.md5(doc_path.name.encode("utf-8")).hexdigest()


def _normalize_level(level_raw) -> int:
    """
    [6] Нормализует section_level → int в диапазоне [0, 6].

    None / 0 / "" → 0  (уровень неизвестен)
    int           → min(int, 6)
    "1"           → 1
    "2.1"         → 2  (глубина = кол-во компонентов через точку)
    "3.1.2"       → 3
    """
    if level_raw is None or level_raw == "":
        return 0
    if isinstance(level_raw, int):
        return min(level_raw, 6)
    s = str(level_raw).strip()
    # Строка вида "2.1.3" — считаем глубину
    parts = s.split(".")
    # Если первая часть — не цифра (e.g. "heading"), ищем первую цифру
    if not parts[0].isdigit():
        m = re.search(r"\d+", s)
        return min(int(m.group()), 6) if m else 0
    return min(len(parts), 6)


def detect_section_type(section_title: Optional[str]) -> str:
    if not section_title:
        return "other"
    t = section_title.lower()
    for keywords, stype in SECTION_TYPE_MAP:
        if any(kw in t for kw in keywords):
            return stype
    return "other"


# ---------------------------------------------------------------------------
# [E] Структурированные таблицы
# ---------------------------------------------------------------------------

def _is_numeric_row(cells: list) -> bool:
    """
    [З-C3] Проверяет, является ли строка «числовой шапкой» — артефактом
    объединённых ячеек в DOCX (строки вида «1 | 2 | 3 | 5 | 5 | 6 | 7»).
    Такие строки появляются как вторая строка в таблицах со сложной шапкой
    (например, «СВЕДЕНИЯ об обеспеченности»), где первая строка — реальные
    заголовки, а вторая — сквозная нумерация столбцов.
    Критерий: все непустые ячейки содержат только цифры.
    """
    non_empty = [c.strip() for c in cells if c.strip()]
    if not non_empty:
        return False
    return all(re.match(r"^\d+$", c) for c in non_empty)


def process_table(table: Table) -> Dict:
    raw_rows = []
    for row in table.rows:
        cells = []
        seen_tcs: set = set()  # [БАГ 1 ИСПРАВЛЕНО]: дедупликация merged-ячеек
        for cell in row.cells:
            if cell._tc in seen_tcs:
                continue
            seen_tcs.add(cell._tc)
            cell_text = " ".join(
                p.text.strip() for p in cell.paragraphs if p.text.strip()
            )
            cells.append(cell_text)
        if any(c.strip() for c in cells):
            raw_rows.append(cells)
    if not raw_rows:
        return {"headers": [], "rows": []}

    # [З-C3] ИСПРАВЛЕНО: первая строка — заголовки. Вторая строка отфильтровывается
    # если она представляет собой числовую нумерацию столбцов (1|2|3|5|5|6|7|9).
    # Такой артефакт типичен для таблиц с объединёнными ячейками в шапке.
    headers   = raw_rows[0]
    data_rows = raw_rows[1:]
    if data_rows and _is_numeric_row(data_rows[0]):
        data_rows = data_rows[1:]

    return {"headers": headers, "rows": data_rows}


def table_to_text(table_data: Dict) -> str:
    headers = table_data.get("headers", [])
    rows    = table_data.get("rows", [])
    if not headers and not rows:
        return ""
    lines = []
    if headers:
        lines.append(" | ".join(str(h) for h in headers))
        lines.append("-" * max(len(lines[0]), 20))
    for row in rows:
        if any(str(c).strip() for c in row):
            lines.append(" | ".join(str(c) for c in row))
    return "\n".join(lines)


def extract_key_table_rows(
    table: Table, section_type: str, doc_name: str, section_title: str,
    document_id: str = "",
    section_level: int = 0,
) -> List[Dict]:
    # [З-C2] ИСПРАВЛЕНО: добавлен тип "bibliography" — строки таблицы литературы
    # (отдельная книга = отдельная строка) теперь извлекаются как отдельные
    # table_row чанки вместо одного монолитного table-блока.
    if section_type not in ("competencies", "learning_outcomes", "content", "assessment",
                            "bibliography"):
        return []
    table_data = process_table(table)
    headers = table_data.get("headers", [])
    rows    = table_data.get("rows", [])
    if not rows:
        return []

    # [З-C4] ИСПРАВЛЕНО: заголовок таблицы добавляется только к первому чанку.
    # Прежде header_line повторялся в КАЖДОЙ строке → 8 чанков content начинались
    # с одинакового «№ пп. | Номер раздела | Название темы | Трудоемкость, часы».
    # Это засоряло retrieval одинаковыми префиксами и снижало эффективность
    # дедупликации по тексту в chunking.py.
    # Новое поведение: первый чанк содержит «заголовок + строка», последующие —
    # только строки данных, без повторного заголовка.
    header_line = " | ".join(str(h) for h in headers) if headers else ""
    blocks = []
    for idx, row_cells in enumerate(rows):
        row_text = " | ".join(str(c) for c in row_cells)
        if not row_text.strip() or len(row_text.split()) < 5:
            continue
        # Заголовок только у первого содержательного чанка
        if header_line and idx == 0:
            text = f"{header_line}\n{row_text}"
        else:
            text = row_text
        row_lower = row_text.lower()
        effective_type = "learning_outcomes" if any(
            kw in row_lower for kw in ("знать:", "уметь:", "владеть:", "з(", "у(", "в(")
        ) else section_type
        blocks.append({
            "document_id":   document_id,
            "title":         doc_name,
            "section_title": section_title,
            "section_level": section_level,
            "section_type":  effective_type,
            "text":          text,
            "type":          "table_row",
        })
    return blocks


def extract_document_metadata(doc: Document) -> Dict:
    """[7] Метаданные документа — хранятся НА ВЕРХНЕМ УРОВНЕ JSON, не в чанке."""
    core = doc.core_properties
    return {
        "title":            core.title    or "",
        "subject":          core.subject  or "",
        "author":           core.author   or "",
        "keywords":         core.keywords or "",
        "created":          str(core.created)  if core.created  else "",
        "modified":         str(core.modified) if core.modified else "",
        "paragraphs_count": len(doc.paragraphs),
        "tables_count":     len(doc.tables),
    }


def is_section_heading(text: str, style_name: Optional[str] = None) -> Tuple[bool, int]:
    text = text.strip()
    if style_name:
        sl = style_name.lower()
        if "heading" in sl or "заголовок" in sl:
            m = re.search(r"(\d+)", style_name)
            return True, int(m.group(1)) if m else 1
    if re.match(r"^\d{2}\.\d{2}\.\d{4}", text): return False, 0
    if re.match(r"^[\d\.]+$", text):             return False, 0
    if len(text) > MAX_HEADING_LENGTH:           return False, 0
    m = SECTION_RE.match(text)
    if m:
        return True, min(m.group(1).count(".") + 1, 6)
    for kw in KEY_HEADERS:
        if text.lower().startswith(kw.lower()):
            return True, 1
    if text.isupper() and len(text) < 100 and not text.endswith("."):
        return True, 2
    return False, 0


def table_to_blocks(
    table: Table,
    doc_name: str,
    section: str,
    section_level: int,
    document_id: str = "",
) -> List[Dict]:
    table_data = process_table(table)
    headers    = table_data.get("headers", [])
    data_rows  = table_data.get("rows", [])
    if not headers and not data_rows:
        return []
    blocks = []
    if not data_rows:
        text = table_to_text(table_data)
        if len(text.split()) >= MIN_CHUNK_WORDS:
            blocks.append({
                "document_id":   document_id,
                "title":         doc_name,
                "section_title": section,
                "section_level": section_level,
                "text":          "[ТАБЛИЦА]\n" + text,
                "type":          "table",
                "table_data":    table_data,
            })
        return blocks
    for i in range(0, max(len(data_rows), 1), MAX_TABLE_ROWS):
        chunk_rows = data_rows[i: i + MAX_TABLE_ROWS]
        chunk_data = {"headers": headers, "rows": chunk_rows}
        text = "[ТАБЛИЦА]\n" + table_to_text(chunk_data)
        if len(text.split()) >= MIN_CHUNK_WORDS:
            blocks.append({
                "document_id":   document_id,
                "title":         doc_name,
                "section_title": section,
                "section_level": section_level,
                "text":          text,
                "type":          "table",
                "table_data":    chunk_data,
            })
    return blocks


def clean_text(text: str) -> str:
    text = text.replace("\xa0", " ").replace("\t", " ")
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_into_chunks(paragraphs: List[str], max_words: int = MAX_CHUNK_WORDS) -> List[str]:
    chunks: List[str] = []
    buffer: List[str] = []
    count = 0
    for p in paragraphs:
        p = clean_text(p)
        if not p:
            continue
        words = p.split()
        wc    = len(words)
        if wc > max_words:
            # [БАГ 6 ИСПРАВЛЕНО]: если buffer не пустой — сначала сбрасываем его,
            # затем обрабатываем большой абзац по предложениям.
            # Раньше: if wc > max_words and not buffer — при непустом buffer
            # огромный абзац добавлялся в buffer без разбивки → чанк мог быть
            # в несколько раз больше max_words.
            if buffer:
                chunks.append("\n\n".join(buffer))
                buffer = []
                count = 0
            sentence_buffer: List[str] = []
            sentence_count = 0
            for sentence in re.split(r"(?<=[.!?])\s+", p):
                sentence = sentence.strip()
                if not sentence:
                    continue
                sentence_buffer.append(sentence)
                sentence_count += len(sentence.split())
                if sentence_count >= max_words:
                    chunks.append(" ".join(sentence_buffer))
                    sentence_buffer = []
                    sentence_count = 0
            if sentence_buffer:
                chunks.append(" ".join(sentence_buffer))
            continue
        count += wc
        buffer.append(p)
        if count >= max_words:
            chunks.append("\n\n".join(buffer))
            buffer = []
            count = 0
    if buffer:
        chunks.append("\n\n".join(buffer))
    return chunks


def iter_block_items(parent):
    body = parent.element.body if isinstance(parent, DocumentClass) else parent._element
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield "table", Table(child, parent)


def process_document(doc_path: Path) -> Dict:
    """
    [7] Возвращает структуру:
      {
        "document_id": "<md5>",
        "metadata":    { title, author, created, ... },
        "chunks":      [ { document_id, title, section_title,
                           section_level (int), section_type,
                           text, type, [table_data] } ]
      }
    """
    doc      = Document(doc_path)
    doc_id   = generate_doc_id(doc_path)
    metadata = extract_document_metadata(doc)

    chunks:          List[Dict]    = []
    current_section: Optional[str] = None
    current_level:   int           = 0
    current_stype:   str           = "other"
    buffer:          List[str]     = []

    def flush_buffer():
        nonlocal buffer
        for chunk in split_into_chunks(buffer):
            if len(chunk.split()) >= MIN_CHUNK_WORDS:
                chunks.append({
                    "document_id":   doc_id,
                    "title":         doc_path.name,
                    "section_title": current_section,
                    "section_level": current_level,   # [6] int
                    "section_type":  current_stype,
                    "text":          chunk,
                    "type":          "text",
                })
        buffer = []

    for item_type, item in iter_block_items(doc):
        if item_type == "paragraph":
            text = item.text.strip()
            if not text:
                continue
            style_name = item.style.name if item.style else None
            is_heading, raw_level = is_section_heading(text, style_name)
            if is_heading:
                if buffer and current_section:
                    flush_buffer()
                m = SECTION_RE.match(text)
                if m:
                    current_section = m.group(3).strip()
                    current_level   = _normalize_level(m.group(1))  # [6]
                else:
                    current_section = text.strip()
                    current_level   = _normalize_level(raw_level)   # [6]
                current_stype = detect_section_type(current_section)
            else:
                buffer.append(text)

        elif item_type == "table":
            # [БАГ 8 ИСПРАВЛЕНО]: убрано условие `and current_section`.
            # Таблицы до первого заголовка (титульная страница РПД) больше не теряются.
            if buffer:
                flush_buffer()
            _sec = current_section or "Введение"
            stype = detect_section_type(_sec)
            row_blocks = extract_key_table_rows(
                item, stype, doc_path.name, _sec,
                document_id=doc_id, section_level=current_level,
            )
            if row_blocks:
                chunks.extend(row_blocks)
            else:
                for b in table_to_blocks(
                    item, doc_path.name, _sec,
                    section_level=current_level,
                    document_id=doc_id,
                ):
                    b["section_type"] = stype
                    chunks.append(b)

    if buffer and current_section:
        flush_buffer()

    # [7] document_metadata — на верхнем уровне, НЕ в chunks[0]
    return {
        "document_id": doc_id,
        "metadata":    metadata,
        "chunks":      chunks,
    }


def main():
    output_dir = Path(RPD_JSON)
    output_dir.mkdir(exist_ok=True)
    docx_files = [
        f for f in Path(RPD_CORPUS).glob("*.docx")
        if not f.name.startswith("~$")
    ]
    ok = errors = 0
    for doc_path in sorted(docx_files):
        try:
            result = process_document(doc_path)
            chunks = result.get("chunks", [])
            if chunks:
                out_path = output_dir / doc_path.with_suffix(".json").name
                with open(out_path, "w", encoding="utf-8") as f:
                    json.dump(result, f, ensure_ascii=False, indent=2)
                ok += 1
                print(f"  ✅ {doc_path.name} → {len(chunks)} блоков")
            else:
                print(f"  ⚠️  {doc_path.name} → блоки не извлечены")
        except Exception as e:
            errors += 1
            print(f"  ❌ {doc_path.name}: {e}")
    print(f"\nКонвертация завершена. Успешно: {ok}, ошибок: {errors}")


if __name__ == "__main__":
    main()
