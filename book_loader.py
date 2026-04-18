"""
book_loader.py — загрузка учебников из rpd_books/ в pipeline.
"""

import argparse
import hashlib
import json
import re
import time
from pathlib import Path

import requests
from docx import Document
from utils import get_embedding as _get_embedding

try:
    import fitz
except ImportError:
    fitz = None

BOOKS_DIR = Path("rpd_books")
CONFIG_PATH = Path("config.json")
# [FIX-SYNC] Синхронизировано с chunking.py (было 300/50 — расхождение ~25%)
CHUNK_TOKENS = 400
OVERLAP_TOKENS = 60

QDRANT_URL = "http://localhost:6333"
COLLECTION = "rpd_rag"
UPSERT_BATCH = 64
RETRY_COUNT = 3
RETRY_DELAY = 2.0

# [FIX-SYNC] Токенайзер bge-m3 вместо len//4, аналогично chunking.py §13.2.
# Fallback: transformers (bge-m3) → len×0.375 (≈ 1.5 слова/токен для русского).
try:
    from transformers import AutoTokenizer as _AutoTok
    _tokenizer = _AutoTok.from_pretrained("BAAI/bge-m3", use_fast=True)
    def _approx_tokens(text: str) -> int:
        return len(_tokenizer.encode(text, add_special_tokens=False))
except Exception:
    _tokenizer = None
    def _approx_tokens(text: str) -> int:  # type: ignore[misc]
        return int(len(text) * 0.375)  # ~2.67 симв./токен для русского

_BIBLIO_RE = re.compile(
    r"(?P<authors>[А-ЯA-Z][^.]+?)\.\s+"
    r"(?P<title>[^/]+?)\s*/\s*"
    r"(?P<rest>.+?)\.\s*—\s*"
    r"(?P<city>[^:]+?)\s*:\s*"
    r"(?P<publisher>[^,]+),\s*"
    r"(?P<year>\d{4})",
    re.DOTALL,
)


def extract_metadata_from_filename(path: Path) -> dict:
    stem = path.stem
    m = _BIBLIO_RE.search(stem)
    if m:
        desc = f"{m['authors']}. {m['title'].strip()} — {m['city'].strip()} : {m['publisher'].strip()}, {m['year']}."
        return {"desc": desc, "year": m["year"], "raw": True}
    return {"desc": stem, "year": "", "raw": False}


def extract_metadata_from_docx(path: Path) -> dict:
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs[:20] if p.text.strip())
    m = _BIBLIO_RE.search(text)
    if m:
        desc = f"{m['authors']}. {m['title'].strip()} — {m['city'].strip()} : {m['publisher'].strip()}, {m['year']}."
        return {"desc": desc, "year": m["year"], "raw": True}
    return extract_metadata_from_filename(path)


def extract_metadata_from_pdf(path: Path) -> dict:
    if fitz is None:
        return extract_metadata_from_filename(path)
    doc = fitz.open(str(path))
    text = doc[0].get_text() if len(doc) > 0 else ""
    doc.close()
    m = _BIBLIO_RE.search(text)
    if m:
        desc = f"{m['authors']}. {m['title'].strip()} — {m['city'].strip()} : {m['publisher'].strip()}, {m['year']}."
        return {"desc": desc, "year": m["year"], "raw": True}
    return extract_metadata_from_filename(path)


def load_book_metadata(path: Path) -> dict:
    ext = path.suffix.lower()
    if ext == ".pdf":
        meta = extract_metadata_from_pdf(path)
    elif ext == ".docx":
        meta = extract_metadata_from_docx(path)
    else:
        meta = extract_metadata_from_filename(path)
    meta["source_file"] = str(path)
    return meta


try:
    import pytesseract as _pytesseract
    from PIL import Image as _PILImage
    import io as _io
    _OCR_AVAILABLE = True
except ImportError:
    _OCR_AVAILABLE = False


def extract_text_pdf(path: Path) -> str:
    if fitz is None:
        print(f"  ⚠️  Пропуск {path.name} — PyMuPDF не установлен")
        return ""
    doc = fitz.open(str(path))
    pages = [page.get_text() for page in doc]
    text = "\n".join(pages)

    # [FIX-OCR] Если fitz вернул пустой текст — PDF сканированный (нет OCR-слоя).
    # Fallback: растеризация страниц через PyMuPDF + pytesseract (rus).
    # Хайкин «Нейронные сети» — сканированный PDF, без fallback = 0 чанков.
    if not text.strip() and _OCR_AVAILABLE:
        print(f"  ⚠️  {path.name}: текстовый слой пуст — пробуем OCR (pytesseract)...")
        ocr_pages = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = _PILImage.open(_io.BytesIO(pix.tobytes("png")))
            try:
                ocr_pages.append(_pytesseract.image_to_string(img, lang="rus"))
            except Exception as e:
                print(f"    ⚠️  OCR страница {page.number}: {e}")
        text = "\n".join(ocr_pages)
        if text.strip():
            print(f"  ✅ OCR: извлечено {len(text)} симв. из {path.name}")
        else:
            print(f"  ❌ OCR не дал результата для {path.name} — нужна DOCX-версия")
    elif not text.strip() and not _OCR_AVAILABLE:
        print(f"  ⚠️  {path.name}: текстовый слой пуст. "
              f"Установите pytesseract+Pillow для OCR или предоставьте DOCX-версию.")

    doc.close()
    return text


def extract_text_docx(path: Path) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_pdf(path)
    elif ext == ".docx":
        return extract_text_docx(path)
    return ""


def chunk_text(text: str, source_file: str) -> list[dict]:
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks = []
    buf, buf_tokens, idx = [], 0, 0

    for para in paragraphs:
        t = _approx_tokens(para)
        if buf_tokens + t > CHUNK_TOKENS and buf:
            chunks.append({
                "id": f"{Path(source_file).stem}_chunk_{idx}",
                "text": " ".join(buf),
                "stype": "book_content",
                "content_type": "book",  # [З-07] для индекса content_type: keyword в Qdrant
                "source_file": source_file,
            })
            overlap_buf, overlap_tok = [], 0
            for sent in reversed(buf):
                overlap_buf.insert(0, sent)
                overlap_tok += _approx_tokens(sent)
                if overlap_tok >= OVERLAP_TOKENS:
                    break
            buf, buf_tokens = overlap_buf, sum(_approx_tokens(s) for s in overlap_buf)
            idx += 1
        buf.append(para)
        buf_tokens += t

    if buf:
        chunks.append({
            "id": f"{Path(source_file).stem}_chunk_{idx}",
            "text": " ".join(buf),
            "stype": "book_content",
            "content_type": "book",  # [З-07]
            "source_file": source_file,
        })
    return chunks


def build_biblio_entry(meta: dict, btype: str = "Основная литература") -> dict:
    return {
        "type": btype,
        "purpose": "Для изучения теории;Для выполнения СРО;",
        "desc": meta["desc"],
        "url": "http://bibl.rusoil.net",
        "coeff": "1.00",
    }


def update_config(entries: list[dict]):
    cfg = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    cfg["main_bibliography"] = entries
    CONFIG_PATH.write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"✅ config.json обновлён: {len(entries)} записей в main_bibliography")


def upsert_batch(ids: list, vectors: list, payloads: list) -> tuple[bool, list]:
    body = {"batch": {"ids": ids, "vectors": vectors, "payloads": payloads}}
    try:
        r = requests.put(f"{QDRANT_URL}/collections/{COLLECTION}/points", json=body, timeout=60)
        if r.status_code == 206:
            failed = r.json().get("result", {}).get("failed", [])
            return True, [f["id"] for f in failed] if failed else []
        if r.status_code != 200:
            print(f"  Ошибка upsert: {r.status_code} {r.text[:300]}")
            return False, ids
        return True, []
    except Exception as e:
        print(f"  Исключение при upsert: {e}")
        return False, ids


def upsert_batch_with_retry(ids: list, vectors: list, payloads: list) -> bool:
    delay = RETRY_DELAY
    for attempt in range(1, RETRY_COUNT + 1):
        ok, failed_ids = upsert_batch(ids, vectors, payloads)
        if ok and not failed_ids:
            return True
        if ok and failed_ids:
            print(f"  Retry {len(failed_ids)} failed points (попытка {attempt}/{RETRY_COUNT})...")
            failed_set = set(failed_ids)
            retry_pairs = [(i, v, p) for i, v, p in zip(ids, vectors, payloads) if i in failed_set]
            ids = [x[0] for x in retry_pairs]
            vectors = [x[1] for x in retry_pairs]
            payloads = [x[2] for x in retry_pairs]
            time.sleep(delay)
            delay *= 2
            continue
        if attempt < RETRY_COUNT:
            print(f"  Retry upsert {attempt}/{RETRY_COUNT}... ждём {delay:.0f}с")
            time.sleep(delay)
            delay *= 2
    print(f"  ❌ {len(ids)} точек не загружены после {RETRY_COUNT} попыток")
    return False


def load_chunks_to_qdrant(all_chunks: list[dict]):
    print(f"  Загрузка {len(all_chunks)} книжных чанков в Qdrant (HTTP)...")
    time.sleep(5)

    points = []
    for i, chunk in enumerate(all_chunks):
        vec = _get_embedding(chunk["text"], prefix="passage", retry=RETRY_COUNT)
        if not vec:
            print(f"  ⚠️  Пропуск чанка {chunk['id']} — embedding не получен")
            continue

        # Числовой ID через SHA256
        point_id = int(hashlib.sha256(chunk["id"].encode()).hexdigest()[:15], 16)

        payload = {
            "chunk_id": chunk["id"],
            "id": point_id,
            "doc_id": Path(chunk["source_file"]).stem,
            "source": chunk["source_file"],
            "source_file": chunk["source_file"],
            "section_title": "",
            "section_level": "",
            "doc_position": 0,
            "text": chunk["text"],
            "section_type": chunk["stype"],
            "content_type": chunk.get("content_type", "book"),  # [З-07]
            "metadata": {"section_type": chunk["stype"]},
            "direction": "",
            "level": "",
            "department": "",
            "embedding_model": "bge-m3",
        }
        points.append({"id": point_id, "vector": vec, "payload": payload})
        if (i + 1) % 10 == 0:
            print(f"    embedded {i+1}/{len(all_chunks)}")

    if not points:
        print("  ⚠️  Нет точек для загрузки")
        return

    uploaded = 0
    for batch_start in range(0, len(points), UPSERT_BATCH):
        batch = points[batch_start: batch_start + UPSERT_BATCH]
        ids = [p["id"] for p in batch]
        vectors = [p["vector"] for p in batch]
        payloads = [p["payload"] for p in batch]

        if upsert_batch_with_retry(ids, vectors, payloads):
            uploaded += len(batch)
            print(f"    Загружено: {uploaded}/{len(points)}")
        else:
            print(f"    ❌ Не удалось загрузить батч {batch_start}-{batch_start+len(batch)}")

    print(f"✅ Загружено в Qdrant: {uploaded} чанков (stype=book_content)")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--meta-only", action="store_true")
    args = parser.parse_args()

    books = sorted(BOOKS_DIR.glob("*.*"))
    books = [b for b in books if b.suffix.lower() in (".pdf", ".docx")]

    if not books:
        print(f"⚠️  Папка {BOOKS_DIR} пуста или не найдена")
        return

    print(f"📚 Найдено книг: {len(books)}")

    all_entries = []
    all_chunks = []

    for path in books:
        print(f"  → {path.name}")
        meta = load_book_metadata(path)
        all_entries.append(build_biblio_entry(meta))

        if not args.meta_only:
            text = extract_text(path)
            chunks = chunk_text(text, str(path))
            all_chunks.extend(chunks)
            print(f"     {len(chunks)} чанков")

    update_config(all_entries)

    if not args.meta_only and all_chunks:
        load_chunks_to_qdrant(all_chunks)


if __name__ == "__main__":
    main()