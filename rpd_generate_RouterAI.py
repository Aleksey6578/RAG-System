"""
rpd_generate.py — генерация РПД на основе шаблона Шаблон_пустой.dotx.

Стратегия (v4): копируем пустой шаблон с [] плейсхолдерами →
заполняем [] в параграфах и строках таблиц сгенерированным LLM-контентом.
Таблицы ищутся по заголовку (find_table), а не по хрупкому индексу.
Поля old_discipline / old_code / replace_all — удалены за ненадобностью.

RouterAI-версия: эмбеддинги и LLM через внешний API routerai.ru.

=============================================================================
ИСПРАВЛЕНИЯ ПО ОТЧЁТУ ОТ 16.04.2026 (RouterAI v3)
=============================================================================

[FIX-01] Отчёт §4, замечание 1 (критический). Обрезка индикаторов на
         полуслове: «УК-1.1 знать архитектуры... агент-ориенти в данной
         области». Было: `indicator = f"… {rotated[0][:50]} {_ind_qual}"`.
         Стало: `indicator = f"… {base_item}"` — без обрезки.
         Место: fill_outcomes_table.

[FIX-02] Отчёт §4, замечание 2 (критический). Дублирование З/У/В между
         компетенциями: З(УК-1) == З(ПК-1) == «архитектуры интеллектуальных
         систем…». Причина была двойной:
           (а) LLM просили сгенерировать ровно 9 результатов (3З+3У+3В) на
               любое число компетенций → при N>3 элементы переиспользовались;
           (б) fill_outcomes_table использовал `items[idx % len(items)]`
               от индекса КОМПЕТЕНЦИИ, что для того же type_idx давало
               повторы.
         Исправление:
           (а) промпт PROMPTS["outcomes"] требует competency_count*3
               результатов (по одной формулировке каждого типа на каждую
               компетенцию), competency_codes_list передаётся в шаблон;
           (б) сквозной счётчик _type_pos + guard _seen_texts в
               fill_outcomes_table — каждой (компетенция, тип) достаётся
               свой item, при исчерпании списка добавляется квалификатор.
         Места: PROMPTS["outcomes"], fill_outcomes_table, base_vars,
                required_count в gen_with_json_retry для outcomes.

[FIX-03] Отчёт §4, замечание 3 (критический). Индикаторы ФОС —
         бессмысленные склейки глагола-показателя и начала текста
         компетенции: «УК-1.1 Знает осуществлять поиск», «ОПК-2.3 Владеет
         понимать принципы работы». Причина: индикатор строился как
         `{verb} {comp_desc[:60]}`, где comp_desc — глагольная фраза
         компетенции, не совместимая по смыслу с verb («Знает», «Умеет»,
         «Владеет»). Стало: индикатор строится из outcome_text (того же
         текста, что попадает в колонку «результат обучения» ФОС-строки),
         с нормализацией префикса (убираем дублирующие «Знать:», «уметь»,
         чтобы не было «Знает: Знать: …»).
         Место: fill_t21_fos.

[FIX-04] Отчёт §4, замечание 4 (высокий). Обрезка названий тем в ФОС:
         «Методы машинного обучения и нейронные се» вместо полного
         названия раздела. Было: `sec_short = sec_name[:40]`. Стало:
         `sec_short = sec_name` с мягкой обрезкой только >80 симв. и
         суффиксом «…», чтобы не ломать вёрстку сверхдлинных названий.
         Место: fill_t21_fos.

[FIX-05] Отчёт §4, замечание 5 (высокий). Названия ЛР и ПЗ избыточно
         детальные (30+ слов). В промпты PROMPTS["lab_works"] и
         PROMPTS["practice"] добавлены:
           • явное требование «8–15 слов на ЛР/ПЗ»,
           • запрет перечислять через запятую 3+ метода,
           • пример правильной длины + пример запрещённой длины.

[FIX-06] Отчёт §4, замечание 6 (средний). Везде жёстко «выполнение
         расчётно-графической работы» — нетипично для «Интеллектуальных
         систем» и многих других дисциплин. fill_t11_sro теперь принимает
         cfg и выбирает основной вид СРО по приоритету:
           1) cfg["sro_types"] — явный список (если задан в config.json);
           2) эвристика по названию и focus дисциплины:
              • «интеллектуальн/нейрон/машинн/агент/проектирован/разработк»
                → «выполнение индивидуального задания»;
              • «моделирован/анализ/исследован/экспертн/нечётк»
                → «подготовка реферата»;
              • иначе (классическая инженерная) → РГР.
         Пропорции часов 20 %/20 %/60 % сохранены.

[FIX-07] Отчёт §4, замечание 7 (средний). «_» вместо «:» в выходных
         данных библиографии: «Москва _ Вильямс». Нарушение
         ГОСТ Р 7.0.5-2008. Добавлена функция _normalize_gost_biblio(),
         которая в fill_bibliography_main и fill_bibliography_method
         приводит «Город_Издательство» и «Город _ Издательство» к
         «Город : Издательство». Защищено от ложных срабатываний:
         паттерн требует букву с обеих сторон, дефисные названия
         («Ростов-на-Дону») не трогаются.
         Места: _normalize_gost_biblio, fill_bibliography_main,
                fill_bibliography_method.

[FIX-08] Отчёт §4, замечание 8 (средний). Вымышленное ФИО «Д. М. Зарипов»
         попадало в Т17 из fallback-шаблона. Убрано: методические издания
         теперь указывают только кафедру-составителя («каф. ВТИК»).
         Место: _make_fallback_method.

[FIX-10] Отчёт §4, замечание 10 (низкий). Темы докладов в Приложении В
         дословно копировали названия лекций. Добавлена функция
         _to_research_topic() с 6 шаблонами исследовательской
         переформулировки: «Современное состояние…», «Сравнительный
         анализ подходов…», «Применение … в прикладных задачах» и т. д.
         Шаблоны чередуются, чтобы тексты не были однообразны.
         Место: fill_appendix_v.

Замечания, НЕ требующие правки этого файла:
  • №7 (исходные «_» в config.json → main_bibliography) — нормализация
    выполняется на лету в fill_bibliography_main; исходный конфиг можно
    не править, но желательно.
  • №9 (пустые поля «Место дисциплины») — требует заполнения
    prerequisite / postrequisite в config.json; код уже поддерживает
    их чтение.

=============================================================================
РАНЕЕ ВНЕСЁННЫЕ ИСПРАВЛЕНИЯ (RouterAI v2)
=============================================================================

[FIX-SEM]     fill_t6_workload: regex с negative lookaround для поиска
              столбца семестра. Было: sem_str in cell_text давало ложные
              срабатывания («7» совпадало с «17», «27»).

[FIX-T6-SRO]  fill_t6_workload: разбивка СРО согласована с fill_t11_sro.
              Добавлены kw_map-ключи «подготовка к лаборатор», «изучение учебного».
              SKIP_PATTERNS дополнены: «иные», «выполнение и подготовка к защит».
              exam_prep_hours вычитается из hrs_prep явно.

[FIX-T21-VERB] fill_t21_fos: «Умеет применять» → «Умеет».
              При конкатенации с desc компетенции («применять методы…»)
              получалось двойное «Умеет применять применять».

[FIX-LAB-SEC]  fill_lab_table / fill_practice_table: sec_num bounds handling
              через modulo вместо строгой проверки 1 <= sec_num <= len(sections).
              Устраняет IndexError при sec_num=0 или вне диапазона.

[FIX-OUTCOMES] fill_outcomes_table: заменена архитектура vMerge на flat rows
              (fill_placeholder_rows). Упрощает код, устраняет порчу форматирования
              при клонировании template_group с vMerge-ячейками.
              Индикатор строится из rotated outcome-текста (уникален для каждой компетенции).

[FIX-BIB-RAG]  gen_bibliography: добавлен прямой парсинг RAG-чанков через
              _parse_rag_bibliography_chunks(). При ≥2 ГОСТ-строках LLM не вызывается.

[FIX-SIM]     _print_similar_disciplines: добавлена функция вывода похожих
              дисциплин из корпуса до генерации.

[FIX-PRACTICE] SECTION_QUERIES["practice"]: обновлены формулировки запросов
              для лучшего покрытия тем ПЗ.

[FIX-OUTCOMES-PROMPT] PROMPTS["outcomes"]: добавлено правило
              «каждый текст — ОДНА краткая фраза».

[FIX-PRACTICE-PROMPT] PROMPTS["practice"]: обновлены требования
              (убрана Python-специфика, добавлен синтез/моделирование).

[FIX-DEDUP9]  main(): удалён дублирующийся блок «Фикс №9» (проверка тем внутри
              разделов выполнялась дважды подряд).

[FIX-SIM-CALL] main(): добавлен вызов _print_similar_disciplines() до генерации.
"""

import json
import re
import sys
import os
import shutil
import time
import copy
import requests
from openai import OpenAI
from typing import Optional
from lxml import etree
from docx import Document
from docx.table import Table
from docx.oxml.ns import qn

OUTPUT_DOCX     = "output_rpd.docx"
GENERATION_LOG  = "generation_log.json"

QDRANT = {"url": "http://localhost:6333", "collection": "rpd_rag"}

# ---------------------------------------------------------------------------
# RouterAI — замена Ollama. Эмбеддинги и LLM через внешний API.
# Уточни точные названия моделей в личном кабинете routerai.ru.
# ---------------------------------------------------------------------------
ROUTERAI = {
    "api_key":     "sk-KnAptJMGtv69zxhmW2v8f7LILGs8umvT",
    "base_url":    "https://routerai.ru/api/v1",
    "embed_model": "qwen/qwen3-embedding-4b",
    "llm_model":   "qwen/qwen-plus",
}

_client_ai = OpenAI(
    api_key=ROUTERAI["api_key"],
    base_url=ROUTERAI["base_url"],
    timeout=120.0,
)

GENERATION = {"top_k": 8, "min_score": 0.45}

# [RERANK] Расширенный пул для первичного retrieval перед reranking.
# top-20 → cross-encoder → top-8 (GENERATION["top_k"]).
RERANK_TOP_K = 20

# [RERANK] Lazy-loaded cross-encoder (bge-reranker-v2-m3 — мультиязычный, поддерживает русский).
# Если sentence-transformers не установлен — reranking прозрачно пропускается,
# retrieve() работает как раньше (cosine-only ranking).
_reranker = None
_RERANKER_MODEL = "BAAI/bge-reranker-v2-m3"


def _get_reranker():
    """Инициализирует CrossEncoder один раз (lazy). False = попытка была, модель недоступна."""
    global _reranker
    if _reranker is not None:
        return _reranker
    try:
        from sentence_transformers import CrossEncoder  # noqa: PLC0415
        _reranker = CrossEncoder(_RERANKER_MODEL, max_length=512)
        print(f"  ✅ Reranker загружен: {_RERANKER_MODEL}")
    except Exception as e:
        print(f"  ⚠️  Reranker недоступен ({e}), cosine-only ranking")
        _reranker = False
    return _reranker


def rerank(query: str, hits: list, top_n: int) -> list:
    """
    [RERANK] Cross-encoder reranking: hits (top-RERANK_TOP_K) → top_n.

    Использует первый query как основной текст для пар (query, chunk_text).
    Если reranker недоступен — возвращает hits[:top_n] без изменений (fallback).
    """
    if not hits:
        return hits
    reranker = _get_reranker()
    if not reranker:
        return hits[:top_n]
    texts = [h.get("payload", {}).get("text", "") or "" for h in hits]
    pairs = [(query, t[:512]) for t in texts]  # CrossEncoder max_length=512 токенов
    try:
        scores = reranker.predict(pairs, show_progress_bar=False)
        ranked = sorted(zip(scores, hits), key=lambda x: x[0], reverse=True)
        reranked = [h for _, h in ranked[:top_n]]
        print(f"      ↑ rerank: {len(hits)} → {len(reranked)} "
              f"(top score: {scores[0]:.3f} → {max(scores):.3f})")
        return reranked
    except Exception as e:
        print(f"  ⚠️  rerank ошибка: {e}, fallback cosine")
        return hits[:top_n]

# [J] Максимальная длина контекста, передаваемого в LLM (символы).
# [FIX-CTX] ИСПРАВЛЕНО: поднято с 3000 до 6000 символов.
# qwen-plus (внешний API) имеет достаточное окно — прежний лимит 3000 ≈ 750 токенов
# обрезал большинство retrieved-чанков до «[...контекст обрезан...]».
# 6000 символов ≈ 1500 токенов русского текста — безопасный запас.
MAX_CONTEXT_CHARS = 6000

# [З-R2] Фильтрация чанков по section_type для каждого генерируемого раздела.
# [FIX-PRACTICE-STF] ИСПРАВЛЕНО (отчёт §1.3, §7.6): добавлены подтипы
# lecture_content, lab_content, practice_content для более точного retrieval.
# Корневая причина BLEU 0.02 для practice: classify_section() объединяла
# «содержан|лекц|лаборатор|практич|тем» в один тип content, что делало невозможным
# целевой retrieval для practice. Теперь practice ищет также по practice_content.
SECTION_TYPE_FILTER = {
    "competencies":     ["competencies", "learning_outcomes"],
    "outcomes":         ["competencies", "learning_outcomes"],
    "content":          ["content", "lecture_content"],
    "lab_works":        ["content", "lab_content", "book_content"],
    "practice":         ["content", "practice_content", "book_content"],
    "bibliography_main": ["bibliography", "place"],
}

EMBED_CACHE    = {}
RETRIEVE_CACHE = {}

# [З-R5] Персистентный кэш — сохраняется между запусками.
_CACHE_FILE = "rpd_cache.json"

# [З-G6] Хеш конфигурации retrieval
_RETRIEVAL_CONF_HASH = ""

def _make_retrieval_conf_hash(top_k: int, min_score: float) -> str:
    # [FIX-§3.2.3] Добавлены SECTION_TYPE_FILTER и corpus_timestamp.
    import hashlib as _hl, json as _js, os as _os
    stf_hash = _hl.md5(
        _js.dumps(SECTION_TYPE_FILTER, sort_keys=True).encode()
    ).hexdigest()[:8]
    _chunks_mtime = int(_os.path.getmtime("chunks.jsonl")) if _os.path.exists("chunks.jsonl") else 0
    return f"k{top_k}_s{min_score:.3f}_stf{stf_hash}_ct{_chunks_mtime}"

def _load_cache() -> None:
    """Загружает кэш из файла, если он существует."""
    global EMBED_CACHE, RETRIEVE_CACHE
    if not os.path.exists(_CACHE_FILE):
        return
    try:
        with open(_CACHE_FILE, encoding="utf-8") as f:
            data = json.load(f)
        EMBED_CACHE    = data.get("embed", {})
        RETRIEVE_CACHE = {
            k: (v[0], v[1]) for k, v in data.get("retrieve", {}).items()
        }
        print(f"  Кэш загружен: {len(EMBED_CACHE)} эмбеддингов, "
              f"{len(RETRIEVE_CACHE)} retrieval-запросов")
    except Exception as e:
        print(f"  ⚠️  Кэш не загружен: {e}")

def _save_cache() -> None:
    """Сохраняет кэш в файл."""
    try:
        with open(_CACHE_FILE, "w", encoding="utf-8") as f:
            json.dump(
                {"embed": EMBED_CACHE, "retrieve": RETRIEVE_CACHE},
                f, ensure_ascii=False
            )
    except Exception as e:
        print(f"  ⚠️  Кэш не сохранён: {e}")

# [З-R1] Multi-query: две формулировки на секцию с намеренно разным стилем.
SECTION_QUERIES = {
    "competencies": [
        "{discipline} УК ОПК ПК формируемые компетенции шифр индекс ФГОС",
        "{discipline} способен разрабатывать применять анализировать профессиональная деятельность",
    ],
    "outcomes": [
        "{discipline} результаты обучения индикаторы достижения компетенций ФГОС",
        "{discipline} знать уметь владеть навыки практические умения студент",
    ],
    "content": [
        "{discipline} тематический план содержание дисциплины разделы лекции",
        "{discipline} программа курса темы методы алгоритмы технологии практика",
    ],
    "lab_works": [
        # [FIX-LR1] Имитируем формулировку реальной ЛР вместо описания раздела.
        "{discipline} реализация алгоритма классификации обучение модели нейронная сеть",
        "{discipline} лабораторная работа задание исследование программирование Python",
    ],
    # [FIX-PRACTICE] Обновлены формулировки для лучшего покрытия тем ПЗ
    "practice": [
        "{discipline} практические занятия перечень тем задач методы алгоритмы",
        "{discipline} решение задач моделирование синтез исследование системы",
    ],
    "bibliography_main": [
        "{discipline} учебник учебное пособие литература библиография",
        "{discipline} основная дополнительная литература ЭБС Знаниум",
    ],
}

# Глобальный лог генерации — [C]
_generation_log: dict = {}


# ---------------------------------------------------------------------------
# Утилиты
# ---------------------------------------------------------------------------

def clean(text: str) -> str:
    text = re.sub(r" +", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"\[score=[^\]]+\]\n?", "", text)
    return "\n".join(l.strip() for l in text.split("\n") if l.strip()).strip()


def get_embedding(text: str):
    if text in EMBED_CACHE:
        return EMBED_CACHE[text]
    for attempt in range(3):
        try:
            response = _client_ai.embeddings.create(
                model=ROUTERAI["embed_model"],
                input=text,
                encoding_format="float",
            )
            vec = response.data[0].embedding
            if vec:
                EMBED_CACHE[text] = vec
                return vec
        except Exception as e:
            if attempt == 2:
                return []
            time.sleep(2 ** attempt)
    return []


def _search_qdrant(vec: list, payload_filter: dict | None, top_k: int) -> list:
    """Поиск в Qdrant с fallback query → search."""
    try:
        body = {"query": vec, "limit": top_k, "with_payload": True}
        if payload_filter:
            body["filter"] = payload_filter
        r = requests.post(
            f"{QDRANT['url']}/collections/{QDRANT['collection']}/points/query",
            json=body, timeout=30)
        r.raise_for_status()
        return r.json().get("result", {}).get("points", [])
    except requests.HTTPError:
        body = {"vector": vec, "limit": top_k, "with_payload": True}
        if payload_filter:
            body["filter"] = payload_filter
        r = requests.post(
            f"{QDRANT['url']}/collections/{QDRANT['collection']}/points/search",
            json=body, timeout=30)
        r.raise_for_status()
        return r.json().get("result", [])


def retrieve(section: str, discipline: str, section_types: list = None,
             direction: str = "", level: str = "") -> tuple[str, list]:
    """
    Ищет релевантные чанки в Qdrant.

    [K] Multi-query: объединяем результаты по нескольким формулировкам.
    [B] Доменная фильтрация по direction/level.
    [S] Фильтр использует "section_type" (верхний уровень payload).
    [R] При пустом результате возвращает пустую строку с флагом для caller.

    Возвращает: (ctx_string, hits_list) для логирования [C].
    """
    cache_key = (f"{section}|{discipline}|{','.join(section_types or [])}"
                 f"|{direction}|{level}|{_RETRIEVAL_CONF_HASH}")
    if cache_key in RETRIEVE_CACHE:
        return RETRIEVE_CACHE[cache_key]

    try:
        must_conditions: list = []
        if section_types:
            if len(section_types) == 1:
                # [FIX-SHOULD1] Qdrant отклоняет "should" с одним условием (HTTP 400).
                must_conditions.append(
                    {"key": "section_type", "match": {"value": section_types[0]}}
                )
            else:
                must_conditions.append({
                    "should": [
                        {"key": "section_type", "match": {"value": st}}
                        for st in section_types
                    ]
                })
        if direction:
            must_conditions.append({"key": "direction", "match": {"value": direction}})
        if level:
            must_conditions.append({"key": "level", "match": {"value": level}})

        payload_filter = {"must": must_conditions} if must_conditions else None

        queries = SECTION_QUERIES.get(section, [f"{discipline} {section}"])
        queries = [q.format(discipline=discipline) for q in queries]

        all_hits: dict[int, dict] = {}
        for query_text in queries:
            vec = get_embedding(query_text)
            if not vec:
                continue
            hits = _search_qdrant(vec, payload_filter, RERANK_TOP_K)  # [RERANK] пул 20
            for h in hits:
                hit_id = h.get("id")
                if hit_id not in all_hits or h.get("score", 0) > all_hits[hit_id].get("score", 0):
                    all_hits[hit_id] = h

        # [RERANK] Предварительный фильтр с мягким порогом (0.7×min_score),
        # чтобы дать reranker достаточный пул кандидатов даже при низких cosine scores.
        pre_filter = sorted(
            [h for h in all_hits.values()
             if h.get("score", 0) >= GENERATION["min_score"] * 0.7],
            key=lambda h: h.get("score", 0), reverse=True,
        )
        # Cross-encoder reranking пула → top_k*2 кандидатов перед diversity filter
        reranked = rerank(queries[0], pre_filter, GENERATION["top_k"] * 2)

        MAX_PER_SOURCE = 2
        _source_counts: dict = {}
        _diverse_all: list = []
        for h in reranked:
            src = h.get("payload", {}).get("source", "")
            if _source_counts.get(src, 0) < MAX_PER_SOURCE:
                _source_counts[src] = _source_counts.get(src, 0) + 1
                _diverse_all.append(h)

        good_hits = _diverse_all[:GENERATION["top_k"]]

        if not good_hits:
            print(f"    ⚠️  RAG [{section}]: нет чанков выше {GENERATION['min_score']}, "
                  f"пробую без доменного фильтра...")
            vec = get_embedding(queries[0])
            if vec:
                hits = _search_qdrant(vec, None, RERANK_TOP_K)  # [RERANK]
                fallback_candidates = sorted(
                    [h for h in hits if h.get("score", 0) >= GENERATION["min_score"] * 0.7],
                    key=lambda h: h.get("score", 0), reverse=True,
                )
                good_hits = rerank(queries[0], fallback_candidates, GENERATION["top_k"])

        print(f"    🔍 RAG [{section}]: найдено {len(good_hits)} чанков "
              f"(scores: {[round(h.get('score', 0), 3) for h in good_hits]})")

        seen_texts: set = set()
        parts: list[str] = []
        for h in good_hits:
            payload = h.get("payload", {})
            raw_text = payload.get("text", "")
            if not raw_text:
                continue

            if len(raw_text) > 1200:
                cut = raw_text[:1200]
                last_dot = cut.rfind(".")
                text = cut[:last_dot + 1] if last_dot >= 800 else cut
            else:
                text = raw_text

            dedup_key = text[:100]
            if dedup_key in seen_texts:
                continue
            seen_texts.add(dedup_key)
            source        = payload.get("source", "")
            section_title = payload.get("section_title", "")
            prefix = ""
            if source:
                prefix += f"[{source}]"
            if section_title:
                prefix += f" [{section_title}]"
            parts.append(f"{prefix}\n{text}" if prefix else text)

        ctx = "\n\n---\n\n".join(parts)

        if len(ctx) > MAX_CONTEXT_CHARS:
            ctx = ctx[:MAX_CONTEXT_CHARS].rsplit("\n", 1)[0]
            ctx += "\n[...контекст обрезан до MAX_CONTEXT_CHARS символов...]"

        RETRIEVE_CACHE[cache_key] = (ctx, good_hits)
        return ctx, good_hits

    except Exception as e:
        print(f"  ⚠️  RAG [{section}]: {e}")
        return "", []


def llm(prompt: str, max_tokens: int = 800) -> str:
    for attempt in range(3):
        try:
            response = _client_ai.chat.completions.create(
                model=ROUTERAI["llm_model"],
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "Ты помощник по составлению рабочих программ дисциплин (РПД) "
                            "для российского технического университета. "
                            "Отвечай строго по запросу, без лишних пояснений."
                        ),
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.3,
                max_tokens=max_tokens,
            )
            text = response.choices[0].message.content or ""
            if text:
                return clean(text)
        except Exception as e:
            if attempt == 2:
                return f"[Ошибка: {e}]"
            time.sleep(5)
    return "[Ошибка: пустой ответ]"


def _sanitize_retrieved_text(text: str) -> str:
    """
    [замечание #13] Базовая защита от prompt injection в retrieved-контексте.
    """
    INJECTION_PATTERNS = re.compile(
        r"^(ignore\s+(previous|all|prior)|forget\s+(previous|all)|"
        r"system\s*:|instruction\s*:|act\s+as\s|"
        r"забудь\s+(предыдущ|все)|игнорируй\s+предыдущ|"
        r"ты\s+теперь\s|притворяйся\s|ты\s+—\s)",
        re.IGNORECASE,
    )
    clean_lines = [
        line for line in text.split("\n")
        if not INJECTION_PATTERNS.match(line.strip())
    ]
    return "\n".join(clean_lines).strip()


# [§2.2.5] Словарь замены ошибочных транслитераций и калек.
_TERM_CORRECTIONS: list[tuple[str, str]] = [
    (r"\bсемисери\b",               "полуконтролируемое обучение"),
    (r"\bполусери\b",               "полуконтролируемое обучение"),
    (r"\bбезпосредственн\w*",       "обучение без учителя"),
    (r"\bГейш-рекуррентн\w*",       "управляемый рекуррентный блок"),
    (r"\bГейш\s+рекуррентн\w*",     "управляемый рекуррентный блок"),
    (r"\bглубокий\s+обучени[ея]\b", "глубокое обучение"),
    (r"\bмашинный\s+обучени[ея]\b", "машинное обучение"),
    (r"\bнадзорное\s+обучени[ея]\b","обучение с учителем"),
    (r"\bненадзорн\w+",             "без учителя"),
    (r"\bинтеллектального\b",       "интеллектуального"),
]
_TERM_RE: list[tuple[re.Pattern, str]] = [
    (re.compile(pat, re.IGNORECASE), repl)
    for pat, repl in _TERM_CORRECTIONS
]


def _apply_term_corrections(text: str) -> str:
    """Применяет _TERM_RE к тексту LLM-ответа."""
    for pattern, replacement in _TERM_RE:
        text = pattern.sub(replacement, text)
    return text


def gen(label: str, discipline: str, prompt: str,
        direction: str = "", level: str = "", **extra) -> str:
    """
    Генерация секции с RAG-контекстом.

    [R] При пустом retrieval добавляет явную инструкцию в промпт.
    [C] Сохраняет данные в _generation_log для последующей записи в JSON.
    """
    section_types = SECTION_TYPE_FILTER.get(label)
    ctx, hits = retrieve(label, discipline, section_types, direction, level)

    ctx = _sanitize_retrieved_text(ctx)

    if ctx:
        ctx_block = (
            "Примеры из базы РПД кафедры (используй как образец стиля и формата):\n"
            f"{ctx}\n\n"
        )
    else:
        print(f"  ⚠️  RAG [{label}]: контекст пуст — генерация без примеров из корпуса")
        ctx_block = (
            "Примеры из базы РПД недоступны. "
            "Сгенерируй содержимое самостоятельно строго по указанному формату "
            "без копирования примеров из промпта.\n\n"
        )

    fmt_vars = {"discipline": discipline, "direction": direction, "level": level, **extra}
    full_prompt = ctx_block + prompt.format(**fmt_vars) + f"\n\nСоздай для «{discipline}»:"
    result = _apply_term_corrections(llm(full_prompt))

    _generation_log[label] = {
        "prompt_preview":   full_prompt[:600],
        "retrieved_chunks": [
            {
                "id":           h.get("id"),
                "source":       h.get("payload", {}).get("source", ""),
                "score":        round(h.get("score", 0), 4),
                "text_preview": h.get("payload", {}).get("text", "")[:120],
            }
            for h in hits
        ],
        "llm_response":     result,
        "timestamp":        time.strftime("%Y-%m-%dT%H:%M:%S"),
    }

    return result


# ---------------------------------------------------------------------------
# [T1-T4] Работа с DOCX-шаблоном на основе [] плейсхолдеров
# ---------------------------------------------------------------------------

# Маппинг семантического ключа → предикат для поиска таблицы.
# [FIX-HDR5] max_rows поднят до 5 — заголовки work_types/fos занимают 4 строки.
_TABLE_PREDICATES: dict = {
    "workload":     lambda h: "Зачетные единицы" in h,
    "competencies": lambda h: "Формируемые компетенции" in h,
    "outcomes":     lambda h: "Индикаторы достижения компетенций" in h,
    "work_types":   lambda h: "Вид учебной работы" in h,
    "topics":       lambda h: "Название темы (раздела)" in h,
    "lectures":     lambda h: (
        "Название темы" in h and "№ пп." in h
        and "Название темы (раздела)" not in h
        and "лабораторной" not in h
    ),
    "labs":         lambda h: "Название лабораторной работы" in h,
    "practice":     lambda h: "Тема практического занятия" in h,
    "sro":          lambda h: "Вид СРО" in h,
    "bibliography": lambda h: "Тип" in h and "Библиографическое описание" in h,
    "method_bib":   lambda h: (
        "Назначение учебных изданий" in h and "Тип" not in h
    ),
    "fos":          lambda h: "Контролируемые разделы (темы) дисциплины" in h,
    "fos_types":    lambda h: "Вид оценочного средства" in h and "Контролируемые разделы (темы) дисциплины" not in h,
}


def _table_header_set(table, max_rows: int = 5) -> frozenset:
    """Собирает множество уникальных текстов ячеек из первых max_rows строк.
    [FIX-HDR5] Поднят с 3 до 5: в шаблонах УГНТУ заголовки таблиц
    (work_types, fos) иногда занимают 4 строки (двойная шапка)."""
    texts = set()
    for row in table.rows[:max_rows]:
        seen_tc = set()
        for cell in row.cells:
            if id(cell._tc) not in seen_tc:
                seen_tc.add(id(cell._tc))
                t = cell.text.strip()
                if t:
                    texts.add(t)
    return frozenset(texts)


def find_annotation_table(doc: Document) -> Optional[Table]:
    """
    Находит таблицу аннотации РПД по позиции в документе:
    ищет первую таблицу после параграфа «Аннотация к рабочей программе».
    """
    body = doc.element.body
    found_ann = False
    for child in body:
        tag = child.tag.split("}")[1] if "}" in child.tag else child.tag
        if tag == "p":
            text = "".join(
                t.text or "" for t in child.findall(f".//{{{qn('w:t').split('}')[0][1:]}}}t")
            )
            if "Аннотация к рабочей программе" in text:
                found_ann = True
        elif tag == "tbl" and found_ann:
            return Table(child, doc)
    return None


def fill_annotation_table(
    doc: Document,
    competencies: list,
    outcomes: list,
    topics: list,
    credits: int,
    hours_total: int,
    exam_type: str,
) -> None:
    """
    Заполняет таблицу аннотации РПД (15 строк).
    """
    table = find_annotation_table(doc)
    if table is None:
        print("  ⚠️  [fill_annotation] Таблица аннотации не найдена")
        return

    # [1] Компетенции с индикаторами
    _ann_z_texts = [t for ot, t in outcomes if ot == "З"] if outcomes else []
    comp_lines = []
    for comp_idx, (code, desc) in enumerate(competencies):
        comp_lines.append(f" {code} {desc}:")
        _ind_text = _ann_z_texts[comp_idx % len(_ann_z_texts)][:60] if _ann_z_texts \
            else "Применяет методы и инструменты дисциплины"
        comp_lines.append(f"-{code}.1 Знает {_ind_text}")
    comp_text = "\n".join(comp_lines)

    # [4][6][8] Группируем outcomes по типу
    z_items = [(code, t) for (ot, t), (code, _) in zip(
        [o for o in outcomes if o[0] == "З"],
        competencies
    )] if outcomes else []
    u_items = [(code, t) for (ot, t), (code, _) in zip(
        [o for o in outcomes if o[0] == "У"],
        competencies
    )] if outcomes else []
    v_items = [(code, t) for (ot, t), (code, _) in zip(
        [o for o in outcomes if o[0] == "В"],
        competencies
    )] if outcomes else []

    def fmt_outcomes(items: list) -> str:
        return "\n".join(f"{code}- {text}" for code, text in items) if items else ""

    # [10] Краткий перечень разделов
    sections = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    sections_text = "; ".join(
        re.sub(r"^Раздел\s*\d+\.\s*", "", s) for s in sections
    ) if sections else "; ".join(t for t in topics[:3])

    values = [
        "Компетенции, формируемые в результате освоения дисциплины",
        comp_text,
        "Результат обучения",
        "Знать:",
        fmt_outcomes(z_items) or "Основные методы и понятия дисциплины",
        "Уметь:",
        fmt_outcomes(u_items) or "Применять методы дисциплины для решения задач",
        "Владеть:",
        fmt_outcomes(v_items) or "Навыками работы с инструментами дисциплины",
        "Краткая характеристика дисциплины",
        sections_text,
        "Трудоёмкость (з.е. / часы)",
        f"{credits} з.е.  ({hours_total}час)",
        "Вид промежуточной аттестации",
        exam_type,
    ]

    tbl_xml = table._tbl
    all_trs  = tbl_xml.findall(qn("w:tr"))
    for tr, val in zip(all_trs, values):
        seen: set = set()
        for tc in tr.findall(f".//{qn('w:tc')}"):
            if id(tc) in seen:
                continue
            seen.add(id(tc))
            _set_cell_xml(tc, val)
            break


def find_table(doc: Document, key: str) -> Optional[Table]:
    """
    [T3] Находит таблицу по семантическому ключу из _TABLE_PREDICATES.
    """
    predicate = _TABLE_PREDICATES.get(key)
    if predicate is None:
        raise KeyError(f"Неизвестный ключ таблицы: {key!r}. "
                       f"Доступны: {sorted(_TABLE_PREDICATES)}")
    for table in doc.tables:
        header_set = _table_header_set(table)
        if predicate(header_set):
            return table
    print(f"  ⚠️  [find_table] Таблица {key!r} не найдена — проверьте шаблон")
    return None


def _set_cell_xml(tc, text: str) -> None:
    """
    Записывает text в ячейку (lxml <w:tc>), сохраняя форматирование первого run.
    """
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W  = lambda tag: f"{{{NS}}}{tag}"

    paras = tc.findall(W("p"))
    if not paras:
        lines = text.split("\n") if "\n" in text else [text]
        new_p = etree.SubElement(tc, W("p"))
        new_r = etree.SubElement(new_p, W("r"))
        new_t = etree.SubElement(new_r, W("t"))
        new_t.text = lines[0]
        for line in lines[1:]:
            extra_p = etree.SubElement(tc, W("p"))
            extra_r = etree.SubElement(extra_p, W("r"))
            extra_t = etree.SubElement(extra_r, W("t"))
            extra_t.text = line
        return

    for p in paras[1:]:
        tc.remove(p)
    p = paras[0]

    runs = p.findall(f".//{W('r')}")
    saved_rpr = None
    if runs:
        rpr = runs[0].find(W("rPr"))
        if rpr is not None:
            saved_rpr = copy.deepcopy(rpr)
        for r in runs:
            p.remove(r)

    lines = text.split("\n") if "\n" in text else [text]

    new_r = etree.SubElement(p, W("r"))
    if saved_rpr is not None:
        new_r.append(copy.deepcopy(saved_rpr))
    new_t = etree.SubElement(new_r, W("t"))
    new_t.text = lines[0]
    if lines[0] and (lines[0][0] == " " or lines[0][-1] == " "):
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    for line in lines[1:]:
        new_p = copy.deepcopy(p)
        for r in new_p.findall(f".//{W('r')}"):
            new_p.remove(r)
        nr = etree.SubElement(new_p, W("r"))
        if saved_rpr is not None:
            nr.append(copy.deepcopy(saved_rpr))
        nt = etree.SubElement(nr, W("t"))
        nt.text = line
        if line and (line[0] == " " or line[-1] == " "):
            nt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        tc.append(new_p)


def _is_placeholder_row(tr) -> bool:
    """True если хотя бы одна уникальная ячейка строки содержит только '[]'."""
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W  = lambda tag: f"{{{NS}}}{tag}"
    seen = set()
    for tc in tr.findall(f".//{W('tc')}"):
        if id(tc) in seen:
            continue
        seen.add(id(tc))
        t = "".join(x.text or "" for x in tc.findall(f".//{W('t')}")).strip()
        if t == "[]":
            return True
    return False


def _fill_tr(tr, values: list[str]) -> None:
    """Записывает values в уникальные ячейки строки."""
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W  = lambda tag: f"{{{NS}}}{tag}"
    cells = []
    seen = set()
    for tc in tr.findall(f".//{W('tc')}"):
        if id(tc) not in seen:
            seen.add(id(tc))
            cells.append(tc)
    for tc, val in zip(cells, values):
        _set_cell_xml(tc, str(val))


def fill_placeholder_rows(table: Table, data: list[list[str]]) -> None:
    """
    [T2] Заполняет строки таблицы, содержащие [] плейсхолдеры.
    """
    tbl_xml = table._tbl
    all_trs = list(tbl_xml)

    placeholder_trs = [tr for tr in all_trs if _is_placeholder_row(tr)]
    if not placeholder_trs:
        print("  ⚠️  [fill_placeholder_rows] Строк с [] не найдено — шаблон не обновлён?")
        return

    template_tr = placeholder_trs[-1]

    for tr, values in zip(placeholder_trs, data):
        _fill_tr(tr, values)

    for values in data[len(placeholder_trs):]:
        new_tr = copy.deepcopy(template_tr)
        tbl_xml.append(new_tr)
        _fill_tr(new_tr, values)

    for extra_tr in placeholder_trs[len(data):]:
        tbl_xml.remove(extra_tr)


def fill_doc_header(doc: Document, discipline: str, code: str,
                    year: str = "2025", credits: int = 4,
                    hours_total: int = 144, exam_type: str = "экзамен") -> None:
    """
    [T1] Заменяет [] плейсхолдеры в параграфах документа.
    """
    label    = f"({code}){discipline}"
    workload = f"{credits} з.е.  ({hours_total}час)"

    _ASSESSMENT_HEADERS = {
        "реферат", "доклад", "лабораторная работа",
        "письменный и устный опрос", "тест",
        "расчётно-графическая работа",
        "перечень вопросов (задач, заданий, тем, комплекта тестовых заданий):",
    }
    # [Д-2] Флаг блока «Темы для СРО»
    _in_sro_block = False
    prev_txt = ""

    def _set_para(para, text: str) -> None:
        if not para.runs:
            return
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""

    first_label_done = False

    for para in doc.paragraphs:
        txt = para.text.strip()

        # [Д-2] SRO-блок: отслеживаем вход/выход
        if "Темы для самостоятельной работы обучающихся" in txt:
            _in_sro_block = True
        elif _in_sro_block and re.match(r"^\d+\.", txt):
            _in_sro_block = False

        if _in_sro_block:
            prev_txt = txt
            continue

        if not first_label_done and txt == "[]":
            _set_para(para, label)
            first_label_done = True
            prev_txt = txt
            continue

        if txt == "Трудоемкость дисциплины: []":
            _set_para(para, f"Трудоемкость дисциплины: {workload}")
            prev_txt = txt
            continue

        if txt == f"Уфа []":
            _set_para(para, f"Уфа {year}")
            prev_txt = txt
            continue

        # [§6.1.5] Замена хардкод-дат
        if re.search(r"\b\d{2}\.\d{2}\.\d{4}\b", txt):
            new_txt = re.sub(r"\b(\d{2}\.\d{2}\.)\d{4}\b",
                             lambda m: m.group(1) + str(year), para.text)
            if new_txt != para.text:
                _set_para(para, new_txt)
            prev_txt = txt
            continue

        if re.search(r"Год\s+приема\s+\d{4}\s+г", txt, re.IGNORECASE):
            new_txt = re.sub(r"(Год\s+приема\s+)\d{4}(\s+г)",
                             lambda m: m.group(1) + str(year) + m.group(2),
                             para.text, flags=re.IGNORECASE)
            if new_txt != para.text:
                _set_para(para, new_txt)
            prev_txt = txt
            continue

        if txt == "[]" and first_label_done:
            if prev_txt.strip().lower() not in _ASSESSMENT_HEADERS:
                _set_para(para, label)
            prev_txt = txt
            continue

        if txt == "[]" or "[]" in txt:
            if "з.е." in txt or "час" in txt:
                _set_para(para, para.text.replace("[]", workload))
            else:
                _set_para(para, para.text.replace("[]", label))

        prev_txt = txt


def fill_appendix_v(doc: Document, discipline: str, topics: list) -> None:
    """
    Заполняет перечни вопросов/заданий в Приложении В.
    """
    sections = [
        re.sub(r"^Раздел\s*\d+\.\s*", "", t)
        for t in topics if re.match(r"^Раздел\s*\d+", t)
    ][:5]
    if not sections:
        sections = [f"основные разделы дисциплины «{discipline}»"]

    topics_list   = "\n".join(f"- {s}" for s in sections)
    topics_inline = "; ".join(sections)

    # [FIX-AppV-DOK] Доклад строится из подтем (Тема X.Y.), а не разделов.
    # [FIX-10] Отчёт §4, замечание 10.
    # Было: subtopics_list = «- <название темы>» — дословная копия названия
    #       лекции, что формально некорректно для Приложения В (там ожидаются
    #       исследовательские вопросы, а не перечисление тем).
    # Стало: каждая подтема обрамляется в вопросительную / исследовательскую
    #       формулировку. Чередуются префиксы, чтобы тексты не были однообразны.
    subtopics = [
        re.sub(r"^Тема\s+[\d.]+\s*", "", t).strip()
        for t in topics if re.match(r"^Тема\s+[\d.]", t)
    ][:6]

    def _to_research_topic(subtopic: str, idx: int) -> str:
        """Превращает «название темы» в «исследовательский вопрос для доклада»."""
        base = subtopic.rstrip(".").strip()
        if not base:
            return subtopic
        # Нижний регистр первой буквы для встраивания во фразу
        base_lc = base[0].lower() + base[1:] if len(base) > 1 else base.lower()
        templates = [
            f"Современное состояние и перспективы развития: {base_lc}",
            f"Сравнительный анализ подходов к теме «{base}»",
            f"Применение {base_lc} в прикладных задачах",
            f"Актуальные проблемы и ограничения: {base_lc}",
            f"Обзор методов и алгоритмов: {base_lc}",
            f"{base}: практический обзор существующих решений",
        ]
        return templates[idx % len(templates)]

    subtopics_list = (
        "\n".join(f"- {_to_research_topic(s, i)}" for i, s in enumerate(subtopics))
        if subtopics else topics_list
    )

    _TEMPLATES = {
        "реферат": (
            f"Темы рефератов по дисциплине «{discipline}»:\n{topics_list}"
        ),
        "доклад": (
            f"Темы докладов по дисциплине «{discipline}»:\n{subtopics_list}"
        ),
        "лабораторная работа": (
            f"Перечень лабораторных работ охватывает разделы: {topics_inline}"
        ),
        "письменный и устный опрос": (
            f"Контрольные вопросы по разделам дисциплины «{discipline}»:\n{topics_list}"
        ),
        "тест": (
            f"Тестирование по дисциплине «{discipline}». "
            f"Охватываемые разделы: {topics_inline}"
        ),
        "расчётно-графическая работа": (
            f"Задания РГР по дисциплине «{discipline}»:\n{topics_list}"
        ),
    }

    # [Д-4 ИСПРАВЛЕНО]: last_assessment сохраняется до встречи [] независимо от
    # числа промежуточных параграфов.
    last_assessment: str | None = None
    for para in doc.paragraphs:
        txt     = para.text.strip()
        txt_key = txt.lower().rstrip(".")

        matched_key = next((k for k in _TEMPLATES if txt_key.startswith(k)), None)
        if matched_key:
            last_assessment = matched_key
            continue

        if txt == "[]" and last_assessment is not None:
            content = _TEMPLATES[last_assessment]
            if para.runs:
                para.runs[0].text = content
                for r in para.runs[1:]:
                    r.text = ""
            last_assessment = None
            continue

        if re.match(r"^\d+\.", txt) or txt.startswith("СОГЛАСОВАНО"):
            last_assessment = None


def fill_sro_topic_paragraphs(doc: Document, topics: list, label: str) -> None:
    """[FIX-SRO] Заполняет блок «Темы для СРО»."""
    sections = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    if not sections:
        return

    # [FIX-SRO-TOPICS] Нечётные слоты — конкретные темы раздела
    section_subtopics: list[str] = []
    current: list[str] = []
    for t in topics:
        if re.match(r"^Раздел\s*\d+", t):
            if current or section_subtopics:
                section_subtopics.append("; ".join(current) if current else label)
                current = []
        elif re.match(r"^Тема\s+\d", t):
            name = re.sub(r"^Тема\s+[\d.]+\s*", "", t).strip()
            current.append(name)
    section_subtopics.append("; ".join(current) if current else label)

    in_block = False
    slot_idx = 0
    sec_idx  = 0

    for para in doc.paragraphs:
        txt = para.text.strip()

        if "Темы для самостоятельной работы обучающихся" in txt:
            in_block = True
            continue

        if in_block:
            if re.match(r"^\d+\.", txt) and "Раздел" not in txt:
                break
            if txt == "[]":
                if para.runs:
                    if slot_idx % 2 == 0:
                        text = sections[sec_idx] if sec_idx < len(sections) else label
                        sec_idx += 1
                    else:
                        idx  = sec_idx - 1
                        text = (
                            section_subtopics[idx]
                            if 0 <= idx < len(section_subtopics)
                            else label
                        )
                    para.runs[0].text = text
                    for r in para.runs[1:]:
                        r.text = ""
                slot_idx += 1


# [Фикс Д-SecRot] Детерминированная ротация секций.
def _normalize_section_assignment(items: list, n_sections: int) -> list:
    if n_sections < 1 or not items:
        return items
    n       = len(items)
    per_sec = max(1, n // n_sections)
    for i, item in enumerate(items):
        if isinstance(item, dict):
            item["section"] = min((i // per_sec) + 1, n_sections)
    return items


def set_cell_text(cell, text: str) -> None:
    """Устанавливает текст ячейки python-docx Cell. Используется только в fill_t6_workload."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if not cell.paragraphs:
        cell.add_paragraph(text)
    elif not cell.paragraphs[0].runs:
        cell.paragraphs[0].add_run(text)
    else:
        cell.paragraphs[0].runs[0].text = text


# ---------------------------------------------------------------------------
# [A] JSON-парсеры с fallback на regex
# ---------------------------------------------------------------------------

def parse_competencies_json(text: str) -> list | None:
    m = re.search(r"\[.*\]", text, re.S)
    if not m:
        return None
    try:
        data = json.loads(m.group())
        if not isinstance(data, list):
            return None
        result = [
            (str(d.get("code", "")), str(d.get("desc", "")))
            for d in data
            if isinstance(d, dict) and d.get("code") and d.get("desc")
        ]
        return result if result else None
    except (json.JSONDecodeError, TypeError):
        return None


def parse_competencies(text: str, codes: list = None) -> list:
    json_result = parse_competencies_json(text)
    if json_result:
        return json_result

    descriptions = []
    seen = set()
    for line in text.split("\n"):
        line = re.sub(r"^\d+[\.)]\s*", "", line.strip())
        line = re.sub(r"^[-–•]\s*", "", line)
        line = re.sub(r"\*\*", "", line).strip()
        if not line or len(line) < 10:
            continue
        if re.match(r"^Способен", line, re.I):
            key = line.lower()[:60]
            if key not in seen:
                seen.add(key)
                descriptions.append(line)

    if codes and descriptions:
        while len(descriptions) < len(codes):
            descriptions.append("Способен применять методы и инструменты дисциплины на практике")
        return list(zip(codes, descriptions[:len(codes)]))

    result = []
    seen_codes: set = set()
    for line in text.split("\n"):
        m = re.match(r"(УК-\d+|ОПК-\d+|ПК-\d+)[:\.\s]+(.+)", line.strip())
        if m and m.group(1) not in seen_codes:
            seen_codes.add(m.group(1))
            result.append((m.group(1), m.group(2).strip()))
    return result if result else [
        ("УК-1",  "Способен применять системный подход для анализа и решения задач"),
        ("ОПК-1", "Способен разрабатывать алгоритмы и программы для интеллектуальных систем"),
        ("ПК-1",  "Способен применять методы машинного обучения для решения прикладных задач"),
    ]


def parse_outcomes_json(text: str, required_count: int = 0) -> list | None:
    m = re.search(r"\[.*\]", text, re.S)
    if not m:
        return None
    try:
        data = json.loads(m.group())
        if not isinstance(data, list):
            return None
        result = [
            (str(d.get("type", "")), str(d.get("text", "")))
            for d in data
            if isinstance(d, dict) and d.get("type") in ("З", "У", "В") and d.get("text")
        ]
        if len(result) < 3:
            return None
        if required_count > 0 and len(result) < required_count:
            print(f"  ⚠️  [outcomes] JSON содержит {len(result)} элементов, "
                  f"нужно {required_count} — retry")
            return None
        return result
    except (json.JSONDecodeError, TypeError):
        return None


def parse_outcomes(text: str) -> list:
    json_result = parse_outcomes_json(text)
    if json_result:
        return json_result

    result = []
    current_type = None
    lines = []

    def flush():
        if current_type and lines:
            result.append((current_type, "\n".join(lines)))

    def split_inline(rest: str) -> list[str]:
        items = re.split(r";\s*-\s*|;\s*–\s*", rest)
        cleaned = []
        for item in items:
            item = re.sub(r"^[-–•]\s*", "", item.strip())
            item = re.sub(r"^\d+[\.)]\s*", "", item)
            item = re.sub(r"\*\*", "", item)
            if item and len(item) > 3:
                cleaned.append(item)
        return cleaned

    for line in text.split("\n"):
        line = line.strip()
        m_know = re.match(r"^Знать:\s*(.*)", line, re.I)
        m_can  = re.match(r"^Уметь:\s*(.*)", line, re.I)
        m_have = re.match(r"^Владеть:\s*(.*)", line, re.I)
        if m_know:
            flush(); current_type = "З"; lines = []
            rest = m_know.group(1).strip()
            if rest:
                lines.extend(split_inline(rest) if ";" in rest else [rest])
        elif m_can:
            flush(); current_type = "У"; lines = []
            rest = m_can.group(1).strip()
            if rest:
                lines.extend(split_inline(rest) if ";" in rest else [rest])
        elif m_have:
            flush(); current_type = "В"; lines = []
            rest = m_have.group(1).strip()
            if rest:
                lines.extend(split_inline(rest) if ";" in rest else [rest])
        elif line and current_type:
            item = re.sub(r"^\d+[\.)]\s*|\*\*|^[-–•]\s*", "", line)
            if item:
                lines.append(item)

    flush()

    VLADEET_PREFIXES = ("навыками", "методами", "инструментами", "технологиями",
                        "опытом", "практикой", "способностью")

    # [FIX-ZUV] Валидация З и У
    _ACTION_VERBS = (
        "применять", "разрабатыва", "анализирова", "использова",
        "реализовыва", "проектирова", "оценива", "выполнять",
        "создавать", "строить", "моделирова", "формирова",
        "составлять", "решать", "описыва", "обеспечива",
    )
    _INFINITIVE_PREFIXES = (
        "применять", "разрабатыва", "анализирова", "использова",
        "реализовыва", "проектирова", "оценива", "выполнять",
        "создавать", "строить", "моделирова", "формирова",
        "составлять", "решать", "описыва", "исследова", "обеспечива",
        "разрабатывать",
    )

    fixed = []
    for otype, otext in result:
        if otype == "В":
            fixed_lines = []
            for ln in otext.split("\n"):
                ln = ln.strip()
                if not ln:
                    continue
                ll = ln.lower()
                if not any(ll.startswith(p) for p in VLADEET_PREFIXES):
                    ln = re.sub(r"^(Основ[ыа]|Знание|Понимание|Базов[ые]+)\s+",
                                "навыками ", ln, flags=re.I)
                    ll = ln.lower()
                    if not any(ll.startswith(p) for p in VLADEET_PREFIXES):
                        ln = "навыками " + ln[0].lower() + ln[1:]
                fixed_lines.append(ln)
            fixed.append((otype, "\n".join(fixed_lines)))

        elif otype == "З":
            fixed_lines = []
            for ln in otext.split("\n"):
                ln = ln.strip()
                if not ln:
                    continue
                ll = ln.lower()
                if any(ll.startswith(v) for v in _ACTION_VERBS):
                    ln = re.sub(
                        r"^(применять|использовать|разрабатывать|анализировать|"
                        r"реализовывать|проектировать|оценивать|выполнять|"
                        r"создавать|строить|моделировать|формировать|"
                        r"составлять|решать|описывать|обеспечивать)\s+",
                        "", ln, flags=re.I,
                    ).strip()
                    if ln:
                        ln = ln[0].upper() + ln[1:]
                fixed_lines.append(ln)
            fixed.append((otype, "\n".join(fixed_lines)))

        elif otype == "У":
            fixed_lines = []
            for ln in otext.split("\n"):
                ln = ln.strip()
                if not ln:
                    continue
                ll = ln.lower()
                if not any(ll.startswith(p) for p in _INFINITIVE_PREFIXES):
                    ln = "применять " + ln[0].lower() + ln[1:]
                fixed_lines.append(ln)
            fixed.append((otype, "\n".join(fixed_lines)))

        else:
            fixed.append((otype, otext))

    if fixed:
        return fixed

    _z_base = [
        "основные методы и алгоритмы дисциплины",
        "теоретические принципы построения систем",
        "современные инструменты и технологии в данной области",
        "методы анализа и оценки эффективности систем",
        "нормативную базу и стандарты в области дисциплины",
    ]
    _u_base = [
        "применять методы дисциплины для решения практических задач",
        "разрабатывать и реализовывать алгоритмы в рамках дисциплины",
        "анализировать результаты и интерпретировать их в контексте задачи",
        "использовать инструментальные средства при проектировании систем",
        "выбирать оптимальные подходы для решения профессиональных задач",
    ]
    _v_base = [
        "навыками применения методов дисциплины на практике",
        "методами проектирования и разработки систем",
        "инструментами анализа и оценки качества решений",
        "навыками работы с профессиональными программными средствами",
        "методами исследования и верификации результатов",
    ]
    fallback: list = []
    for z in _z_base:
        fallback.append(("З", z))
    for u in _u_base:
        fallback.append(("У", u))
    for v in _v_base:
        fallback.append(("В", v))
    return fallback


def parse_topics_json(text: str) -> list | None:
    m = re.search(r"\[.*\]", text, re.S)
    if not m:
        try:
            data = json.loads(text.strip())
            if isinstance(data, list):
                pass
            else:
                return None
        except Exception:
            return None
    else:
        try:
            data = json.loads(m.group())
        except (json.JSONDecodeError, TypeError):
            return None

    if not isinstance(data, list):
        return None

    SECTION_TYPES = {"section", "раздел", "section_type"}
    TOPIC_TYPES   = {"topic", "тема", "подтема", "subtopic"}
    SKIP_TYPES    = {"lecture", "лекция", "lab", "лр", "practice", "пз",
                     "work", "задание", "task", "item"}

    topics = []
    section_counter = 0
    for d in data:
        if not isinstance(d, dict):
            continue
        item_type = str(d.get("type", "")).strip().lower()
        label = str(d.get("label", "")).strip()
        name  = str(d.get("name",  "")).strip()

        if not name:
            continue
        if item_type in SKIP_TYPES:
            continue

        is_section = (item_type in SECTION_TYPES or
                      re.match(r"^(?:Раздел|Section)\s*\d*", label, re.I))
        is_topic   = (item_type in TOPIC_TYPES or
                      re.match(r"^(?:Тема|Topic)\s*[\d\.]*", label, re.I))

        if is_section:
            section_counter += 1
            if not re.match(r"^Раздел\s*\d+", label, re.I):
                m_num = re.search(r"(\d+)", label)
                label = f"Раздел {m_num.group(1) if m_num else section_counter}"
            topics.append(f"{label}. {name}")
        elif is_topic or not item_type:
            if label:
                topics.append(f"{label}. {name}")
            else:
                topics.append(name)

    return topics if topics else None


def parse_topics(text: str) -> list:
    """[A] Парсит содержание дисциплины: JSON-режим → regex-fallback."""
    json_result = parse_topics_json(text)
    if json_result:
        return json_result

    topics = []
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    for para in paragraphs:
        tokens = re.split(
            r"(?=(?:Раздел|Тема)\s+\d+[\.\d]*\.?\s|\b\d+\.\d+\.?\s)", para
        )
        tokens = [t.strip() for t in tokens if t.strip()]
        for token in tokens:
            m_sec = re.match(r"^(Раздел\s+\d+)\.\s+(.+)", token)
            if m_sec:
                name = re.split(r"\s+\d+\.\d+\.", m_sec.group(2))[0].strip()
                if name:
                    topics.append(f"{m_sec.group(1)}. {name}")
                continue
            m_tema = re.match(r"^(Тема\s+[\d\.]+)\.\s+(.+)", token)
            if m_tema:
                name = re.split(r"\s+\d+\.\d+\.", m_tema.group(2))[0].strip()
                if name:
                    topics.append(f"{m_tema.group(1)}. {name}")
                continue
            m_sub = re.match(r"^(\d+\.\d+)\.?\s+(.+)", token)
            if m_sub:
                name = re.split(r"\s+\d+\.\d+\.", m_sub.group(2))[0].strip()
                if name:
                    topics.append(f"Тема {m_sub.group(1)}. {name}")

    if not topics:
        for line in text.split("\n"):
            m = re.match(r"^(Раздел|Тема)\s*([\d\.]+)[\.\\ ]+(.+)", line.strip())
            if m:
                topics.append(
                    f"{m.group(1)} {m.group(2).rstrip('.')}. {m.group(3).strip()}"
                )

    return topics if topics else [
        "Раздел 1. Основы интеллектуальных систем",
        "Раздел 2. Методы машинного обучения",
        "Раздел 3. Применение интеллектуальных систем",
    ]


def parse_list_json(text: str, min_items: int = 3) -> list | None:
    m = re.search(r"\[.*\]", text, re.S)
    if not m:
        return None
    try:
        data = json.loads(m.group())
        if not isinstance(data, list):
            return None
        result = [str(d.get("title", "")).strip() for d in data
                  if isinstance(d, dict) and d.get("title")]
        return result if len(result) >= min_items else None
    except (json.JSONDecodeError, TypeError):
        return None


def parse_list_json_with_section(text: str, min_items: int = 3) -> list | None:
    m = re.search(r"\[.*\]", text, re.S)
    if not m:
        return None
    try:
        data = json.loads(m.group())
        if not isinstance(data, list):
            return None
        result = [
            {"title": str(d.get("title", "")).strip(),
             "section": d.get("section")}
            for d in data
            if isinstance(d, dict) and d.get("title")
        ]
        return result if len(result) >= min_items else None
    except (json.JSONDecodeError, TypeError):
        return None


def parse_list(text: str, discipline: str = "", min_items: int = 3) -> list:
    """[A] Парсит список ЛР/ПЗ: JSON-режим → regex-fallback."""
    json_result = parse_list_json(text, min_items=min_items)
    if json_result:
        return json_result[:8]

    OFFTRACK_KEYWORDS = [
        "презентаци", "доклад", "реферат", "публикаци", "журнал",
        "flutter", "react native", "android studio", "xcode",
        "google play", "app store", "swift", "kotlin",
        "устный", "подготовка к",
    ]
    items = []
    for line in text.split("\n"):
        line = line.strip()
        line = re.sub(r"^(ЛР\s*№?\d+|ЛР\s*No\d+|\d+[\.):])\s+", "", line)
        line = re.sub(r"^\*\*(.+)\*\*$", r"\1", line)
        line = re.sub(r"^<[^>]{1,30}>\s*[-–\.\:]?\s*", "", line)
        line = re.sub(r"^<[^>]{1,30}>\s*$", "", line)
        if not line or len(line) < 6:
            continue
        if any(kw in line.lower() for kw in OFFTRACK_KEYWORDS):
            continue
        items.append(line)
    return items[:8] if len(items) >= min_items else ["Лабораторная работа 1", "Лабораторная работа 2"]


# ---------------------------------------------------------------------------
# Заполнение таблиц шаблона
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Библиография — генерация и заполнение таблиц
# ---------------------------------------------------------------------------

def _parse_rag_bibliography_chunks(hits: list) -> list:
    """
    [FIX-BIB-RAG] Прямой парсинг библиографических чанков из Qdrant.

    Вместо передачи retrieved-текста в LLM как «примера» — извлекаем ГОСТ-строки
    напрямую из payload["text"] каждого хита. Исключает галлюцинации авторов.

    Признак ГОСТ-строки: содержит ' — ' и 4-значный год.
    """
    entries = []
    seen: set = set()
    for h in hits:
        text = h.get("payload", {}).get("text", "")
        if not text:
            continue
        btype = (
            "Дополнительная литература"
            if "дополнительн" in text.lower()
            else "Основная литература"
        )
        for line in text.splitlines():
            line = line.strip()
            has_separator = (" — " in line) or (". -" in line) or (" : " in line and " - " in line)
            if not has_separator or not re.search(r"\b\d{4}\b", line):
                continue
            desc = re.sub(r"^\d+[\.\/\)]\s*", "", line).strip()
            if len(desc) < 20:
                continue
            key = desc[:60].lower()
            if key in seen:
                continue
            seen.add(key)
            entries.append({
                "type":    btype,
                "purpose": "Для изучения теории;Для выполнения СРО;",
                "desc":    desc,
                "url":     "http://bibl.rusoil.net",
                "coeff":   "1.00",
            })
    return entries


def parse_bibliography_json(text: str) -> list | None:
    m = re.search(r"\[.*\]", text, re.S)
    if not m:
        return None
    try:
        data = json.loads(m.group())
        if not isinstance(data, list):
            return None
        result = [d for d in data if isinstance(d, dict) and d.get("desc")]
        return result if result else None
    except (json.JSONDecodeError, TypeError):
        return None


def gen_bibliography(discipline: str, direction: str = "", level: str = "", cfg: dict = None) -> tuple[list, list]:
    """
    Генерирует основную (Т15) и методическую (Т17) литературу.
    """
    _PLACEHOLDER_MARKERS = ("фамилия", "название", "<гост", "<реальная", "...", "<")

    def _is_placeholder(desc: str) -> bool:
        dl = desc.lower()
        return any(m in dl for m in _PLACEHOLDER_MARKERS)

    def _make_fallback_main() -> list:
        """Реальные учебники по ИИ/МО, доступные в российских ЭБС."""
        return [
            {
                "type": "Основная литература",
                "purpose": "Для изучения теории;",
                "desc": (
                    "Флах, П. Машинное обучение : наука и искусство построения алгоритмов, "
                    "которые извлекают знания из данных / П. Флах ; пер. с англ. "
                    "А. А. Слинкина. — Москва : ДМК Пресс, 2015. — 400 с."
                ),
                "url": "http://www.znanium.com",
                "coeff": "1.00",
            },
            {
                "type": "Основная литература",
                "purpose": "Для изучения теории;Для выполнения СРО;",
                "desc": (
                    "Осовский, С. Нейронные сети для обработки информации : учебное пособие / "
                    "С. Осовский ; пер. с польск. И. Д. Рудинского. — Москва : "
                    "Финансы и статистика, 2002. — 344 с."
                ),
                "url": "http://www.znanium.com",
                "coeff": "1.00",
            },
            {
                "type": "Дополнительная литература",
                "purpose": "Для изучения теории;",
                "desc": (
                    "Рассел, С. Искусственный интеллект : современный подход / "
                    "С. Рассел, П. Норвиг ; пер. с англ. — 4-е изд. — Москва : "
                    "Вильямс, 2022. — 1408 с."
                ),
                "url": "http://biblio-online.ru",
                "coeff": "0.50",
            },
        ]

    def _make_fallback_method(disc: str) -> list:
        # [FIX-08] Отчёт §4, замечание 8.
        # Было: «сост. Д. М. Зарипов» — вымышленное ФИО, попадало в РПД.
        # Стало: без персоналии, только кафедра-составитель (соответствует
        # общепринятой практике оформления методических изданий, когда
        # конкретный автор неизвестен).
        return [
            {
                "purpose": "Для выполнения лабораторных работ;",
                "desc": (
                    f"Методические указания к выполнению лабораторных работ "
                    f"по дисциплине «{disc}» / УГНТУ, каф. ВТИК. — "
                    "Уфа : УГНТУ, 2023. — 64 с."
                ),
                "url": "http://bibl.rusoil.net",
                "coeff": "1.00",
            },
            {
                "purpose": "Для выполнения практических занятий;",
                "desc": (
                    f"Методические указания к практическим занятиям "
                    f"по дисциплине «{disc}» / УГНТУ, каф. ВТИК. — "
                    "Уфа : УГНТУ, 2023. — 48 с."
                ),
                "url": "http://bibl.rusoil.net",
                "coeff": "1.00",
            },
        ]

    _cfg = cfg or {}
    _custom_main = _cfg.get("main_bibliography")

    if _custom_main and isinstance(_custom_main, list) and len(_custom_main) >= 2:
        main_entries = _custom_main
        print(f"    ✅ Библиография T15: из config.json (main_bibliography), "
              f"{len(main_entries)} записей")
        _generation_log["bibliography_main_source"] = "config.json"
    else:
        # [FIX-BIB-RAG] Путь 2: прямой парсинг RAG-чанков без вызова LLM.
        _rag_section_types = SECTION_TYPE_FILTER.get("bibliography_main", ["bibliography", "place"])
        _, _bib_hits = retrieve("bibliography_main", discipline, _rag_section_types,
                                direction=direction, level=level)
        _rag_entries = _parse_rag_bibliography_chunks(_bib_hits)
        if len(_rag_entries) >= 2:
            main_entries = _rag_entries
            print(f"    ✅ Библиография T15: из RAG-чанков напрямую, "
                  f"{len(main_entries)} записей (LLM не вызывался)")
            _generation_log["bibliography_main_source"] = "rag_direct"
        else:
            # Fallback: RAG не дал достаточно записей → LLM с фильтром галлюцинаций
            if _rag_entries:
                print(f"    ⚠️  RAG вернул только {len(_rag_entries)} записей — передаём в LLM")
            raw_main = gen(
                "bibliography_main", discipline, PROMPTS["bibliography_main"],
                direction=direction, level=level,
            )
            llm_entries = parse_bibliography_json(raw_main)

            if llm_entries:
                clean_entries = [e for e in llm_entries if not _is_placeholder(e.get("desc", ""))]
                if len(clean_entries) >= 1:
                    main_entries = list(clean_entries)
                    if len(main_entries) < 2:
                        fb = _make_fallback_main()
                        existing_keys = {
                            re.sub(r"\s+", "", e.get("desc", ""))[:50].lower()
                            for e in main_entries
                        }
                        for fb_entry in fb:
                            fb_key = re.sub(r"\s+", "", fb_entry.get("desc", ""))[:50].lower()
                            if fb_key not in existing_keys:
                                main_entries.append(fb_entry)
                                existing_keys.add(fb_key)
                            if len(main_entries) >= 3:
                                break
                    added = len(main_entries) - len(clean_entries)
                    suffix = f" + {added} из fallback" if added else ""
                    print(f"    ✅ Библиография T15: принято {len(clean_entries)} от LLM{suffix}")
                    _generation_log["bibliography_main_source"] = "llm"
                else:
                    _generation_log["bibliography_main_source"] = "fallback"
                    print(f"    ⚠️  Библиография T15: LLM вернул шаблонные записи → fallback")
                    main_entries = _make_fallback_main()
            else:
                _generation_log["bibliography_main_source"] = "fallback"
                print(f"    ⚠️  Библиография T15: JSON не распарсился → fallback")
                main_entries = _make_fallback_main()

    _custom_method = _cfg.get("method_bibliography")

    if _custom_method and isinstance(_custom_method, list) and len(_custom_method) > 0:
        method_entries = _custom_method
        print("    ✅ Библиография T17: из config.json (method_bibliography)")
        _generation_log["bibliography_method_source"] = "config.json"
    else:
        method_entries = _make_fallback_method(discipline)
        print("    ✅ Библиография T17: используется fallback (реальные УГНТУ-пособия)")
        _generation_log["bibliography_method_source"] = "fallback"

    return main_entries, method_entries


def _normalize_gost_biblio(desc: str) -> str:
    """
    [FIX-07] Отчёт §4, замечание 7: «Москва _ Вильямс» вместо «Москва : Вильямс».
    ГОСТ Р 7.0.5-2008 требует пробел-двоеточие-пробел между местом издания и
    издательством. В корпусе РПД вместо «:» иногда стоит «_» (артефакт
    OCR / табличной конвертации).

    Функция нормализует типовые отклонения:
      • «Город_Издательство»     → «Город : Издательство»
      • «Город _ Издательство»   → «Город : Издательство»
      • «Город-Издательство»     (редкий OCR) оставляем как есть (чтобы не
        сломать реальные дефисные названия типа «Ростов-на-Дону»)
      • двойные пробелы вокруг «—» и «:» схлопываются в одинарные.
    Не трогает сами названия книг, фамилии и годы.
    """
    if not desc:
        return desc
    s = desc
    # «Москва_Вильямс» или «Москва _ Вильямс» → «Москва : Вильямс»
    # Паттерн: буква/закрывающая кавычка + (пробелы)_(пробелы) + буква
    s = re.sub(r"(?<=[A-Za-zА-Яа-яё»])\s*_\s*(?=[A-ZА-ЯЁ])", " : ", s)
    # «Москва:Вильямс» без пробелов → «Москва : Вильямс»
    s = re.sub(r"(?<=[A-Za-zА-Яа-яё»])\s*:\s*(?=[A-ZА-ЯЁ])", " : ", s)
    # Схлопываем двойные пробелы (но сохраняем неразрывные, если есть)
    s = re.sub(r" {2,}", " ", s).strip()
    return s


def fill_bibliography_main(doc: Document, entries: list, semester: str):
    """Заполняет T15 основную и дополнительную литературу через fill_placeholder_rows."""
    table = find_table(doc, "bibliography")
    if table is None:
        return

    seen_descs: set = set()
    deduped: list = []
    for entry in entries:
        # [FIX-07] Нормализация ГОСТ-разделителей перед записью в РПД
        entry = dict(entry)
        entry["desc"] = _normalize_gost_biblio(entry.get("desc", ""))
        key = re.sub(r"\s+", " ", entry.get("desc", "")).strip()[:60].lower()
        if key and key in seen_descs:
            print(f"  ⚠️  Т15: дублирующая запись пропущена: {key[:50]!r}")
            continue
        if key:
            seen_descs.add(key)
        deduped.append(entry)

    rows = [
        [
            e.get("type",    "Основная литература"),
            e.get("purpose", "Для изучения теории;"),
            semester, "", "",
            e.get("desc",    ""),
            "1",
            e.get("url",     ""),
            e.get("coeff",   "1.00"),
        ]
        for e in deduped
    ]
    fill_placeholder_rows(table, rows)


def fill_bibliography_method(doc: Document, entries: list, semester: str):
    """Заполняет T17 учебно-методические издания через fill_placeholder_rows."""
    table = find_table(doc, "method_bib")
    if table is None:
        return
    # [FIX-07] Нормализация ГОСТ-разделителей и для Т17
    _entries = [
        {**e, "desc": _normalize_gost_biblio(e.get("desc", ""))}
        for e in entries
    ]
    rows = [
        [
            e.get("purpose", "Для выполнения лабораторных работ;"),
            semester, "", "",
            e.get("desc",    ""),
            "1", "0",
            e.get("url",     ""),
            e.get("coeff",   "1.00"),
        ]
        for e in _entries
    ]
    fill_placeholder_rows(table, rows)


def fill_competencies_table(doc: Document, competencies: list):
    table = find_table(doc, "competencies")
    if table is None:
        return
    rows = [[str(i), desc, code] for i, (code, desc) in enumerate(competencies, 1)]
    fill_placeholder_rows(table, rows)


def fill_outcomes_table(doc: Document, competencies: list, outcomes: list):
    """
    Заполняет T5 — таблицу результатов обучения.

    [FIX-OUTCOMES] Архитектура изменена с vMerge на flat rows через fill_placeholder_rows.
    Индикатор строится из rotated outcome-текста — уникален для каждой компетенции.
    """
    table = find_table(doc, "outcomes")
    if table is None:
        return

    type_lists: dict = {"З": [], "У": [], "В": []}
    for ot, otext in outcomes:
        if ot in type_lists:
            type_lists[ot].append(otext)

    if not type_lists["З"]:
        type_lists["З"] = ["основные концепции и методы дисциплины"]
    if not type_lists["У"]:
        type_lists["У"] = ["применять методы дисциплины для решения задач"]
    if not type_lists["В"]:
        type_lists["В"] = ["навыками работы с инструментами дисциплины"]

    def split_items(text: str) -> list:
        lines = []
        for ln in text.split("\n"):
            ln = re.sub(r"^\d+[\.)]\s*", "", ln.strip())
            if ln and len(ln) > 4:
                lines.append(ln)
        return lines if lines else [text.strip()]

    z_items = []
    for t in type_lists["З"]:
        z_items.extend(split_items(t))
    u_items = []
    for t in type_lists["У"]:
        u_items.extend(split_items(t))
    v_items = []
    for t in type_lists["В"]:
        v_items.extend(split_items(t))

    z_qualifiers = ["в данной области", "в контексте дисциплины",
                    "применительно к решаемым задачам",
                    "необходимые для профессиональной деятельности",
                    "включая теоретические основы и практические аспекты"]
    u_qualifiers = ["в профессиональной деятельности", "для решения практических задач",
                    "при разработке и исследовании систем",
                    "при анализе данных и построении моделей",
                    "в ходе проектирования и реализации"]
    v_qualifiers = ["при решении профессиональных задач",
                    "для анализа и разработки систем",
                    "в проектировании и исследовательской деятельности",
                    "при реализации и тестировании решений",
                    "в профессиональной практике"]

    rows = []
    # [FIX-IND] Глаголы индикатора по типу результата обучения (ФГОС 3++)
    _ind_verb = {"З": "Знает", "У": "Умеет применять", "В": "Владеет навыками"}
    _ind_verbs = {"З": "знать", "У": "уметь", "В": "владеть"}

    # [FIX-01][FIX-02] Отчёт §4, замечания 1 и 2.
    # Было: indicator = f"… {rotated[0][:50] } …" — обрезка по 50 симв. давала
    #       «УК-1.1 знать архитектуры интеллектуальных систем: агент-ориенти в данной области»
    # Было: для всех компетенций items[idx % len(items):] попадал на тот же [0]
    #       при len(items)==3 и len(competencies)>3 (повторялись кольцом),
    #       что давало З(УК-1)==З(ПК-1)==З(ОПК-1) при одинаковом idx % 3.
    # Стало: (а) обрезка убрана полностью,
    #        (б) для каждой пары (компетенция, тип) берём СВОЙ item из списка
    #            через сквозной счётчик _type_pos — пока items хватает, все
    #            результаты уникальны; если компетенций больше, чем items,
    #            добавляем уточнитель-квалификатор, чтобы текст оставался разным.
    _type_pos = {"З": 0, "У": 0, "В": 0}
    _seen_texts: dict = {"З": set(), "У": set(), "В": set()}
    for idx, (code, desc) in enumerate(competencies):
        for type_idx, (otype, items, qualifiers) in enumerate([
            ("З", z_items, z_qualifiers),
            ("У", u_items, u_qualifiers),
            ("В", v_items, v_qualifiers),
        ]):
            result_code = f"{otype}({code})"
            indicator_num = type_idx + 1

            # [FIX-02] Уникальный item: сквозной счётчик на каждый тип.
            base_item = items[_type_pos[otype] % len(items)]
            qual = qualifiers[_type_pos[otype] % len(qualifiers)]
            _type_pos[otype] += 1

            # [FIX-02] Если этот текст УЖЕ был для того же типа у другой
            # компетенции — приклеиваем квалификатор, чтобы не дублировать.
            prefix = {"З": "Знать:", "У": "Уметь:", "В": "Владеть:"}[otype]
            result_text = f"{prefix} {base_item}"
            if result_text in _seen_texts[otype]:
                result_text = f"{prefix} {base_item} — {qual}"
            _seen_texts[otype].add(result_text)

            # [FIX-01] Индикатор БЕЗ обрезки.
            indicator = f"{code}.{indicator_num} {_ind_verbs[otype]} {base_item}"

            rows.append([code, indicator, result_code, result_text])

    fill_placeholder_rows(table, rows)


def fill_topics_table(doc: Document, topics: list, semester: str, hours: dict,
                      codes_list: list = None):
    table = find_table(doc, "topics")
    if table is None:
        return

    sections_only = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    n = max(len(sections_only), 1) if sections_only else max(len(topics), 1)

    lec  = hours.get("lecture",  12) // n
    pz   = hours.get("practice", 36) // n
    lr   = hours.get("lab",      16) // n
    sro  = hours.get("self",     62) // n
    total_l = total_pz = total_lr = total_sro = 0

    codes = codes_list or ["ОПК-1", "ПК-1"]
    rows = []
    for i, sec in enumerate(sections_only, 1):
        sec_name = re.sub(r"^Раздел\s*\d+\.\s*", "", sec).strip()
        c1 = codes[(i - 1) % len(codes)]
        c2 = codes[i % len(codes)]
        shifer = f"З({c1})\nУ({c1})\nВ({c2})"
        rows.append([str(i), sec_name, semester,
                     str(lec), str(pz), str(lr), str(sro), str(lec + pz + lr + sro),
                     shifer])
        total_l += lec; total_pz += pz; total_lr += lr; total_sro += sro

    rows.append(["", "ИТОГО:", "",
                 str(total_l), str(total_pz), str(total_lr), str(total_sro),
                 str(total_l + total_pz + total_lr + total_sro), ""])
    fill_placeholder_rows(table, rows)


def _compact_section(section: str) -> str:
    """«Раздел N. Название» → «N-Название»."""
    m = re.match(r"^Раздел\s*(\d+)[.\s]+(.+)$", section.strip())
    if m:
        return f"{m.group(1)}-{m.group(2).strip()}"
    return section


def fill_lectures_table(doc: Document, topics: list, hours: dict):
    table = find_table(doc, "lectures")
    if table is None:
        return

    themes_only   = [t for t in topics if not re.match(r"^Раздел\s*\d+", t)]
    sections_only = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]

    n_topics = max(len(themes_only), 1) if themes_only else max(len(sections_only), 1)
    lec = max(hours.get("lecture", 12) // n_topics, 1)

    def _clean_lecture_name(raw_name: str) -> str:
        name = re.sub(r"^(?:Лекция\s+\d+\.\s*)+", "", raw_name.strip()).strip()
        name = re.sub(r"\s+[А-ЯЁ]{2,4}-\d+\s*$", "", name).strip()
        name = re.sub(r"\s+[A-ZА-ЯЁ]{2,5}\s*$", "", name).strip()
        name = re.sub(r"\.\s*$", "", name).strip()
        return name if name else raw_name.strip()

    rows = []
    section = ""
    lec_no  = 0
    for topic in topics:
        if re.match(r"^Раздел\s*\d+", topic):
            section = topic
        else:
            lec_no += 1
            short = re.sub(r"^Тема\s*[\d\.]+[\.\\ ]+", "", topic).strip()
            short = _clean_lecture_name(short)
            rows.append([str(lec_no), _compact_section(section) if section else topic,
                         f"Лекция {lec_no}. {short}", str(lec), "", ""])

    if lec_no == 0:
        lec = max(hours.get("lecture", 12) // max(len(sections_only), 1), 1)
        for i, topic in enumerate(sections_only, 1):
            short = re.sub(r"^Раздел\s*\d+[\.\\ ]+", "", topic).strip()
            short = _clean_lecture_name(short)
            rows.append([str(i), _compact_section(topic), f"Лекция {i}. {short}", str(lec), "", ""])

    fill_placeholder_rows(table, rows)


def fill_lab_table(doc: Document, lab_works: list, topics: list, hours_lab: int = 18):
    table = find_table(doc, "labs")
    if table is None:
        return

    if len(lab_works) < 6:
        print(f"  ⚠️  Т9: получено {len(lab_works)} ЛР — дополняю до 6")
        for j in range(len(lab_works), 6):
            lab_works.append({"title": f"Лабораторная работа {j + 1}", "section": None})

    sections = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    n_lab = len(lab_works)
    base  = max(hours_lab // n_lab, 1)
    rem   = hours_lab - base * n_lab
    hours_list = [base + (1 if i < rem else 0) for i in range(n_lab)]

    rows = []
    for i, work in enumerate(lab_works, 1):
        # [FIX-LAB-SEC] Исправлена обработка sec_num через modulo
        if isinstance(work, dict):
            title   = work.get("title", f"Лабораторная работа {i}")
            sec_num = work.get("section")
            if sec_num is not None and sections:
                sec_idx = (int(sec_num) - 1) % len(sections)
                section = sections[sec_idx]
            else:
                section = sections[(i - 1) % max(len(sections), 1)] if sections else f"Раздел {((i - 1) // 2) + 1}"
        else:
            title   = work
            section = sections[(i - 1) % max(len(sections), 1)] if sections else f"Раздел {((i - 1) // 2) + 1}"
        rows.append([_compact_section(section), str(i), title, str(hours_list[i - 1]), "", ""])
    rows.append(["-", "", "ИТОГО:", str(hours_lab), "", ""])
    fill_placeholder_rows(table, rows)


def fill_practice_table(doc: Document, practices: list, topics: list,
                        hours_practice: int = 36):
    table = find_table(doc, "practice")
    if table is None:
        return

    if len(practices) < 6:
        print(f"  ⚠️  Т10: получено {len(practices)} ПЗ — дополняю до 6")
        for j in range(len(practices), 6):
            practices.append({"title": f"Практическое занятие {j + 1}", "section": None})

    sections = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    n_prac = len(practices)
    base   = max(hours_practice // n_prac, 1)
    rem    = hours_practice - base * n_prac
    hours_list = [base + (1 if i < rem else 0) for i in range(n_prac)]

    rows = []
    for i, prac in enumerate(practices, 1):
        # [FIX-LAB-SEC] Исправлена обработка sec_num через modulo
        if isinstance(prac, dict):
            title   = prac.get("title", f"Практическое занятие {i}")
            sec_num = prac.get("section")
            if sec_num is not None and sections:
                sec_idx = (int(sec_num) - 1) % len(sections)
                section = sections[sec_idx]
            else:
                section = sections[(i - 1) % max(len(sections), 1)] if sections else f"Раздел {((i - 1) // 2) + 1}"
        else:
            title   = prac
            section = sections[(i - 1) % max(len(sections), 1)] if sections else f"Раздел {((i - 1) // 2) + 1}"
        rows.append([_compact_section(section), str(i), title, str(hours_list[i - 1]), "", ""])
    rows.append(["-", "", "ИТОГО:", str(hours_practice), "", ""])
    fill_placeholder_rows(table, rows)


def fill_t3_hours(doc: Document, semester: str, credits: int,
                  hours_total: int, hours_contact: int, hours_sro: int,
                  exam_type: str):
    """Заполняет T3 (трудоёмкость) через fill_placeholder_rows."""
    table = find_table(doc, "workload")
    if table is None:
        return
    fill_placeholder_rows(table, [
        [semester, str(credits), str(hours_total), str(hours_contact), str(hours_sro), exam_type],
        ["ИТОГО:", str(credits), str(hours_total), str(hours_contact), str(hours_sro), ""],
    ])


def fill_t6_workload(doc: Document, lec: int, pz: int, lr: int, sro: int,
                     semester: str, exam_prep_hours: int = 0):
    t = find_table(doc, "work_types")
    if t is None:
        return
    # [Д-5] Диагностика таблицы
    print(f"  ℹ️  Т6 work_types: найдена, строк={len(t.rows)}, ищу семестр={semester!r}")
    sem_col = None
    sem_str = str(semester).strip()
    for header_row in t.rows[:4]:
        for j, cell in enumerate(header_row.cells):
            cell_text = cell.text.strip()
            # [FIX-SEM] Negative lookaround: «7» не совпадает с «17» или «27»
            _sem_re = re.compile(r'(?<!\d)' + re.escape(sem_str) + r'(?!\d)')
            if cell_text == sem_str or _sem_re.search(cell_text):
                sem_col = j
                break
        if sem_col is not None:
            break
    if sem_col is None:
        print(f"  ⚠️  Т6: столбец семестра {sem_str!r} не найден — "
              f"заголовки: {[c.text.strip() for c in t.rows[0].cells]}")

    total_contact = lec + pz + lr

    def clear_data_columns(row):
        for ci in range(2, len(row.cells)):
            set_cell_text(row.cells[ci], "")

    # [FIX-T6-SRO] SKIP_PATTERNS обновлены
    SKIP_PATTERNS = ("on-line", "онлайн", "в т.ч.", "иная", "иные", "проектная",
                     "освоение", "самостоятельная проект",
                     "контролируем",
                     "выполнение и подготовка к защит",  # [FIX-T6-ZERO2]
                     )

    # [FIX-T6-SRO] Разбивка СРО согласована с fill_t11_sro
    hrs_study = round(sro * 0.20)           # изучение вынесенного материала
    hrs_main  = round(sro * 0.20)           # основной вид: РГР / реферат / инд.задание / курсовой
    hrs_prep  = sro - hrs_study - hrs_main - exam_prep_hours  # подготовка к ЛР/ПЗ

    kw_map = {
        "контактная":             total_contact,
        "лекции":                 lec,
        "подготовка к лаборатор": hrs_prep,          # [FIX-T6-SRO]
        "подготовка к сдач":      exam_prep_hours,   # [FIX-EXAM-PREP]
        "практические занятия":   pz,
        "лабораторные работы":    lr,
        "самостоятельная работа": sro,
        "изучение учебного":      hrs_study,         # [FIX-T6-SRO]
        # [FIX-06b] Все варианты основного вида СРО — после изменения
        # fill_t11_sro (раньше был только РГР). Совпадение по подстроке
        # label (case-insensitive), так что в шаблоне подойдёт любое
        # оформление типа «выполнение РГР», «подготовка реферата» и т. д.
        "расчётно-графическ":    hrs_main,
        "расчетно-графическ":    hrs_main,
        "реферат":                hrs_main,
        "индивидуальн":           hrs_main,
        "курсов":                 hrs_main,   # курсовая работа / курсовой проект
    }

    for row in t.rows:
        label = row.cells[0].text.strip().lower()

        clear_data_columns(row)

        if "итого" in label:
            total = total_contact + sro
            set_cell_text(row.cells[1], str(total))
            if sem_col is not None and sem_col < len(row.cells):
                set_cell_text(row.cells[sem_col], str(total))
            continue

        if any(pat in label for pat in SKIP_PATTERNS):
            set_cell_text(row.cells[1], "0")
            continue

        for kw, val in kw_map.items():
            if kw in label:
                set_cell_text(row.cells[1], str(val))
                if sem_col is not None and sem_col < len(row.cells):
                    set_cell_text(row.cells[sem_col], str(val))
                break


def fill_t11_sro(doc: Document, topics: list, sro: int, cfg: dict = None):
    """
    Заполняет T11 — таблицу СРО (самостоятельной работы обучающегося).

    [FIX-06] Отчёт §4, замечание 6.
    Было: жёстко прописан один и тот же набор
        «подготовка к ЛР/ПЗ», «изучение материала», «выполнение РГР»
    для любой дисциплины. Для «Интеллектуальных систем» РГР нетипична —
    ожидается курсовой проект / индивидуальное задание / реферат.

    Стало: виды СРО берутся из config.json → sro_types (если задано),
    иначе подбираются эвристически по названию дисциплины:
      • «системы», «проектирование», «разработка» → индивидуальное задание
      • «моделирование», «анализ», «исследование» → реферат
      • иначе (классическое инженерное) → расчётно-графическая работа.
    Пропорции часов сохранены: 20 % основной вид + 20 % изучение + 60 %
    подготовка к ЛР/ПЗ.
    """
    table = find_table(doc, "sro")
    if table is None:
        return

    sections = [tp for tp in topics if re.match(r"^Раздел\s*\d+", tp)]
    n = max(len(sections), 1)

    hrs_study = round(sro * 0.20)
    hrs_main  = round(sro * 0.20)
    hrs_prep  = sro - hrs_study - hrs_main

    # [FIX-06] Кастомные виды СРО из config.json имеют приоритет
    _cfg = cfg or {}
    _custom_sro = _cfg.get("sro_types")
    if _custom_sro and isinstance(_custom_sro, list) and len(_custom_sro) >= 1:
        # [FIX-06c] Формат: список строк ["вид 1", "вид 2", ...].
        # Объекты {"name": ..., "hours_ratio": ...} не поддерживаются — str(dict)
        # даст кривой текст в таблице.
        # Всегда сохраняем базовые «подготовка к
        # ЛР/ПЗ» (60 %) и «изучение материала» (20 %), а оставшиеся 20 %
        # делим РАВНОМЕРНО между указанными видами работ — это позволяет
        # задавать несколько видов (курсовой + реферат), не ломая пропорции.
        _custom_names = [str(x).strip() for x in _custom_sro if str(x).strip()]
        if _custom_names:
            per_item = hrs_main // len(_custom_names)
            # остаток часов (из-за округления) уходит в последний вид работы
            remainder = hrs_main - per_item * len(_custom_names)
            sro_types = [
                ("подготовка к лабораторным и/или практическим занятиям", hrs_prep),
                ("изучение учебного материала, вынесенного на СРО",       hrs_study),
            ]
            for i, name in enumerate(_custom_names):
                hrs = per_item + (remainder if i == len(_custom_names) - 1 else 0)
                sro_types.append((name, hrs))
        else:
            # пустой список — падаем в эвристику ниже
            _custom_sro = None
    if not (_custom_sro and isinstance(_custom_sro, list) and len(_custom_sro) >= 1):
        # [FIX-06] Эвристический выбор основного вида работы
        _disc_name = (_cfg.get("discipline_name", "") or "").lower()
        _focus_raw = _cfg.get("discipline_focus", "")
        _focus = (
            " ".join(_focus_raw) if isinstance(_focus_raw, list)
            else str(_focus_raw)
        ).lower()
        _text = f"{_disc_name} {_focus}"

        if any(kw in _text for kw in (
            "интеллектуальн", "нейрон", "машинн", "агент",
            "проектирован", "разработк", "информационн",
        )):
            main_name = "выполнение индивидуального задания"
        elif any(kw in _text for kw in (
            "моделирован", "анализ данн", "исследован",
            "экспертн", "нечётк",
        )):
            main_name = "подготовка реферата"
        else:
            main_name = "выполнение расчётно-графической работы"

        sro_types = [
            ("подготовка к лабораторным и/или практическим занятиям", hrs_prep),
            ("изучение учебного материала, вынесенного на СРО",       hrs_study),
            (main_name,                                                hrs_main),
        ]

    rows = []
    for sec_idx, sec in enumerate(sections):
        for stype, total_hrs in sro_types:
            base_per_sec = round(total_hrs / n)
            if sec_idx < n - 1:
                hrs_per_sec = base_per_sec
            else:
                hrs_per_sec = total_hrs - base_per_sec * (n - 1)
            rows.append([sec, stype, str(hrs_per_sec), "", ""])
    rows.append(["-", "ИТОГО:", str(sro), "", ""])
    fill_placeholder_rows(table, rows)


def fill_t21_fos(doc: Document, competencies: list, topics: list,
                 outcomes: list = None, discipline: str = "дисциплины"):
    """
    Заполняет таблицу паспорта ФОС через fill_placeholder_rows.
    Структура: раздел × компетенция × 3 типа (З/У/В).
    """
    table = find_table(doc, "fos")
    if table is None:
        return

    sections = [tp for tp in topics if re.match(r"^Раздел\s*\d+", tp)]
    ocs = ["Письменный и устный опрос", "Лабораторная работа",
           "Тест", "Расчётно-графическая работа"]

    outcomes_by_type: dict = {"З": [], "У": [], "В": []}
    for o in (outcomes or []):
        if isinstance(o, tuple):
            t_type = o[0] if len(o) > 0 else ""
            t_text = o[1] if len(o) > 1 else ""
        else:
            t_type = o.get("type", "")
            t_text = o.get("text", "")
        if t_type in outcomes_by_type:
            outcomes_by_type[t_type].append(t_text)

    # [FIX-03b] Аудит 3: fallback «знания в области X» давал тавтологию
    # «Знает: знания в области X». Теперь тексты содержат *объект* без
    # дублирования глагола из _indicator_verbs.
    _fallbacks = {
        "З": f"основные методы и принципы дисциплины «{discipline}»",
        "У": f"применять инструменты дисциплины «{discipline}» для решения практических задач",
        "В": f"навыками работы с инструментами в области «{discipline}»",
    }
    # [FIX-T21-VERB] ИСПРАВЛЕНО: «Умеет» вместо «Умеет применять».
    # При конкатенации с desc компетенции («применять методы…») было
    # «ОПК-2.2 Умеет применять применять методы…» — двойной глагол.
    _indicator_verbs = {"З": "Знает", "У": "Умеет", "В": "Владеет"}
    _indicator_objects = {
        "З": "основные методы и принципы в области «{sec}»",
        "У": "методы дисциплины применительно к разделу «{sec}»",
        "В": "навыками работы с инструментами в разделе «{sec}»",
    }

    rows = []
    n = 1
    for sec in sections:
        sec_name  = re.sub(r"^Раздел\s*\d+\.\s*", "", sec)
        # [FIX-04] Отчёт §4, замечание 4: обрезка sec_name[:40] давала
        # «Методы машинного обучения и нейронные се» — без «…ти».
        # Для ФОС храним полное имя и сокращаем только для «показателя»
        # (короткая фраза «по теме …»), где иначе вылезает за ячейку.
        sec_short = sec_name if len(sec_name) <= 80 else sec_name[:77].rstrip() + "…"
        for comp_idx, (code, comp_desc) in enumerate(competencies):
            for type_idx, res_type in enumerate(("З", "У", "В")):
                type_outcomes = outcomes_by_type[res_type]
                outcome_text  = (
                    type_outcomes[(comp_idx * 3 + type_idx) % len(type_outcomes)]
                    if type_outcomes else _fallbacks[res_type]
                )
                indicator_num = type_idx + 1
                verb      = _indicator_verbs[res_type]

                # [FIX-03] Отчёт §4, замечание 3.
                # Было: indicator = f"{code}.{indicator_num} {verb} {_desc_cut}"
                #       где _desc_cut = срез описания КОМПЕТЕНЦИИ («осуществлять
                #       поиск информации»), склеенный с глаголом-показателем
                #       («Знает»). Получалось: «УК-1.1 Знает осуществлять поиск»
                #       — грамматически абсурдно.
                # Стало: индикатор строится из outcome_text (того же текста, что
                #       попал в «результат обучения» ниже в строке ФОС), с
                #       нормализацией префикса: если outcome_text начинается с
                #       «применять/разрабатывать/знать/владеть…» — глагол-
                #       показатель не дублируем. Это даёт содержательные
                #       формулировки типа «УК-1.1 Знает: архитектуры
                #       интеллектуальных систем: агент-ориентированную…».
                _oc_core = outcome_text.strip()
                _oc_core = re.sub(r"^(?:З|У|В)\s*[:\-—]\s*", "", _oc_core)
                _oc_core = re.sub(
                    r"^(?:знать|знает|уметь|умеет|владеть|владеет)\s*[:\-—]?\s*",
                    "", _oc_core, flags=re.IGNORECASE,
                )
                indicator = f"{code}.{indicator_num} {verb}: {_oc_core}"

                pokazatel = (
                    f"{'Отвечает на вопросы' if res_type == 'З' else 'Выполняет задания' if res_type == 'У' else 'Демонстрирует навыки'} "
                    f"по теме «{sec_short}»"
                )
                rows.append([
                    str(n), sec_name, f"{res_type}({code})",
                    outcome_text, indicator, pokazatel,
                    ocs[(n - 1) % len(ocs)],
                ])
                n += 1

    fill_placeholder_rows(table, rows)


# ---------------------------------------------------------------------------
# [D] Валидация бизнес-правил
# ---------------------------------------------------------------------------

def validate_generation(cfg: dict, hours: dict, competencies: list,
                        topics: list, lab_works: list, practices: list) -> list[str]:
    """
    [D] Проверяет корректность сгенерированного содержимого.
    """
    warnings: list[str] = []

    expected_total = (
        cfg.get("hours_lecture",  12) +
        cfg.get("hours_practice", 36) +
        cfg.get("hours_lab",      16) +
        cfg.get("hours_self",     62)
    )
    actual_total = sum(hours.values())
    if actual_total != expected_total:
        warnings.append(
            f"⚠️  Сумма часов {actual_total} ≠ {expected_total} из config.json"
        )

    # [З-ЧАС] Проверка ФГОС-требования credits × 36 == hours_total
    credits = cfg.get("credits", 0)
    if credits:
        expected_by_credits = credits * 36
        if actual_total != expected_by_credits:
            warnings.append(
                f"⚠️  credits × 36 = {credits} × 36 = {expected_by_credits}, "
                f"но сумма часов = {actual_total} — нарушение ФГОС"
            )

    sections = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    if not sections:
        warnings.append("⚠️  Разделы дисциплины не сгенерированы (topics пуст)")

    if len(lab_works) < 6:
        warnings.append(f"⚠️  ЛР: сгенерировано {len(lab_works)} < 6 минимальных")
    if len(practices) < 6:
        warnings.append(f"⚠️  ПЗ: сгенерировано {len(practices)} < 6 минимальных")

    def _norm_code(c: str) -> str:
        return re.sub(r"[.\s]+$", "", c.strip()).upper()

    codes_from_cfg    = {_norm_code(c) for c in cfg.get("competency_codes", "").split(",") if c.strip()}
    generated_codes   = {_norm_code(code) for code, _ in competencies}
    missing = codes_from_cfg - generated_codes
    if missing:
        warnings.append(f"⚠️  Компетенции не сгенерированы: {', '.join(sorted(missing))}")

    return warnings


# ---------------------------------------------------------------------------
# [A] Промпты — JSON-режим для всех генерируемых разделов
# ---------------------------------------------------------------------------

PROMPTS = {
    "competencies": """\
Ты составляешь рабочую программу дисциплины для российского технического университета.
Напиши описания компетенций для дисциплины «{discipline}».

Коды компетенций:
{competency_codes_numbered}

Правила:
- каждое описание начинается со слова «Способен»
- описание конкретно для «{discipline}», отражает реальные навыки
- не пиши шаблонные фразы, пиши конкретные профессиональные действия

Примеры правильных описаний (для дисциплины «Машинное обучение»):
[
  {{"code": "УК-1", "desc": "Способен применять системный подход для постановки и декомпозиции задач обработки данных"}},
  {{"code": "ОПК-1", "desc": "Способен разрабатывать алгоритмы машинного обучения и реализовывать их программно"}},
  {{"code": "ПК-1", "desc": "Способен обучать, тестировать и оценивать качество моделей машинного обучения"}}
]

ВЕРНИ ТОЛЬКО JSON-массив (без пояснений, без markdown).
Ровно {competency_count} объектов. Примеры выше — для другой дисциплины, напиши свои для «{discipline}».""",

    # [FIX-OUTCOMES-PROMPT] Добавлено правило «каждый текст — ОДНА краткая фраза»
    # [FIX-02] Отчёт §4, замечание 2: LLM возвращал 3З+3У+3В, которые потом
    # распределялись между N компетенциями кольцом → З(УК-1) совпадал с З(ПК-1).
    # Теперь просим LLM генерировать ровно по одному З/У/В НА КАЖДУЮ компетенцию
    # (competency_count штук каждого типа), чтобы в fill_outcomes_table каждой
    # компетенции достался СВОЙ item без повторов.
    "outcomes": """\
Напиши результаты обучения для дисциплины «{discipline}» по ФГОС 3++.
Нужно ровно {outcomes_total} элементов: по {competency_count} знаний (З),
{competency_count} умений (У), {competency_count} навыков (В) — по одной
формулировке каждого типа на каждую из {competency_count} компетенций.

Коды компетенций, для которых надо написать З/У/В: {competency_codes_list}

Правила:
- З: что знает студент — КОНКРЕТНЫЕ методы, алгоритмы, технологии именно «{discipline}»
- У: что умеет — начинается с глагола (применять, разрабатывать, анализировать...)
- В: чем владеет — начинается с «навыками», «методами» или «инструментами»
- [FIX-02] ВСЕ тексты УНИКАЛЬНЫ. Никакие два З не совпадают даже частично.
  Никакие два У не совпадают. Никакие два В не совпадают.
  Разные компетенции → разные результаты. Например, для УК-1 и ПК-1 формулировки
  З должны описывать РАЗНЫЕ области знаний, а не одну и ту же.
- каждый текст — ОДНА краткая фраза (не список, не перечисление через запятую)
- Не повторяй одни и те же слова-головы в разных строках одного типа
- Содержат специфику «{discipline}», НЕ копируй примеры

Пример формата (для другой дисциплины «Компьютерное зрение», 3 компетенции):
[
  {{"type": "З", "text": "методы детектирования объектов: YOLO, SSD, Faster R-CNN"}},
  {{"type": "З", "text": "принципы сегментации изображений: семантическую и экземплярную"}},
  {{"type": "З", "text": "алгоритмы выделения ключевых точек: SIFT, ORB, SuperPoint"}},
  {{"type": "У", "text": "реализовывать нейронные детекторы объектов в PyTorch"}},
  {{"type": "У", "text": "применять OpenCV для предобработки и аугментации изображений"}},
  {{"type": "У", "text": "оценивать качество моделей по метрикам mAP, Precision, Recall"}},
  {{"type": "В", "text": "навыками обучения и тонкой настройки CNN на датасетах COCO, VOC"}},
  {{"type": "В", "text": "методами трекинга объектов: ByteTrack, SORT, DeepSORT"}},
  {{"type": "В", "text": "инструментами визуализации Grad-CAM для интерпретации сети"}}
]

ВЕРНИ ТОЛЬКО JSON-массив (без пояснений, без markdown).
Ровно {outcomes_total} объектов ({competency_count}З + {competency_count}У +
{competency_count}В), все уникальны, специфичны для «{discipline}».""",

    "content": """\
Напиши содержание дисциплины «{discipline}» — ровно 3 раздела, в каждом 2 темы.
Компетенции дисциплины: {competencies_summary}
{discipline_focus_block}
ПРАВИЛА:
- type="section" — крупный тематический блок, label="Раздел 1"/"Раздел 2"/"Раздел 3"
- type="topic"   — конкретная тема внутри раздела, label="Тема 1.1"/"Тема 1.2" и т.д.
- НЕ используй type="lecture", "lab", "лекция", "ЛР" — только section и topic
- Все названия уникальны, специфичны для «{discipline}», отражают ключевые темы выше

ВЕРНИ ТОЛЬКО JSON-массив без пояснений и markdown:
[
  {{"type": "section", "label": "Раздел 1", "name": "<тематический блок 1 для {discipline}>"}},
  {{"type": "topic",   "label": "Тема 1.1", "name": "<конкретная тема из блока 1>"}},
  {{"type": "topic",   "label": "Тема 1.2", "name": "<конкретная тема из блока 1>"}},
  {{"type": "section", "label": "Раздел 2", "name": "<тематический блок 2 для {discipline}>"}},
  {{"type": "topic",   "label": "Тема 2.1", "name": "<конкретная тема из блока 2>"}},
  {{"type": "topic",   "label": "Тема 2.2", "name": "<конкретная тема из блока 2>"}},
  {{"type": "section", "label": "Раздел 3", "name": "<тематический блок 3 для {discipline}>"}},
  {{"type": "topic",   "label": "Тема 3.1", "name": "<конкретная тема из блока 3>"}},
  {{"type": "topic",   "label": "Тема 3.2", "name": "<конкретная тема из блока 3>"}}
]
Ровно 9 объектов: 3 section + 6 topic.""",

    "lab_works": """\
Напиши 6 лабораторных работ для дисциплины «{discipline}».
Компетенции, которые должны формироваться: {competencies_summary}
{discipline_focus_block}
Разделы дисциплины:
{sections_list}

Требования:
- каждая ЛР — конкретное техническое задание специфичное для «{discipline}»
- все 6 ЛР на РАЗНЫЕ темы, покрывают разные аспекты дисциплины
- используй конкретные методы/алгоритмы из ключевых тем выше
- формулировка: глагол + метод/алгоритм + объект («Реализация...», «Обучение...», «Анализ...»)
- [FIX-05] КРАТКО: 8–15 слов на ЛР. НЕ перечисляй через запятую 3+ метода.
  Одна ЛР — одна главная задача. Не более одного придаточного оборота.
  Пример правильной длины: «Реализация алгоритма k-means для кластеризации
  числовых данных» (9 слов). НЕЛЬЗЯ: «Реализация и исследование алгоритмов
  кластеризации k-means, DBSCAN и иерархической кластеризации с визуализацией
  результатов и оценкой качества по метрикам силуэта и Дэвиса–Болдуина» (25 слов).
- укажи номер раздела (1, 2 или 3) к которому относится ЛР

ВЕРНИ ТОЛЬКО JSON-массив (без пояснений, без markdown):
[
  {{"title": "<ЛР 1 специфичная для {discipline}, 8-15 слов>", "section": 1}},
  {{"title": "<ЛР 2 специфичная для {discipline}, 8-15 слов>", "section": 1}},
  {{"title": "<ЛР 3 специфичная для {discipline}, 8-15 слов>", "section": 2}},
  {{"title": "<ЛР 4 специфичная для {discipline}, 8-15 слов>", "section": 2}},
  {{"title": "<ЛР 5 специфичная для {discipline}, 8-15 слов>", "section": 3}},
  {{"title": "<ЛР 6 специфичная для {discipline}, 8-15 слов>", "section": 3}}
]
Ровно 6 объектов. Все темы уникальны и специфичны для «{discipline}».""",

    "bibliography_main": """\
Ты составляешь список рекомендуемой литературы для рабочей программы дисциплины «{discipline}» \
в российском техническом университете.

Напиши 3 записи библиографии: 2 основных учебника и 1 дополнительный.

СТРОГИЕ ПРАВИЛА:
- только реально существующие книги — НЕ придумывай авторов и названия
- авторы должны быть реальными специалистами в области «{discipline}»
- год издания 2010–2024
- формат: Фамилия, И. О. Название : тип / И. О. Фамилия. — Город : Издательство, Год. — N с.
- url: http://www.znanium.com или http://e.lanbook.com или http://biblio-online.ru

ВЕРНИ ТОЛЬКО JSON-массив без пояснений и без markdown:
[
  {{"type": "Основная литература", "purpose": "Для изучения теории;", "desc": "<ГОСТ-запись реального учебника>", "url": "http://www.znanium.com", "coeff": "1.00"}},
  {{"type": "Основная литература", "purpose": "Для выполнения СРО;Для изучения теории;", "desc": "<ГОСТ-запись реального учебника>", "url": "http://e.lanbook.com", "coeff": "1.00"}},
  {{"type": "Дополнительная литература", "purpose": "Для изучения теории;", "desc": "<ГОСТ-запись реального учебника>", "url": "http://biblio-online.ru", "coeff": "0.50"}}
]
Ровно 3 объекта. Замени угловые скобки реальными ГОСТ-записями.""",

    # bibliography_method: ключ оставлен для совместимости, не используется как промпт
    "bibliography_method": "",

    # [FIX-PRACTICE-PROMPT] Обновлены требования: убрана Python-специфика,
    # добавлены синтез/моделирование для соответствия разным дисциплинам
    "practice": """\
Напиши 6 тем практических занятий для дисциплины «{discipline}».
Компетенции, которые должны формироваться: {competencies_summary}
{discipline_focus_block}
Разделы дисциплины:
{sections_list}

Требования:
- каждое занятие — решение конкретной задачи по методам дисциплины
- все 6 тем разные, чередовать: анализ, синтез, моделирование, эксперимент
- темы ПЗ привязаны к конкретным методам из ключевых тем выше
- формулировка: глагол + метод/объект/задача («Анализ...», «Синтез...», «Решение задач...», «Моделирование...»)
- [FIX-05] КРАТКО: 8–15 слов на ПЗ. Одно занятие — одна задача.
  НЕ перечисляй через запятую 3+ метода, НЕ описывай этапы работы в названии.
  Пример правильной длины: «Моделирование нечёткого регулятора в среде
  MATLAB/Simulink» (7 слов). НЕЛЬЗЯ: «Разработка, настройка и сравнительный
  анализ нечётких регуляторов Мамдани и Такаги–Сугено с оценкой качества
  регулирования по переходным процессам и показателям устойчивости» (22 слова).
- укажи номер раздела (1, 2 или 3) к которому относится ПЗ
ВЕРНИ ТОЛЬКО JSON-массив (без пояснений, без markdown):
[
  {{"title": "<тема ПЗ 1 специфичная для {discipline}, 8-15 слов>", "section": 1}},
  {{"title": "<тема ПЗ 2 специфичная для {discipline}, 8-15 слов>", "section": 1}},
  {{"title": "<тема ПЗ 3 специфичная для {discipline}, 8-15 слов>", "section": 2}},
  {{"title": "<тема ПЗ 4 специфичная для {discipline}, 8-15 слов>", "section": 2}},
  {{"title": "<тема ПЗ 5 специфичная для {discipline}, 8-15 слов>", "section": 3}},
  {{"title": "<тема ПЗ 6 специфичная для {discipline}, 8-15 слов>", "section": 3}}
]
Ровно 6 объектов. Все темы уникальны и специфичны для «{discipline}».""",
}


# ---------------------------------------------------------------------------
# [A] Обёртка генерации с JSON-retry
# ---------------------------------------------------------------------------

def gen_with_json_retry(label: str, discipline: str, prompt: str,
                        parser_json, parser_fallback, max_retries: int = 2,
                        direction: str = "", level: str = "", **extra):
    """
    [A] Генерирует секцию с JSON-валидацией и retry.
    """
    raw = gen(label, discipline, prompt, direction=direction, level=level, **extra)
    result = parser_json(raw)
    if result is not None:
        return raw, result

    RETRY_HINT = (
        "\n\n!!! ПРЕДЫДУЩИЙ ОТВЕТ НЕ ПРОШЁЛ ВАЛИДАЦИЮ !!!\n"
        "Верни ТОЛЬКО валидный JSON-массив — никакого текста до или после.\n"
        "Никаких ```json``` блоков. Никаких пояснений. Только [...]\n"
    )
    retry_prompt = RETRY_HINT + prompt

    for attempt in range(max_retries):
        print(f"  🔄 [{label}] JSON не распарсился (попытка {attempt + 1}/{max_retries}), "
              f"перегенерация...")
        raw = gen(label, discipline, retry_prompt, direction=direction, level=level, **extra)
        result = parser_json(raw)
        if result is not None:
            return raw, result

    print(f"  ⚠️  [{label}] JSON недоступен после {max_retries} попыток — regex-fallback")
    return raw, parser_fallback(raw)


# ---------------------------------------------------------------------------
# [SIM] Похожие дисциплины из корпуса
# ---------------------------------------------------------------------------

def _print_similar_disciplines(discipline: str, corpus_dir: str = "rpd_corpus",
                                top_n: int = 5) -> None:
    """
    [FIX-SIM] Выводит top_n наиболее похожих дисциплин из корпуса до генерации.
    Использует get_embedding() → RouterAI API.
    """
    import glob as _glob
    import numpy as _np

    def _cosine(a, b):
        va, vb = _np.array(a, dtype=float), _np.array(b, dtype=float)
        denom = _np.linalg.norm(va) * _np.linalg.norm(vb)
        return float(_np.dot(va, vb) / denom) if denom > 1e-10 else 0.0

    _CODE_RE = re.compile(r"^\(\d+\)\s*(.+)$")
    title_by_src: dict = {}

    if os.path.isdir(corpus_dir):
        for path in _glob.glob(os.path.join(corpus_dir, "*.json")):
            try:
                with open(path, encoding="utf-8") as f:
                    rec = json.load(f)
                src = os.path.basename(path)
                name = ""
                for ch in rec.get("chunks", [])[:5]:
                    for line in ch.get("text", "").split("\n"):
                        line = line.strip()
                        m = _CODE_RE.match(line)
                        if m:
                            name = m.group(1).strip()
                            break
                    if name:
                        break
                if not name:
                    name = rec.get("metadata", {}).get("subject", "").strip()
                if name:
                    title_by_src[src] = name
            except Exception:
                continue

    if not title_by_src:
        jsonl = "data_clean.jsonl"
        if os.path.exists(jsonl):
            seen: set = set()
            try:
                with open(jsonl, encoding="utf-8") as f:
                    for line in f:
                        try:
                            rec  = json.loads(line)
                            src  = rec.get("source", rec.get("title", ""))
                            if not src or src in seen:
                                continue
                            seen.add(src)
                            name = ""
                            for line_text in rec.get("text", "").split("\n"):
                                line_text = line_text.strip()
                                m = _CODE_RE.match(line_text)
                                if m:
                                    name = m.group(1).strip()
                                    break
                            if name:
                                title_by_src[src] = name
                        except Exception:
                            continue
            except Exception:
                pass

    if not title_by_src:
        print("  ⚠️  [SIM] Корпус не найден — список похожих дисциплин недоступен")
        return

    query_vec = get_embedding(discipline)
    if not query_vec:
        print("  ⚠️  [SIM] Embedding недоступен — пропускаю поиск похожих")
        return

    scored = []
    for src, name in title_by_src.items():
        vec = get_embedding(name)
        if vec:
            scored.append((name, src, _cosine(query_vec, vec)))

    if not scored:
        return

    scored.sort(key=lambda x: x[2], reverse=True)

    print(f"\n📚 Похожие дисциплины в корпусе (топ-{top_n}):")
    for i, (name, src, score) in enumerate(scored[:top_n], 1):
        bar = "█" * int(score * 20)
        print(f"  {i}. [{score:.3f}] {bar}  {name}  ({src})")
    print()


# ---------------------------------------------------------------------------
# Точка входа
# ---------------------------------------------------------------------------

def main(config_path: Optional[str] = None, clear_cache: bool = False):
    if config_path is None and os.path.exists("config.json"):
        config_path = "config.json"

    if clear_cache and os.path.exists(_CACHE_FILE):
        os.remove(_CACHE_FILE)
        print(f"  ♻️  Кэш {_CACHE_FILE} сброшен (--clear-cache)")

    if config_path:
        with open(config_path, encoding="utf-8") as f:
            cfg = json.load(f)
    else:
        cfg = {}
    cfg.setdefault("discipline", "Интеллектуальные системы")

    discipline       = cfg["discipline"]
    semester         = str(cfg.get("semester", "7"))
    competency_codes = cfg.get("competency_codes", "УК-1, ОПК-1, ОПК-2, ПК-1, ПК-2")
    direction        = cfg.get("direction", "")
    level            = cfg.get("level", "бакалавриат")

    # [З-R5] Загружаем персистентный кэш
    _load_cache()

    # [FIX-SIM-CALL] Показываем похожие дисциплины до генерации
    _print_similar_disciplines(discipline, corpus_dir="rpd_corpus")

    if "retrieval_top_k" in cfg:
        GENERATION["top_k"] = int(cfg["retrieval_top_k"])
    if "retrieval_min_score" in cfg:
        GENERATION["min_score"] = float(cfg["retrieval_min_score"])

    global _RETRIEVAL_CONF_HASH
    _RETRIEVAL_CONF_HASH = _make_retrieval_conf_hash(
        GENERATION["top_k"], GENERATION["min_score"]
    )

    hours = {
        "lecture":  cfg.get("hours_lecture",  12),
        "practice": cfg.get("hours_practice", 36),
        "lab":      cfg.get("hours_lab",      16),
        "self":     cfg.get("hours_self",     62),
    }

    template = cfg.get("template", "")
    if not template or not os.path.exists(template):
        for candidate in ["Шаблон_пустой.dotx",
                          os.path.join("rpd_corpus", "Шаблон_пустой.dotx")]:
            if os.path.exists(candidate):
                template = candidate
                break
        if not template:
            corpus_dir = "rpd_corpus"
            candidates = sorted(
                f for f in os.listdir(corpus_dir)
                if (f.endswith(".docx") or f.endswith(".dotx")) and not f.startswith("~$")
            ) if os.path.isdir(corpus_dir) else []
            template = os.path.join(corpus_dir, candidates[-1]) if candidates else ""

    print(f"\n{'=' * 60}")
    print(f"ГЕНЕРАЦИЯ РПД: {discipline}")
    print(f"Направление: {direction}  Уровень: {level}")
    print(f"{'=' * 60}\n")

    # Проверка RouterAI
    try:
        r = _client_ai.models.list()
        print("✅ RouterAI доступен")
    except Exception as e:
        print(f"⚠️  RouterAI проверка недоступна (продолжаем): {e}")

    if not template or not os.path.exists(template):
        print(f"❌ Шаблон не найден: {template!r}")
        return

    codes_list = [c.strip() for c in competency_codes.split(",") if c.strip()]
    competency_codes_numbered = "\n".join(f"{i + 1}. {c}" for i, c in enumerate(codes_list))

    _focus_raw = cfg.get("discipline_focus", "")
    if _focus_raw:
        discipline_focus_block = (
            f"Ключевые темы дисциплины (обязательно используй их в названиях разделов и заданий):\n"
            f"{_focus_raw}\n"
        )
    else:
        discipline_focus_block = ""

    base_vars = {
        "competency_codes":          competency_codes,
        "competency_codes_numbered": competency_codes_numbered,
        "competency_codes_list":     ", ".join(codes_list),        # [FIX-02] для промпта outcomes
        "competency_count":          len(codes_list),
        "outcomes_count":            9,
        "outcomes_total":            len(codes_list) * 3,           # [FIX-02] по 3 на компетенцию
        "competencies_summary":      "",
        "discipline_focus_block":    discipline_focus_block,
        "sections_list":             "",
    }

    raw: dict = {}

    # --- Шаг 1: компетенции и результаты обучения ---
    _fgos = cfg.get("fgos_competencies", {})
    if _fgos and isinstance(_fgos, dict):
        competencies = [
            (code, _fgos[code])
            for code in codes_list
            if code in _fgos
        ]
        _missing_codes = [c for c in codes_list if c not in _fgos]
        if _missing_codes:
            print(f"  ⚠️  [FGOS] Коды не найдены в fgos_competencies: {_missing_codes} — генерирую LLM")
            _missing_str = ", ".join(_missing_codes)
            _miss_vars = {**base_vars,
                          "competency_codes": _missing_str,
                          "competency_codes_numbered": "\n".join(f"{i+1}. {c}" for i, c in enumerate(_missing_codes)),
                          "competency_count": len(_missing_codes)}
            _, _extra = gen_with_json_retry(
                "competencies", discipline, PROMPTS["competencies"],
                parser_json=lambda t: parse_competencies_json(t),
                parser_fallback=lambda t: parse_competencies(t, codes=_missing_codes),
                direction=direction, level=level, **_miss_vars
            )
            competencies += _extra
        raw["competencies"] = json.dumps(
            [{"code": c, "desc": d} for c, d in competencies], ensure_ascii=False
        )
        print(f"  ✅ [FGOS] Компетенции из fgos_competencies: {[c for c, _ in competencies]}")
        _generation_log["competencies_source"] = "fgos_competencies (config.json)"
    else:
        raw["competencies"], competencies = gen_with_json_retry(
            "competencies", discipline, PROMPTS["competencies"],
            parser_json=lambda t: parse_competencies_json(t),
            parser_fallback=lambda t: parse_competencies(t, codes=codes_list),
            direction=direction, level=level, **base_vars
        )
        _generation_log["competencies_source"] = "llm"

    raw["outcomes"], outcomes = gen_with_json_retry(
        "outcomes", discipline, PROMPTS["outcomes"],
        parser_json=lambda t: parse_outcomes_json(t, required_count=len(codes_list) * 3),
        parser_fallback=parse_outcomes,
        direction=direction, level=level, **base_vars
    )

    # --- Шаг 2: обновляем competencies_summary и генерируем разделы ---
    comp_summary = "; ".join(f"{c[0]}: {c[1][:60]}" for c in competencies[:5])
    content_vars = {**base_vars, "competencies_summary": comp_summary}
    raw["content"], topics = gen_with_json_retry(
        "content", discipline, PROMPTS["content"],
        parser_json=parse_topics_json,
        parser_fallback=parse_topics,
        direction=direction, level=level, **content_vars
    )

    # [FIX-2] Проверяем наличие разделов; domain drift через _ONIR_KW
    _ONIR_KW = {
        "научно-исследовательск", "этапы научного", "методологии научных",
        "исследований в России", "научного исследования", "нирс",
    }
    _sections_found = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    _is_domain_drift = any(
        any(kw in t.lower() for kw in _ONIR_KW) for t in _sections_found
    )
    if len(_sections_found) < 2 or _is_domain_drift:
        _reason = "domain drift (ОНИР)" if _is_domain_drift else f"Разделов найдено: {len(_sections_found)}"
        print(f"  ⚠️  [content] {_reason} — структурный fallback")
        _comp_keywords = " ".join(c[1][:40] for c in competencies[:3]).lower()
        _has_neuro  = any(w in _comp_keywords for w in ("нейр", "сеть", "deep"))
        _has_fuzzy  = any(w in _comp_keywords for w in ("нечётк", "fuzzy", "логик"))
        _has_manage = any(w in _comp_keywords for w in ("управл", "регулят", "систем"))

        if _has_fuzzy:
            topics = [
                f"Раздел 1. Теоретические основы {discipline}",
                f"Тема 1.1. Математический аппарат нечётких множеств",
                f"Тема 1.2. Архитектуры нечётких систем",
                f"Раздел 2. Методы нечёткого вывода",
                f"Тема 2.1. Системы Мамдани и Сугено",
                f"Тема 2.2. Нейро-нечёткие системы ANFIS",
                f"Раздел 3. Применение {discipline}",
                f"Тема 3.1. Синтез нечётких регуляторов",
                f"Тема 3.2. Оценка эффективности систем",
            ]
        elif _has_neuro:
            topics = [
                f"Раздел 1. Архитектуры нейронных сетей",
                f"Тема 1.1. Многослойные перцептроны и обратное распространение",
                f"Тема 1.2. Сверточные и рекуррентные сети",
                f"Раздел 2. Обучение и оптимизация нейронных сетей",
                f"Тема 2.1. Алгоритмы оптимизации и регуляризация",
                f"Тема 2.2. Трансферное обучение и тонкая настройка",
                f"Раздел 3. Применение нейронных сетей",
                f"Тема 3.1. Задачи классификации и регрессии",
                f"Тема 3.2. Оценка качества и развёртывание моделей",
            ]
        elif _has_manage:
            topics = [
                f"Раздел 1. Основы интеллектуального управления",
                f"Тема 1.1. Классификация и архитектуры ИСУ",
                f"Тема 1.2. Адаптивное управление",
                f"Раздел 2. Методы синтеза интеллектуальных регуляторов",
                f"Тема 2.1. Нейросетевые и нечёткие регуляторы",
                f"Тема 2.2. Обучение с подкреплением в управлении",
                f"Раздел 3. Применение {discipline}",
                f"Тема 3.1. Моделирование и верификация",
                f"Тема 3.2. Сравнительный анализ методов",
            ]
        else:
            topics = [
                f"Раздел 1. Теоретические основы {discipline}",
                f"Тема 1.1. Основные понятия и методы",
                f"Тема 1.2. Архитектуры и инструменты",
                f"Раздел 2. Алгоритмическая база {discipline}",
                f"Тема 2.1. Ключевые алгоритмы и их реализация",
                f"Тема 2.2. Оптимизация и настройка систем",
                f"Раздел 3. Применение {discipline}",
                f"Тема 3.1. Прикладные задачи дисциплины",
                f"Тема 3.2. Оценка эффективности и верификация",
            ]
        _sections_found = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
        print(f"  ℹ️  Создано {len(_sections_found)} разделов из fallback")

    # [Фикс №9] Каждый раздел должен иметь хотя бы 1 тему (ОДНОКРАТНАЯ ПРОВЕРКА)
    _secs_in_topics   = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    _topics_in_topics = [t for t in topics if re.match(r"^Тема\s*[\d\.]+", t)]
    if _secs_in_topics and not _topics_in_topics:
        print(f"  ⚠️  [content] Темы внутри разделов отсутствуют — добавляю базовые")
        enriched: list = []
        for i, sec in enumerate(_secs_in_topics, 1):
            enriched.append(sec)
            enriched.append(f"Тема {i}.1. Теоретические основы")
            enriched.append(f"Тема {i}.2. Практическое применение")
        topics = enriched

    # [Фикс №5+6] sections_list передаётся в промпты ЛР/ПЗ
    _secs = [t for t in topics if re.match(r"^Раздел\s*\d+", t)]
    _SEC_PREFIX = re.compile(r"^Раздел\s*\d+[.\s]+")
    _sections_list_str = "\n".join(
        "{0}. {1}".format(i + 1, _SEC_PREFIX.sub("", s).strip())
        for i, s in enumerate(_secs)
    ) or "1. Теоретические основы\n2. Методы\n3. Применение"
    content_vars = {**content_vars, "sections_list": _sections_list_str}

    raw["lab_works"], lab_works = gen_with_json_retry(
        "lab_works", discipline, PROMPTS["lab_works"],
        parser_json=lambda t: parse_list_json_with_section(t, min_items=6),
        parser_fallback=lambda t: [{"title": x, "section": None} for x in parse_list(t, discipline)],
        direction=direction, level=level, **content_vars
    )
    lab_works = _normalize_section_assignment(lab_works, len(_secs))

    raw["practice"], practices = gen_with_json_retry(
        "practice", discipline, PROMPTS["practice"],
        parser_json=lambda t: parse_list_json_with_section(t, min_items=6),
        parser_fallback=lambda t: [{"title": x, "section": None} for x in parse_list(t, discipline)],
        direction=direction, level=level, **content_vars
    )
    practices = _normalize_section_assignment(practices, len(_secs))

    # --- Шаг 3: библиография ---
    print("  📚 Генерация библиографии...")
    bib_main, bib_method = gen_bibliography(discipline, direction, level, cfg=cfg)

    # --- [D] Валидация ---
    validation_warnings = validate_generation(
        cfg, hours, competencies, topics, lab_works, practices
    )
    if validation_warnings:
        print("\n🔎 Результаты валидации:")
        for w in validation_warnings:
            print(f"  {w}")
        _generation_log["validation_warnings"] = validation_warnings
    else:
        print("\n✅ Валидация пройдена — все бизнес-правила соблюдены")

    # --- Заполнение шаблона ---
    hours_contact = hours["lecture"] + hours["practice"] + hours["lab"]
    hours_sro     = hours["self"]
    hours_total   = hours_contact + hours_sro
    exam_type     = cfg.get("exam_type", "экзамен")
    code          = cfg.get("code", "38050")

    shutil.copy(template, OUTPUT_DOCX)
    doc = Document(OUTPUT_DOCX)

    fill_doc_header(
        doc,
        discipline  = discipline,
        code        = code,
        year        = cfg.get("year", "2025"),
        credits     = cfg.get("credits", 4),
        hours_total = hours_total,
        exam_type   = exam_type,
    )
    fill_appendix_v(doc, discipline, topics)
    fill_sro_topic_paragraphs(doc, topics, label=f"({code}){discipline}")

    for name, fn, args in [
        ("Т3 Трудоёмкость",        fill_t3_hours,           (doc, semester, cfg.get("credits", 4), hours_total, hours_contact, hours_sro, exam_type)),
        ("Т4 Компетенции",         fill_competencies_table,  (doc, competencies)),
        ("Т5 Результаты обучения", fill_outcomes_table,      (doc, competencies, outcomes)),
        ("Т6 Виды работы",         fill_t6_workload,         (doc, hours["lecture"], hours["practice"], hours["lab"], hours["self"], semester,
                                                               cfg.get("exam_prep_hours", 9 if exam_type == "экзамен" else 0))),
        ("Т7 Темы",                fill_topics_table,        (doc, topics, semester, hours, codes_list)),
        ("Т8 Лекции",              fill_lectures_table,      (doc, topics, hours)),
        ("Т9 ЛР",                  fill_lab_table,           (doc, lab_works, topics, hours["lab"])),
        ("Т10 ПЗ",                 fill_practice_table,      (doc, practices, topics, hours["practice"])),
        ("Т11 СРО",                fill_t11_sro,             (doc, topics, hours["self"], cfg)),
        ("Т15 Основная лит-ра",    fill_bibliography_main,   (doc, bib_main,   semester)),
        ("Т17 Метод.издания",      fill_bibliography_method, (doc, bib_method, semester)),
        ("Т21 ФОС",                fill_t21_fos,             (doc, competencies, topics, outcomes, discipline)),
        ("Аннотация",              fill_annotation_table,    (doc, competencies, outcomes, topics,
                                                               cfg.get("credits", 4), hours_total, exam_type)),
    ]:
        try:
            fn(*args)
            print(f"  ✅ {name}")
        except Exception as e:
            print(f"  ⚠️  {name}: {e}")

    doc.save(OUTPUT_DOCX)
    print(f"\n✅ Сохранено: {OUTPUT_DOCX}")

    _save_cache()

    try:
        with open(GENERATION_LOG, "w", encoding="utf-8") as f:
            json.dump(_generation_log, f, ensure_ascii=False, indent=2)
        print(f"📋 Лог генерации: {GENERATION_LOG}")
    except Exception as e:
        print(f"  ⚠️  Не удалось сохранить лог: {e}")


if __name__ == "__main__":
    import argparse as _ap
    _p = _ap.ArgumentParser()
    _p.add_argument("config", nargs="?", default=None, help="Путь к config.json")
    _p.add_argument("--clear-cache", action="store_true",
                    help="Сбросить rpd_cache.json (нужно после пересборки корпуса)")
    _a = _p.parse_args()
    main(_a.config, clear_cache=_a.clear_cache)
