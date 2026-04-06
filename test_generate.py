"""
test_generate.py — генерация тестовых заданий (ФОС) по РПД.

Архитектура: аналог rpd_generate.py — те же Qdrant/Ollama/кэш,
новый pipeline: RPD → разделы → RAG → LLM → вопросы → DOCX.

Формат нумерации (ГОСТ):
    КодДисц.Раздел.Тема.ТипВопроса.НомерВопроса.Ранг.Послед(КолПравильных)
    38050.1.1.1.1001.1.0(1)

Требования из методических указаний:
    - Минимум 30 вопросов на раздел
    - Минимум 100 вопросов на компетенцию
    - Минимум 3 ранга сложности

Запуск:
    python test_generate.py                    # полная генерация
    python test_generate.py --section 1        # только раздел 1
    python test_generate.py --no-rag           # без Qdrant (offline-режим)
    python test_generate.py --questions-per-rank 15  # 15 вопросов на ранг

Выходные файлы:
    output_tests.docx      — тесты в ГОСТ-формате
    coverage_report.json   — покрытие компетенций
    test_cache.json        — кэш эмбеддингов/retrieval
"""

import argparse
import json
import re
import sys
import time
import hashlib
from pathlib import Path
from typing import Optional

import requests
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Конфигурация (зеркало rpd_generate.py) ───────────────────────────────────
CONFIG_PATH  = Path("config.json")
RPD_PATH     = Path("output_rpd.docx")
OUTPUT_DOCX  = "output_tests.docx"
COVERAGE_LOG = "coverage_report.json"
_CACHE_FILE  = "test_cache.json"

QDRANT = {"url": "http://localhost:6333", "collection": "rpd_rag"}
OLLAMA = {
    "embed_url":    "http://localhost:11434/api/embed",
    "generate_url": "http://localhost:11434/api/generate",
    "embed_model":  "bge-m3",
    "llm_model":    "qwen2.5:14b",
}

GENERATION = {
    "top_k":     6,
    "min_score": 0.40,
}

MAX_CONTEXT_CHARS = 4000   # меньше, чем в rpd_generate — промпт для тестов длиннее

# Ранги сложности
RANKS = {
    1: "знание фактов и определений",
    2: "понимание и применение понятий",
    3: "анализ и синтез (несколько правильных ответов)",
}

# Минимальное число вопросов
MIN_PER_RANK    = 10   # базовый минимум; итого 30 на раздел при 3 рангах
MIN_PER_SECTION = 30
MIN_PER_COMP    = 100

# Типы вопросов (поле в нумерации):
#   1 — один правильный ответ, порядок не важен
#   2 — несколько правильных ответов
#   3 — установить последовательность
QTYPE_SINGLE   = 1
QTYPE_MULTIPLE = 2
QTYPE_ORDER    = 3

# ── Кэш (аналогично rpd_generate.py) ─────────────────────────────────────────
EMBED_CACHE:    dict = {}
RETRIEVE_CACHE: dict = {}


def _load_cache() -> None:
    global EMBED_CACHE, RETRIEVE_CACHE
    if Path(_CACHE_FILE).exists():
        try:
            data = json.loads(Path(_CACHE_FILE).read_text(encoding="utf-8"))
            EMBED_CACHE    = data.get("embed",    {})
            RETRIEVE_CACHE = data.get("retrieve", {})
            print(f"📦 Кэш загружен: {len(EMBED_CACHE)} эмбеддингов, "
                  f"{len(RETRIEVE_CACHE)} retrieval-записей")
        except Exception as e:
            print(f"⚠️  Ошибка загрузки кэша: {e}")


def _save_cache() -> None:
    try:
        Path(_CACHE_FILE).write_text(
            json.dumps({"embed": EMBED_CACHE, "retrieve": RETRIEVE_CACHE},
                       ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
    except Exception as e:
        print(f"⚠️  Ошибка сохранения кэша: {e}")


# ── Ollama / Qdrant ───────────────────────────────────────────────────────────
def clean(text: str) -> str:
    """Базовая очистка ответа LLM."""
    return text.strip()


def get_embedding(text: str) -> list:
    if text in EMBED_CACHE:
        return EMBED_CACHE[text]
    for attempt in range(3):
        try:
            r = requests.post(
                OLLAMA["embed_url"],
                json={"model": OLLAMA["embed_model"], "input": f"query: {text}"},
                timeout=60)
            r.raise_for_status()
            d = r.json()
            embeddings = d.get("embeddings")
            if embeddings and isinstance(embeddings, list) and embeddings[0]:
                vec = embeddings[0]
                EMBED_CACHE[text] = vec
                return vec
            vec = d.get("embedding")
            if not vec:
                data_list = d.get("data") or []
                vec = data_list[0].get("embedding") if data_list else None
            if vec:
                EMBED_CACHE[text] = vec
                return vec
        except Exception as e:
            if attempt == 2:
                print(f"  ⚠️  Ошибка эмбеддинга: {e}")
                return []
            time.sleep(2 ** attempt)
    return []


def _search_qdrant(vec: list, payload_filter: Optional[dict], top_k: int) -> list:
    try:
        body = {"query": vec, "limit": top_k, "with_payload": True}
        if payload_filter:
            body["filter"] = payload_filter
        r = requests.post(
            f"{QDRANT['url']}/collections/{QDRANT['collection']}/points/query",
            json=body, timeout=30)
        r.raise_for_status()
        return r.json().get("result", {}).get("points", [])
    except requests.HTTPError:
        # Fallback для старых версий Qdrant
        body = {"vector": vec, "limit": top_k, "with_payload": True}
        if payload_filter:
            body["filter"] = payload_filter
        r = requests.post(
            f"{QDRANT['url']}/collections/{QDRANT['collection']}/points/search",
            json=body, timeout=30)
        r.raise_for_status()
        return r.json().get("result", [])


def retrieve_for_section(section_name: str, discipline: str,
                          no_rag: bool = False) -> str:
    """
    RAG-поиск релевантного контекста для раздела.
    Ищет в book_content и content.
    Возвращает строку контекста для вставки в промпт.
    """
    if no_rag:
        return ""

    cache_key = f"tests|{section_name}|{discipline}"
    if cache_key in RETRIEVE_CACHE:
        return RETRIEVE_CACHE[cache_key]

    queries = [
        f"{discipline}: {section_name}",
        f"тестовые вопросы {section_name}",
        f"определения понятия {section_name}",
    ]

    section_types = ["book_content", "content"]
    payload_filter = {
        "must": [{
            "should": [
                {"key": "section_type", "match": {"value": st}}
                for st in section_types
            ]
        }]
    }

    all_hits: dict = {}
    for query_text in queries:
        vec = get_embedding(query_text)
        if not vec:
            continue
        hits = _search_qdrant(vec, payload_filter, GENERATION["top_k"])
        for h in hits:
            hid = h.get("id")
            if hid not in all_hits or h.get("score", 0) > all_hits[hid].get("score", 0):
                all_hits[hid] = h

    # Дедупликация по источнику (max 2 чанка на источник)
    MAX_PER_SOURCE = 2
    source_counts: dict = {}
    good_hits: list = []
    for h in sorted(all_hits.values(), key=lambda h: h.get("score", 0), reverse=True):
        if h.get("score", 0) < GENERATION["min_score"]:
            continue
        src = h.get("payload", {}).get("source_file", h.get("payload", {}).get("source", ""))
        if source_counts.get(src, 0) < MAX_PER_SOURCE:
            source_counts[src] = source_counts.get(src, 0) + 1
            good_hits.append(h)
        if len(good_hits) >= GENERATION["top_k"]:
            break

    # Если ничего не нашли — fallback без фильтра
    if not good_hits:
        vec = get_embedding(queries[0])
        if vec:
            hits = _search_qdrant(vec, None, GENERATION["top_k"])
            good_hits = sorted(
                [h for h in hits if h.get("score", 0) >= GENERATION["min_score"] * 0.7],
                key=lambda h: h.get("score", 0), reverse=True
            )[:GENERATION["top_k"]]

    print(f"    🔍 RAG [{section_name[:40]}]: {len(good_hits)} чанков "
          f"(scores: {[round(h.get('score', 0), 3) for h in good_hits]})")

    # Сборка контекстной строки
    seen: set = set()
    parts: list = []
    total = 0
    for h in good_hits:
        text = h.get("payload", {}).get("text", "").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        chunk = text[:800]
        if total + len(chunk) > MAX_CONTEXT_CHARS:
            break
        parts.append(chunk)
        total += len(chunk)

    ctx = "\n\n---\n\n".join(parts)
    RETRIEVE_CACHE[cache_key] = ctx
    return ctx


# ── LLM ───────────────────────────────────────────────────────────────────────
def llm(prompt: str, max_tokens: int = 1200) -> str:
    for attempt in range(3):
        try:
            r = requests.post(
                OLLAMA["generate_url"],
                json={
                    "model": OLLAMA["llm_model"],
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.4,
                        "num_predict": max_tokens,
                        "num_ctx": 8192,
                    }
                },
                timeout=300)
            r.raise_for_status()
            text = r.json().get("response", "")
            if text:
                return clean(text)
        except Exception as e:
            if attempt == 2:
                return f"[Ошибка LLM: {e}]"
            time.sleep(5)
    return "[Ошибка: пустой ответ]"


# ── Парсинг РПД ───────────────────────────────────────────────────────────────

# Маппинг раздел → компетенции (из таблицы FOS в РПД)
# Парсится автоматически, но если не удалось — fallback
_DEFAULT_SECTION_COMP = {
    1: ["УК-1", "ОПК-1"],
    2: ["ОПК-1", "ОПК-2"],
    3: ["ОПК-2", "ПК-1", "ПК-2"],
}


def parse_rpd_sections(rpd_path: Path) -> list[dict]:
    """
    Извлекает разделы дисциплины из output_rpd.docx.

    Возвращает список:
    [
        {
            "num":          1,
            "name":         "Нейронные сети и глубокое обучение",
            "competencies": ["УК-1", "ОПК-1"],
            "topics":       ["Архитектуры нейронных сетей...", "Глубокое обучение..."],
        },
        ...
    ]
    """
    doc = Document(str(rpd_path))

    # Собираем весь текст параграфов
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    sections = []
    current_section = None
    section_pattern = re.compile(
        r"Раздел\s+(\d+)[.\s:]+(.+)", re.IGNORECASE
    )
    topic_pattern = re.compile(
        r"^(?:Тема\s+\d+[.\s:]+|[-–•]\s*)(.+)$"
    )

    for text in paragraphs:
        m = section_pattern.match(text)
        if m:
            if current_section:
                sections.append(current_section)
            current_section = {
                "num":          int(m.group(1)),
                "name":         m.group(2).strip().rstrip("."),
                "competencies": _DEFAULT_SECTION_COMP.get(int(m.group(1)), ["УК-1"]),
                "topics":       [],
            }
            continue

        # Попытка поймать темы внутри раздела
        if current_section:
            tm = topic_pattern.match(text)
            if tm and len(text) > 15:
                current_section["topics"].append(tm.group(1).strip())
            elif len(text) > 30 and not section_pattern.match(text):
                # Контентные строки без маркера — могут быть темами
                if not any(kw in text for kw in
                           ["Трудоем", "Семестр", "Форма", "Кафедр", "Зачетн",
                            "ИТОГО", "подготовка", "выполнение", "изучение"]):
                    current_section["topics"].append(text[:120])

    if current_section:
        sections.append(current_section)

    # Если разделы не нашлись стандартным способом — ищем в таблицах
    if not sections:
        sections = _parse_sections_from_tables(doc)

    # Добавляем компетенции из таблиц РПД (перезаписывает дефолтный маппинг)
    _enrich_with_comp_from_tables(doc, sections)

    print(f"📋 РПД: найдено {len(sections)} разделов")
    for s in sections:
        print(f"   Раздел {s['num']}: {s['name'][:60]}")
        print(f"           компетенции: {', '.join(s['competencies'])}")
        print(f"           тем: {len(s['topics'])}")

    return sections


def _parse_sections_from_tables(doc: Document) -> list[dict]:
    """Fallback: ищем разделы в таблицах РПД."""
    sections = []
    sec_re = re.compile(r"Раздел\s+(\d+)[.\s:–]+(.+)", re.IGNORECASE)
    seen_nums: set = set()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                m = sec_re.match(text)
                if m:
                    num = int(m.group(1))
                    if num not in seen_nums:
                        seen_nums.add(num)
                        sections.append({
                            "num":          num,
                            "name":         m.group(2).strip().rstrip("."),
                            "competencies": _DEFAULT_SECTION_COMP.get(num, ["УК-1"]),
                            "topics":       [],
                        })

    return sorted(sections, key=lambda s: s["num"])


def _enrich_with_comp_from_tables(doc: Document, sections: list[dict]) -> None:
    """
    Читает таблицу FOS/тематического плана и сопоставляет разделы с компетенциями.
    Работает по эвристике: если в строке таблицы есть номер раздела и коды компетенций.
    """
    comp_re = re.compile(r"\b([УО]К-\d+|ОПК-\d+|ПК-\d+)\b")

    # Строим словарь num → set компетенций
    comp_map: dict[int, set] = {s["num"]: set() for s in sections}
    sec_nums = {s["num"] for s in sections}

    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells)
            # Ищем номер раздела в строке
            sec_m = re.search(r"\b(\d+)\b", row_text)
            if not sec_m:
                continue
            num = int(sec_m.group(1))
            if num not in sec_nums:
                continue
            comps = comp_re.findall(row_text)
            if comps:
                comp_map[num].update(comps)

    # Применяем найденные компетенции
    for s in sections:
        found = sorted(comp_map.get(s["num"], set()))
        if found:
            s["competencies"] = found


# ── Генерация вопросов ────────────────────────────────────────────────────────

RANK_PROMPTS = {
    1: """Сгенерируй {n} тестовых вопросов РАНГА 1 (знание фактов и определений).
Правила для ранга 1:
- Вопрос требует вспомнить определение, термин или факт
- Ровно ОДИН правильный ответ из 3-х вариантов
- Варианты ответов короткие (1-2 предложения)
- Неправильные ответы правдоподобны, но однозначно ошибочны
""",
    2: """Сгенерируй {n} тестовых вопросов РАНГА 2 (понимание и применение).
Правила для ранга 2:
- Вопрос требует объяснить принцип, выбрать метод или сравнить подходы
- Ровно ОДИН правильный ответ из 4-х вариантов
- Варианты ответов могут быть развёрнутыми (2-3 предложения)
- Дистракторы содержат правдоподобные но неверные утверждения
""",
    3: """Сгенерируй {n} тестовых вопросов РАНГА 3 (анализ и синтез).
Правила для ранга 3:
- Вопрос требует анализа, синтеза или выбора нескольких верных утверждений
- НЕСКОЛЬКО правильных ответов из 4-5 вариантов (2-4 правильных)
- Формулировка вопроса явно указывает «Выберите ВСЕ верные утверждения» или аналог
- Дистракторы — частично верные или инвертированные утверждения
""",
}

ANSWER_LABELS = ["А", "Б", "В", "Г", "Д"]

_PROMPT_TEMPLATE = """\
Ты — преподаватель, составляющий фонд оценочных средств (ФОС) для дисциплины «{discipline}».

Раздел: {section_name}
Тема раздела: {topic}
Компетенции, закрываемые разделом: {competencies}

{rank_prompt}

Контекст из учебников (используй для формулировок):
{context}

=== ФОРМАТ ОТВЕТА — строго соблюдать ===
Для каждого вопроса выводи блок:

ЗАДАНИЕ: <текст вопроса>
А) <вариант А>
Б) <вариант Б>
В) <вариант В>
Г) <вариант Г>
ПРАВИЛЬНЫЙ: <буква(ы) через запятую, например: А или А, В>
===

Выведи ровно {n} таких блоков. Без нумерации, без пояснений вне блоков.
"""


def _parse_questions_from_llm(raw: str, rank: int,
                               section_num: int, topic_num: int,
                               discipline_code: str,
                               start_idx: int) -> list[dict]:
    """
    Парсит ответ LLM в список вопросов.
    Каждый вопрос — dict со всеми полями для нумерации и вывода в DOCX.
    """
    questions = []

    # Разбиваем по «ЗАДАНИЕ:» как разделителю блоков
    blocks = re.split(r"(?=ЗАДАНИЕ\s*:)", raw, flags=re.IGNORECASE)
    blocks = [b.strip() for b in blocks if b.strip() and "ЗАДАНИЕ" in b.upper()]

    q_idx = start_idx
    for block in blocks:
        lines = [l.strip() for l in block.splitlines() if l.strip()]
        if not lines:
            continue

        # Задание
        task_lines = []
        answer_lines = {}
        correct_raw = ""
        mode = "task"

        for line in lines:
            up = line.upper()
            if up.startswith("ЗАДАНИЕ:") or up.startswith("ЗАДАНИЕ :"):
                mode = "task"
                rest = re.sub(r"^ЗАДАНИЕ\s*:\s*", "", line, flags=re.IGNORECASE).strip()
                if rest:
                    task_lines.append(rest)
                continue

            # Варианты ответов А) Б) В) Г) Д)
            ans_m = re.match(r"^([АБВГДАБВГДабвгд])\s*[)\.]\s*(.+)$", line)
            if ans_m:
                mode = "answers"
                label = ans_m.group(1).upper()
                # Нормализуем к русской кириллице
                label = _normalize_label(label)
                answer_lines[label] = ans_m.group(2).strip()
                continue

            if re.match(r"^ПРАВИЛЬН", line, re.IGNORECASE):
                mode = "correct"
                rest = re.sub(r"^ПРАВИЛЬН[Ы]?[ЙЕ]?\s*[:\-]?\s*", "", line,
                               flags=re.IGNORECASE).strip()
                correct_raw = rest
                continue

            if mode == "task":
                task_lines.append(line)
            elif mode == "correct":
                correct_raw += " " + line

        task = " ".join(task_lines).strip()
        if not task or len(answer_lines) < 2:
            continue  # пропускаем невалидные блоки

        # Парсим правильные ответы
        correct_letters = [
            _normalize_label(l.strip().upper())
            for l in re.split(r"[,;и\s]+", correct_raw)
            if l.strip() and re.match(r"^[АБВГДАБВГДабвгдABCDEabcde]$", l.strip())
        ]
        correct_letters = [l for l in correct_letters if l in answer_lines]

        if not correct_letters:
            # Попытка найти букву в ответе (fallback)
            m = re.search(r"[АБВГДабвгдABCDE]", correct_raw)
            if m:
                correct_letters = [_normalize_label(m.group(0).upper())]

        if not correct_letters and answer_lines:
            correct_letters = [list(answer_lines.keys())[0]]

        n_correct = len(correct_letters)
        is_ordered = 0
        if rank == 3 and n_correct > 1:
            qtype = QTYPE_MULTIPLE
        else:
            qtype = QTYPE_SINGLE
            if n_correct != 1:
                # Приводим к одному правильному для ранга 1-2
                correct_letters = correct_letters[:1]
                n_correct = 1

        # Нумерация: КодДисц.Раздел.Тема.ТипВопроса.НомерВопроса.Ранг.Послед(КолПравильных)
        number = (f"{discipline_code}.{section_num}.{topic_num}.{qtype}."
                  f"{q_idx:04d}.{rank}.{is_ordered}({n_correct})")

        correct_texts = [answer_lines[l] for l in correct_letters if l in answer_lines]

        questions.append({
            "number":          number,
            "task":            task,
            "answers":         answer_lines,   # dict label → text
            "correct_letters": correct_letters,
            "correct_texts":   correct_texts,
            "rank":            rank,
            "qtype":           qtype,
            "n_correct":       n_correct,
            "section_num":     section_num,
            "topic_num":       topic_num,
        })
        q_idx += 1

    return questions


def _normalize_label(label: str) -> str:
    """Нормализует метку ответа к русской букве."""
    mapping = {
        "A": "А", "B": "Б", "C": "В", "D": "Г", "E": "Д",
        "а": "А", "б": "Б", "в": "В", "г": "Г", "д": "Д",
    }
    return mapping.get(label, label)


def generate_questions_for_section(
        section: dict,
        discipline: str,
        code: str,
        n_per_rank: int = MIN_PER_RANK,
        no_rag: bool = False,
) -> list[dict]:
    """
    Генерирует вопросы для одного раздела по всем 3 рангам.

    Стратегия: для каждого ранга отправляем один промпт.
    Если тем несколько — распределяем вопросы по темам.
    """
    sec_num  = section["num"]
    sec_name = section["name"]
    comps    = ", ".join(section["competencies"])
    topics   = section["topics"] or [sec_name]
    # Используем максимум 3 темы для структуры
    topics = topics[:3] if len(topics) > 3 else topics

    all_questions: list[dict] = []
    global_idx = sec_num * 1000 + 1  # Нумерация вопросов: 1001, 2001, 3001 по разделам

    print(f"\n  📝 Раздел {sec_num}: «{sec_name[:50]}»")
    print(f"     Компетенции: {comps} | Тем: {len(topics)}")

    # RAG-контекст для всего раздела (один запрос)
    ctx = retrieve_for_section(sec_name, discipline, no_rag)
    if not ctx:
        ctx = f"Содержание раздела «{sec_name}» дисциплины «{discipline}»."

    for rank in [1, 2, 3]:
        rank_questions: list[dict] = []
        n_target = n_per_rank  # минимум вопросов на ранг

        # Распределяем вопросы по темам
        n_topics = len(topics)
        n_per_topic = max(3, (n_target + n_topics - 1) // n_topics)

        for t_idx, topic in enumerate(topics, start=1):
            print(f"     Ранг {rank} | тема {t_idx}/{n_topics}: {topic[:50]}")

            prompt = _PROMPT_TEMPLATE.format(
                discipline=discipline,
                section_name=sec_name,
                topic=topic,
                competencies=comps,
                rank_prompt=RANK_PROMPTS[rank].format(n=n_per_topic),
                context=ctx[:MAX_CONTEXT_CHARS] if ctx else "(контекст недоступен)",
                n=n_per_topic,
            )

            raw = llm(prompt, max_tokens=1400)

            parsed = _parse_questions_from_llm(
                raw,
                rank=rank,
                section_num=sec_num,
                topic_num=t_idx,
                discipline_code=code,
                start_idx=global_idx,
            )

            print(f"       → распознано вопросов: {len(parsed)}")
            rank_questions.extend(parsed)
            global_idx += len(parsed)

        # Если вопросов меньше минимума — делаем ещё один запрос
        if len(rank_questions) < n_per_rank:
            shortage = n_per_rank - len(rank_questions)
            print(f"     ⚠️  Ранг {rank}: нехватка {shortage} вопросов, дозапрос...")
            prompt = _PROMPT_TEMPLATE.format(
                discipline=discipline,
                section_name=sec_name,
                topic=sec_name,
                competencies=comps,
                rank_prompt=RANK_PROMPTS[rank].format(n=shortage + 2),
                context=ctx[:MAX_CONTEXT_CHARS],
                n=shortage + 2,
            )
            raw = llm(prompt, max_tokens=1400)
            extra = _parse_questions_from_llm(
                raw, rank=rank, section_num=sec_num,
                topic_num=len(topics) + 1,
                discipline_code=code, start_idx=global_idx,
            )
            rank_questions.extend(extra[:shortage + 2])
            global_idx += len(extra)

        print(f"     ✅ Ранг {rank}: итого {len(rank_questions)} вопросов")
        all_questions.extend(rank_questions)

    print(f"  ✅ Раздел {sec_num}: всего {len(all_questions)} вопросов")
    return all_questions


# ── Запись DOCX ───────────────────────────────────────────────────────────────

def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_bold_label(doc: Document, label: str, value: str) -> None:
    """Добавляет строку вида «МЕТКА: значение» с жирной меткой."""
    p = doc.add_paragraph()
    run_label = p.add_run(label + ": ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_val = p.add_run(value)
    run_val.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)


def _add_separator(doc: Document) -> None:
    p = doc.add_paragraph("─" * 60)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(8)


def write_tests_docx(
        sections: list[dict],
        all_questions: list[dict],
        discipline: str,
        code: str,
        output_path: str,
) -> None:
    """
    Записывает тестовые задания в DOCX в формате, совместимом с методичкой.

    Структура файла:
    - Шапка (дисциплина, код)
    - Таблица разделов
    - Вопросы, сгруппированные по разделам и рангам
    """
    doc = Document()

    # Поля страницы
    section_prop = doc.sections[0]
    section_prop.left_margin   = Cm(2.5)
    section_prop.right_margin  = Cm(1.5)
    section_prop.top_margin    = Cm(2.0)
    section_prop.bottom_margin = Cm(2.0)

    # ── Шапка ──
    title = doc.add_heading("ТЕСТОВЫЕ ЗАДАНИЯ", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run(f"Дисциплина: {discipline}")
    r.bold = True
    r.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_bold_label(doc, "Шифр дисциплины", code)

    # ── Таблица разделов ──
    doc.add_paragraph()
    _add_heading(doc, "Разделы дисциплины", level=2)

    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "№ раздела"
    hdr[1].text = "Название раздела"
    hdr[2].text = "Компетенции"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for s in sections:
        row = t.add_row().cells
        row[0].text = str(s["num"])
        row[1].text = s["name"]
        row[2].text = ", ".join(s["competencies"])

    doc.add_paragraph()

    # ── Вопросы по разделам ──
    questions_by_section: dict[int, list] = {}
    for q in all_questions:
        questions_by_section.setdefault(q["section_num"], []).append(q)

    for s in sections:
        sec_num = s["num"]
        qs = questions_by_section.get(sec_num, [])
        if not qs:
            continue

        doc.add_page_break()
        _add_heading(doc, f"Раздел {sec_num}. {s['name']}", level=2)
        _add_bold_label(doc, "Компетенции", ", ".join(s["competencies"]))
        _add_bold_label(doc, "Всего вопросов", str(len(qs)))

        # Группируем по рангу
        by_rank: dict[int, list] = {}
        for q in qs:
            by_rank.setdefault(q["rank"], []).append(q)

        for rank in sorted(by_rank.keys()):
            rank_qs = by_rank[rank]
            _add_heading(doc,
                          f"Ранг {rank} — {RANKS[rank]} ({len(rank_qs)} вопросов)",
                          level=3)

            for q in rank_qs:
                # Номер задания
                p_num = doc.add_paragraph()
                r_label = p_num.add_run("Номер: ")
                r_label.bold = True
                r_label.font.size = Pt(10)
                r_code = p_num.add_run(q["number"])
                r_code.font.size = Pt(10)
                r_code.font.color.rgb = RGBColor(0x17, 0x5C, 0xC4)
                p_num.paragraph_format.space_before = Pt(8)
                p_num.paragraph_format.space_after  = Pt(0)

                # Задание
                p_task = doc.add_paragraph()
                r_task_label = p_task.add_run("Задание: ")
                r_task_label.bold = True
                r_task_label.font.size = Pt(11)
                r_task_text = p_task.add_run(q["task"])
                r_task_text.font.size = Pt(11)
                p_task.paragraph_format.space_before = Pt(2)
                p_task.paragraph_format.space_after  = Pt(4)

                # Заголовок «Ответы:»
                p_ah = doc.add_paragraph()
                r_ah = p_ah.add_run("Ответы:")
                r_ah.bold = True
                r_ah.font.size = Pt(11)
                p_ah.paragraph_format.space_before = Pt(0)
                p_ah.paragraph_format.space_after  = Pt(0)

                # Варианты ответов
                for label, text in q["answers"].items():
                    p_ans = doc.add_paragraph(style="List Bullet")
                    r_lbl = p_ans.add_run(f"{label}) ")
                    r_lbl.bold = True
                    r_lbl.font.size = Pt(11)
                    r_txt = p_ans.add_run(text)
                    r_txt.font.size = Pt(11)
                    p_ans.paragraph_format.space_before = Pt(0)
                    p_ans.paragraph_format.space_after  = Pt(0)

                # Разделитель
                doc.add_paragraph()
                p_sep = doc.add_paragraph("─" * 55)
                p_sep.paragraph_format.space_before = Pt(2)
                p_sep.paragraph_format.space_after  = Pt(0)
                for run in p_sep.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

                # Правильный ответ
                p_correct = doc.add_paragraph()
                r_cl = p_correct.add_run("Правильный ответ: ")
                r_cl.bold = True
                r_cl.font.size = Pt(11)
                correct_display = " | ".join(
                    f"{l}) {q['answers'].get(l, '')}"
                    for l in q["correct_letters"]
                )
                r_cv = p_correct.add_run(correct_display)
                r_cv.font.size = Pt(11)
                p_correct.paragraph_format.space_before = Pt(0)
                p_correct.paragraph_format.space_after  = Pt(10)

    doc.save(output_path)
    print(f"\n✅ Тесты сохранены: {output_path}")


# ── Отчёт о покрытии компетенций ──────────────────────────────────────────────

def build_coverage_report(
        sections: list[dict],
        all_questions: list[dict],
        cfg: dict,
) -> dict:
    """
    Строит отчёт о покрытии компетенций по требованиям методички:
    - Минимум 30 вопросов на раздел
    - Минимум 100 вопросов на компетенцию
    """
    questions_by_section: dict[int, list] = {}
    for q in all_questions:
        questions_by_section.setdefault(q["section_num"], []).append(q)

    # Подсчёт по компетенциям
    comp_counts: dict[str, int] = {}
    for s in sections:
        sec_num = s["num"]
        n_qs = len(questions_by_section.get(sec_num, []))
        for comp in s["competencies"]:
            comp_counts[comp] = comp_counts.get(comp, 0) + n_qs

    report = {
        "discipline":      cfg.get("discipline", ""),
        "code":            cfg.get("code", ""),
        "generated_at":    time.strftime("%Y-%m-%dT%H:%M:%S"),
        "total_questions": len(all_questions),
        "sections": [],
        "competencies": [],
        "requirements": {
            "min_per_section": MIN_PER_SECTION,
            "min_per_comp":    MIN_PER_COMP,
        },
        "status": "OK",
    }

    warnings = []

    # По разделам
    for s in sections:
        sec_num = s["num"]
        qs = questions_by_section.get(sec_num, [])
        by_rank = {}
        for q in qs:
            by_rank[q["rank"]] = by_rank.get(q["rank"], 0) + 1
        ok = len(qs) >= MIN_PER_SECTION

        entry = {
            "section":      sec_num,
            "name":         s["name"],
            "competencies": s["competencies"],
            "total":        len(qs),
            "by_rank":      by_rank,
            "ok":           ok,
        }
        if not ok:
            warnings.append(
                f"Раздел {sec_num}: {len(qs)} вопросов (требуется ≥{MIN_PER_SECTION})"
            )
        report["sections"].append(entry)

    # По компетенциям
    for comp, count in sorted(comp_counts.items()):
        ok = count >= MIN_PER_COMP
        entry = {
            "code":  comp,
            "desc":  cfg.get("fgos_competencies", {}).get(comp, ""),
            "total": count,
            "ok":    ok,
        }
        if not ok:
            warnings.append(
                f"Компетенция {comp}: {count} вопросов (требуется ≥{MIN_PER_COMP})"
            )
        report["competencies"].append(entry)

    if warnings:
        report["status"] = "WARNINGS"
        report["warnings"] = warnings

    return report


def print_coverage_summary(report: dict) -> None:
    print("\n" + "=" * 60)
    print("📊 ПОКРЫТИЕ КОМПЕТЕНЦИЙ")
    print("=" * 60)
    print(f"Всего вопросов: {report['total_questions']}\n")

    print("Разделы:")
    for s in report["sections"]:
        status = "✅" if s["ok"] else "❌"
        ranks_str = " | ".join(f"R{r}:{n}" for r, n in sorted(s["by_rank"].items()))
        print(f"  {status} Раздел {s['section']}: {s['total']} вопросов [{ranks_str}]")

    print("\nКомпетенции:")
    for c in report["competencies"]:
        status = "✅" if c["ok"] else "❌"
        print(f"  {status} {c['code']}: {c['total']} вопросов")

    if report.get("warnings"):
        print("\n⚠️  Предупреждения:")
        for w in report["warnings"]:
            print(f"   • {w}")
        print("\n  → Запустите повторно с --questions-per-rank 20 для дозаполнения")
    else:
        print("\n✅ Все требования выполнены!")
    print("=" * 60)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Генерация тестовых заданий (ФОС) по РПД"
    )
    parser.add_argument(
        "--section", type=int, default=None,
        help="Генерировать только указанный раздел (номер)"
    )
    parser.add_argument(
        "--questions-per-rank", type=int, default=MIN_PER_RANK,
        help=f"Вопросов на ранг на раздел (по умолчанию {MIN_PER_RANK})"
    )
    parser.add_argument(
        "--no-rag", action="store_true",
        help="Отключить Qdrant (генерация без контекста из учебников)"
    )
    parser.add_argument(
        "--rpd", type=str, default=str(RPD_PATH),
        help="Путь к output_rpd.docx"
    )
    parser.add_argument(
        "--output", type=str, default=OUTPUT_DOCX,
        help="Путь к выходному DOCX"
    )
    parser.add_argument(
        "--config", type=str, default=str(CONFIG_PATH),
        help="Путь к config.json"
    )
    args = parser.parse_args()

    # Загрузка конфига
    cfg_path = Path(args.config)
    if not cfg_path.exists():
        print(f"❌ config.json не найден: {cfg_path}")
        sys.exit(1)

    cfg = json.loads(cfg_path.read_text(encoding="utf-8"))
    discipline = cfg["discipline"]
    code       = cfg["code"]

    print(f"🎓 Дисциплина: {discipline} (код: {code})")
    print(f"   Компетенции: {cfg.get('competency_codes', '')}")

    # Кэш
    _load_cache()

    # Парсинг РПД
    rpd_path = Path(args.rpd)
    if not rpd_path.exists():
        print(f"❌ РПД не найден: {rpd_path}")
        sys.exit(1)

    sections = parse_rpd_sections(rpd_path)
    if not sections:
        print("❌ Не удалось извлечь разделы из РПД")
        sys.exit(1)

    # Фильтрация по --section
    if args.section is not None:
        sections = [s for s in sections if s["num"] == args.section]
        if not sections:
            print(f"❌ Раздел {args.section} не найден в РПД")
            sys.exit(1)
        print(f"🎯 Режим: только раздел {args.section}")

    # Генерация вопросов
    print(f"\n🚀 Начало генерации: {args.questions_per_rank} вопросов/ранг × 3 ранга × "
          f"{len(sections)} разделов = ~{args.questions_per_rank * 3 * len(sections)} вопросов")

    all_questions: list[dict] = []

    for section in sections:
        qs = generate_questions_for_section(
            section=section,
            discipline=discipline,
            code=code,
            n_per_rank=args.questions_per_rank,
            no_rag=args.no_rag,
        )
        all_questions.extend(qs)
        _save_cache()

    print(f"\n📦 Итого сгенерировано вопросов: {len(all_questions)}")

    # Запись DOCX
    write_tests_docx(
        sections=sections,
        all_questions=all_questions,
        discipline=discipline,
        code=code,
        output_path=args.output,
    )

    # Отчёт о покрытии
    report = build_coverage_report(sections, all_questions, cfg)
    Path(COVERAGE_LOG).write_text(
        json.dumps(report, ensure_ascii=False, indent=2),
        encoding="utf-8"
    )
    print(f"📋 Отчёт покрытия: {COVERAGE_LOG}")

    print_coverage_summary(report)
    _save_cache()


if __name__ == "__main__":
    main()
