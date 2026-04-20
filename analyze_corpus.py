"""
analyze_corpus.py — анализ существующего корпуса RPD и рекомендации по расширению.

Что делает:
  1. Читает все .docx в rpd_corpus/ → извлекает название дисциплины, коды компетенций,
     семестр, трудоёмкость.
  2. Кластеризует дисциплины по тематическим доменам (ключевые слова).
  3. Показывает распределение корпуса по доменам.
  4. Выявляет слабо покрытые домены (< 2 РПД).
  5. Выдаёт конкретный список рекомендуемых дисциплин для синтетических РПД.
  6. Сохраняет corpus_analysis.json для последующего использования.

Запуск:
  python analyze_corpus.py
  python analyze_corpus.py --corpus-dir path/to/rpd_corpus
  python analyze_corpus.py --jsonl data_clean.jsonl   # режим jsonl (быстрее)
"""

import argparse
import json
import os
import re
import sys
from collections import defaultdict

# ---------------------------------------------------------------------------
# Тематические домены — ключевые слова для классификации
# ---------------------------------------------------------------------------
DOMAINS: dict[str, list[str]] = {
    # ==========================================================================
    # ВАЖНО: порядок доменов определяет приоритет классификации.
    # Специфичные домены — выше; широкие («системн», «управлен») — ниже.
    # [FIX-#22]

    # --- 1. Искусственный интеллект (самый специфичный — первым) ---
    "Искусственный интеллект / МО": [
        # Ядро ML/DL
        "машинн", "нейрон", "глубок обучен", "deep learning",
        "классификаци", "регрессия", "кластериз", "распознавани",
        # Компьютерное зрение / NLP
        "зрени", "nlp", "обработка текст", "обработка естествен",
        # Символьный / гибридный ИИ — [FIX-#22] новые ключевые слова
        "интеллект", "онтологи", "представлен знани", "нейросимвол",
        "гибридн", "агентн", "мультиагентн", "экспертн систем",
        "логическ программирован", "нечётк", "эволюцион", "метаэвристик",
        "поддержк принят решени", "компьютерн моделирован интеллект",
        "глубок",
    ],

    # --- 2. Специфичные технологические домены ---
    "Веб-разработка": [
        "веб", "web", "html", "css", "javascript", "react", "frontend",
        "backend", "api", "rest", "http", "браузер",
    ],
    "Компьютерная графика / Игры": [
        "компьютерн график", "opengl", "directx", "игр", "3d",
        "рендеринг", "анимац", "геймдев",
        # «визуализац» убрана отсюда — перехватывала «анализ данных»
    ],
    "Облачные технологии / DevOps": [
        "облачн", "cloud", "docker", "kubernetes", "контейнер",
        "ci/cd", "развёртыван", "инфраструктур",
    ],
    "Встроенные / IoT системы": [
        "встроенн", "iot", "интернет вещей", "arduino", "raspberry",
        "fpga", "плис", "микропроцессор", "цифров обработк сигнал",
    ],
    "Анализ данных / BI": [
        "анализ данн", "data mining", "интеллектуальн анализ данн",
        "аналитик", "bi", "визуализац данн", "визуализац",
        "статистик", "прогнозирован", "pandas", "tableau", "power bi",
        "большие данн", "big data",
    ],
    "Кибербезопасность": [
        "безопасност", "криптограф", "защит", "угроз", "уязвимост",
        "аутентификац", "шифрован", "атак", "penetration",
    ],
    "Сети / Телекоммуникации": [
        "протокол", "tcp", "ip", "маршрутиз", "коммутац",
        "телекоммуник", "беспровод", "wi-fi", "vpn",
        "сетев", "компьютерные сет",
    ],

    # --- 3. Общие домены (широкие ключевые слова — проверяются последними) ---
    "Базы данных / Хранилища": [
        "баз данн", "sql", "субд", "хранилищ", "запрос",
        "реляцион", "nosql", "mongodb", "администрирован",
        # «баз» и «данн» убраны отдельно — слишком широкие
    ],
    "Математика / Алгоритмы": [
        "математик", "алгоритм", "теория", "граф", "дискретн",
        "вычислительн", "числен", "оптимизац", "параллельн вычислен",
        "распределённ систем",
    ],
    "Системное / Низкоуровневое ПО": [
        "операционные систем", "ядро", "компилятор", "ассемблер",
        "микроконтроллер", "системное программ",
        # «систем», «процесс», «поток», «памят» убраны — слишком широкие
    ],
    "Программная инженерия": [
        "программн инженери", "software", "методолог", "тестирован",
        "жизненн цикл", "agile", "архитектур программ", "паттерн",
    ],
    "Управление проектами / Экономика ИТ": [
        "управлен ит-проект", "управлен проект", "it-проект", "менеджмент", "экономик",
        "стоимост", "бюджет", "риск", "планирован", "организац",
    ],
    "Научно-исследовательская деятельность": [
        "исследовательск", "научн деятельност", "квалификацион",
        "публикац", "методологи научн",
    ],
    "Философия науки / Педагогика": [
        "философи", "педагогик", "история науки", "дидактик",
        "преподаван", "академическ письм", "научн коммуникац",
    ],
    "Продвинутые методы ИИ": [
        "трансформер", "diffusion", "reinforcement", "генеративн",
        "self-supervised", "federated", "llm", "foundation model",
        "мультимодальн", "объяснимый ии", "xai",
    ],
}

# Минимальное число РПД в домене, при котором домен считается «достаточно покрытым».
# Значение по умолчанию — 3 (было 2, но порог 2 скрывал реальный дисбаланс корпуса).
# Можно переопределить через --min-coverage.
MIN_COVERAGE = 3

# ---------------------------------------------------------------------------
# Рекомендуемые дисциплины для слабо покрытых доменов
# ---------------------------------------------------------------------------
RECOMMENDATIONS: dict[str, list[dict]] = {
    "Веб-разработка": [
        {"name": "Веб-программирование", "focus": "HTML5, CSS3, JavaScript ES6+, REST API, Node.js"},
        {"name": "Разработка веб-приложений", "focus": "React, Vue.js, SPA, webpack, TypeScript"},
        {"name": "Серверное программирование", "focus": "Python Flask/Django, JWT, PostgreSQL, REST"},
    ],
    "Кибербезопасность": [
        {"name": "Информационная безопасность", "focus": "криптография, PKI, протоколы TLS, угрозы OWASP Top 10"},
        {"name": "Защита программного обеспечения", "focus": "анализ уязвимостей, SAST/DAST, безопасная разработка"},
        {"name": "Сетевая безопасность", "focus": "межсетевые экраны, IDS/IPS, VPN, анализ трафика"},
    ],
    "Облачные технологии / DevOps": [
        {"name": "Облачные вычисления", "focus": "AWS/Azure/GCP, IaaS/PaaS/SaaS, виртуализация, Terraform"},
        {"name": "Технологии контейнеризации", "focus": "Docker, Kubernetes, CI/CD, GitLab, Jenkins"},
        {"name": "DevOps-практики", "focus": "Git flow, автоматизация, мониторинг, Prometheus, Grafana"},
    ],
    "Компьютерная графика / Игры": [
        {"name": "Компьютерная графика", "focus": "OpenGL, GLSL, растеризация, трассировка лучей, Three.js"},
        {"name": "Разработка игр", "focus": "Unity3D, C#, игровые паттерны, физический движок, UI"},
    ],
    "Встроенные / IoT системы": [
        {"name": "Программирование микроконтроллеров", "focus": "STM32, C/C++, RTOS, SPI/I2C/UART, прерывания"},
        {"name": "Интернет вещей (IoT)", "focus": "MQTT, Zigbee, LoRa, облачные платформы, Edge Computing"},
    ],
    "Анализ данных / BI": [
        {"name": "Анализ и визуализация данных", "focus": "Python pandas, matplotlib, seaborn, дашборды, Power BI"},
        {"name": "Большие данные (Big Data)", "focus": "Hadoop, Spark, HDFS, MapReduce, потоковая обработка"},
    ],
    "Системное / Низкоуровневое ПО": [
        {"name": "Операционные системы", "focus": "процессы, потоки, синхронизация, файловые системы, Linux API"},
        {"name": "Системное программирование", "focus": "C/C++, POSIX, межпроцессное взаимодействие, память"},
    ],
    "Математика / Алгоритмы": [
        {"name": "Алгоритмы и структуры данных", "focus": "сортировка, деревья, графы, динамическое программирование"},
        {"name": "Дискретная математика", "focus": "булева алгебра, графы, комбинаторика, логика предикатов"},
    ],
    "Базы данных / Хранилища": [
        {"name": "Проектирование баз данных", "focus": "ER-диаграммы, нормализация, индексы, транзакции ACID"},
        {"name": "NoSQL базы данных", "focus": "MongoDB, Redis, Cassandra, документно-ориентированные СУБД"},
    ],
    "Сети / Телекоммуникации": [
        {"name": "Компьютерные сети", "focus": "модель OSI, TCP/IP, маршрутизация, VLAN, Cisco IOS"},
        {"name": "Протоколы передачи данных", "focus": "HTTP/2, WebSocket, gRPC, AMQP, DNS, DHCP"},
    ],
}


# ---------------------------------------------------------------------------
# Требования к покрытию по направлениям / уровням подготовки
# ---------------------------------------------------------------------------
DIRECTION_REQUIREMENTS: dict[str, dict] = {
    "09.04.01": {
        "label":    "Магистратура — 09.04.01 ИВТ",
        "min_rpd":  2,
        "required_domains": [
            "Искусственный интеллект / МО",
            "Продвинутые методы ИИ",
            "Математика / Алгоритмы",
            "Облачные технологии / DevOps",
            "Базы данных / Хранилища",
            "Управление проектами / Экономика ИТ",
            "Научно-исследовательская деятельность",
        ],
        "recommendations": {
            "Продвинутые методы ИИ": [
                {"name": "Современные архитектуры глубокого обучения",
                 "focus": "Transformer, Diffusion models, RL, LLM fine-tuning, PEFT"},
                {"name": "Объяснимый искусственный интеллект",
                 "focus": "SHAP, LIME, attention visualization, fairness metrics"},
            ],
            "Научно-исследовательская деятельность": [
                {"name": "Методология научных исследований в ИТ",
                 "focus": "постановка гипотезы, эксперимент, метрики, публикации IEEE/ACM"},
                {"name": "Научно-исследовательская работа (магистратура)",
                 "focus": "НИР 1–2 семестр, подготовка к защите ВКР"},
            ],
            "Облачные технологии / DevOps": [
                {"name": "Высокопроизводительные вычисления и GPU",
                 "focus": "CUDA, OpenCL, распределённое обучение, Horovod, Ray"},
            ],
        },
    },
    "1.2.2": {
        "label":    "Аспирантура — 1.2.2 Компьютерные науки и информатика",
        "min_rpd":  1,
        "required_domains": [
            "Философия науки / Педагогика",
            "Научно-исследовательская деятельность",
            "Математика / Алгоритмы",
            "Продвинутые методы ИИ",
        ],
        "recommendations": {
            "Философия науки / Педагогика": [
                {"name": "Философия науки и методология исследований",
                 "focus": "логика научного познания, Поппер, Кун, фальсификационизм"},
                {"name": "Педагогика высшей школы",
                 "focus": "дидактика, проектирование учебных курсов, ФОС"},
                {"name": "Иностранный язык (академическое письмо)",
                 "focus": "написание статей IEEE, подготовка докладов, IELTS/TOEFL"},
            ],
            "Научно-исследовательская деятельность": [
                {"name": "Научно-исследовательская работа (аспирантура)",
                 "focus": "НИР 1–4 семестр, подготовка кандидатской диссертации"},
            ],
            "Математика / Алгоритмы": [
                {"name": "Специальные разделы математики для ИИ",
                 "focus": "топология, теория меры, стохастические процессы, оптимизация"},
            ],
            "Продвинутые методы ИИ": [
                {"name": "Современные проблемы информатики",
                 "focus": "актуальные направления CS: квантовые вычисления, neuromorphic"},
            ],
        },
    },
}


# ---------------------------------------------------------------------------
# Извлечение данных из DOCX
# ---------------------------------------------------------------------------

_CODE_RE    = re.compile(r"^\s*\(([А-ЯA-Z0-9Б1-9\.]{2,12})\)\s*(.+)", re.IGNORECASE)
_DIR_RE     = re.compile(r"\b(09\.\d{2}\.\d{2}|1\.\d+\.\d+)\b")
_COMP_RE    = re.compile(r"\b(УК-\d+|ОПК-\d+|ПК-\d+)\b")
_SEM_RE     = re.compile(r"\b([1-9]|10)\s*(?:семестр|сем\.)", re.IGNORECASE)
_HOURS_RE   = re.compile(r"\b(\d{2,3})\s*час", re.IGNORECASE)
_CREDITS_RE = re.compile(r"\b([1-9])\s*з\.е\.", re.IGNORECASE)


def _extract_from_docx(path: str) -> dict:
    """Извлекает метаданные из .docx файла."""
    try:
        from docx import Document
        doc = Document(path)
    except Exception as e:
        return {"error": str(e)}

    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # Добавляем текст из таблиц — код направления обычно там
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text += "\n" + cell.text

    # Название дисциплины
    name = ""
    for line in full_text.split("\n"):
        m = _CODE_RE.match(line.strip())
        if m:
            candidate = m.group(2).strip()
            # Фильтруем вузовскую шапку
            if (len(candidate) > 4 and
                    not re.search(r"кафедр|университет|втик|уфа", candidate, re.I)):
                name = candidate
                break

    # Коды компетенций
    comp_codes = sorted(set(_COMP_RE.findall(full_text)))

    # Семестр
    sem_matches = _SEM_RE.findall(full_text)
    semester = sem_matches[0] if sem_matches else ""

    # Часы и з.е.
    hours_matches   = _HOURS_RE.findall(full_text)
    credits_matches = _CREDITS_RE.findall(full_text)
    hours   = int(hours_matches[0])   if hours_matches   else 0
    credits = int(credits_matches[0]) if credits_matches else 0

    # Направление подготовки
    dir_matches = _DIR_RE.findall(full_text)
    direction = dir_matches[0] if dir_matches else "unknown"

    return {
        "name":       name,
        "file":       os.path.basename(path),
        "comp_codes": comp_codes,
        "semester":   semester,
        "hours":      hours,
        "credits":    credits,
        "direction":  direction,
    }


def _extract_from_jsonl(jsonl_path: str) -> list[dict]:
    """Быстрый режим: читает data_clean.jsonl вместо DOCX."""
    raw: dict = {}
    with open(jsonl_path, encoding="utf-8") as f:
        for line in f:
            try:
                r = json.loads(line)
            except Exception:
                continue
            src = r.get("source", "")
            if not src or src in raw:
                continue
            text = r.get("text", "")
            m = _CODE_RE.match(text.strip().split("\n")[0])
            name = m.group(2).strip() if m else ""
            if not name:
                dm = r.get("document_meta") or {}
                name = dm.get("discipline", "") or dm.get("title", "")
            comp_codes = sorted(set(_COMP_RE.findall(text)))
            dir_matches = _DIR_RE.findall(text)
            raw[src] = {
                "name":       name,
                "file":       src,
                "comp_codes": comp_codes,
                "semester":   "",
                "hours":      0,
                "credits":    0,
                "direction":  dir_matches[0] if dir_matches else "unknown",
            }
    return list(raw.values())


# ---------------------------------------------------------------------------
# Классификация по домену
# ---------------------------------------------------------------------------

def classify_domain(name: str) -> str:
    """Возвращает название домена или 'Прочее'."""
    name_lower = name.lower()
    for domain, keywords in DOMAINS.items():
        if any(kw in name_lower for kw in keywords):
            return domain
    return "Прочее"


# ---------------------------------------------------------------------------
# Основная логика
# ---------------------------------------------------------------------------

def analyze(items: list[dict]) -> dict:
    by_domain: dict[str, list[dict]] = defaultdict(list)
    for item in items:
        if not item.get("name"):
            continue
        domain = classify_domain(item["name"])
        by_domain[domain].append(item)

    return dict(by_domain)


def print_report(by_domain: dict[str, list[dict]], items: list[dict]) -> None:
    total = len(items)
    print(f"\n{'='*65}")
    print(f"  АНАЛИЗ КОРПУСА РПД  ({total} документов)")
    print(f"{'='*65}")
    print(f"\n{'Домен':<42} {'РПД':>4}  Дисциплины")
    print("-"*65)

    # Сортируем по убыванию числа РПД
    for domain, docs in sorted(by_domain.items(), key=lambda x: -len(x[1])):
        names = ", ".join(d["name"] for d in docs[:3])
        if len(docs) > 3:
            names += f" ... (+{len(docs)-3})"
        flag = "✅" if len(docs) >= MIN_COVERAGE else "❌"
        print(f"  {flag} {domain:<40} {len(docs):>3}  {names}")

    # Непокрытые домены из рекомендаций
    covered = set(by_domain.keys())
    missing_domains = [d for d in RECOMMENDATIONS if d not in covered
                       or len(by_domain.get(d, [])) < MIN_COVERAGE]

    print(f"\n{'='*65}")
    print(f"  СЛАБО ПОКРЫТЫЕ ДОМЕНЫ (< {MIN_COVERAGE} РПД):")
    print(f"{'='*65}")
    if not missing_domains:
        print("  Все ключевые домены покрыты.")
    else:
        for d in missing_domains:
            have = len(by_domain.get(d, []))
            print(f"  ⚠️  {d}  (текущих: {have})")

    print(f"\n{'='*65}")
    print(f"  РЕКОМЕНДУЕМЫЕ ДИСЦИПЛИНЫ ДЛЯ СИНТЕТИЧЕСКИХ РПД:")
    print(f"{'='*65}")
    total_recs = 0
    for domain in missing_domains:
        recs = RECOMMENDATIONS.get(domain, [])
        if not recs:
            continue
        have = len(by_domain.get(domain, []))
        need = max(0, MIN_COVERAGE - have)
        print(f"\n  [{domain}]")
        for i, r in enumerate(recs[:need + 1]):
            marker = "★" if i < need else " "
            print(f"   {marker} {r['name']}")
            print(f"       focus: {r['focus']}")
            total_recs += 1
    if not total_recs:
        print("  Нет рекомендаций — корпус сбалансирован.")

    print(f"\n{'='*65}")
    print(f"  СВОДКА: нужно добавить ~{total_recs} синтетических РПД")
    print(f"  для достижения минимального покрытия по всем доменам.")
    print(f"{'='*65}\n")


def print_direction_report(by_domain: dict[str, list[dict]], items: list[dict]) -> None:
    """Выводит покрытие корпуса по уровням подготовки (магистратура, аспирантура)."""
    from collections import Counter

    dir_counts = Counter(i.get("direction", "unknown") for i in items)

    print(f"\n{'='*65}")
    print(f"  ПОКРЫТИЕ ПО НАПРАВЛЕНИЯМ ПОДГОТОВКИ")
    print(f"{'='*65}")

    for dir_code, req in DIRECTION_REQUIREMENTS.items():
        have_total = dir_counts.get(dir_code, 0)
        label = req["label"]
        print(f"\n  📋 {label}  — в корпусе: {have_total} РПД")
        print(f"  {'Домен':<42} {'Нужно':>6}  {'Есть':>5}  Статус")
        print(f"  {'-'*60}")

        missing_recs = []
        for domain in req["required_domains"]:
            have = len(by_domain.get(domain, []))
            need = req["min_rpd"]
            ok = have >= need
            flag = "✅" if ok else "❌"
            print(f"  {flag} {domain:<42} {need:>5}  {have:>5}")
            if not ok:
                missing_recs.append(domain)

        if missing_recs:
            print(f"\n  ⚠️  Рекомендуется добавить РПД в {len(missing_recs)} доменах:")
            for domain in missing_recs:
                recs = req["recommendations"].get(domain, [])
                need = req["min_rpd"] - len(by_domain.get(domain, []))
                for i, r in enumerate(recs[:need + 1]):
                    marker = "★" if i < need else " "
                    print(f"     {marker} [{domain}] {r['name']}")
                    print(f"         focus: {r['focus']}")
        else:
            print(f"\n  ✅ Все обязательные домены покрыты.")

    print(f"\n{'='*65}\n")


def save_json(by_domain: dict, output_path: str, items: list[dict]) -> None:
    report = {
        "total_documents": len(items),
        "domains": {
            domain: {
                "count": len(docs),
                "covered": len(docs) >= MIN_COVERAGE,
                "disciplines": [d["name"] for d in docs],
            }
            for domain, docs in sorted(by_domain.items(), key=lambda x: -len(x[1]))
        },
        "recommendations": {
            domain: [
                {"name": r["name"], "focus": r["focus"]}
                for r in recs
            ]
            for domain, recs in RECOMMENDATIONS.items()
            if domain not in by_domain or len(by_domain.get(domain, [])) < MIN_COVERAGE
        },
        "all_disciplines": [
            {"file": i["file"], "name": i["name"], "domain": classify_domain(i["name"]),
             "direction": i.get("direction", "unknown")}
            for i in items if i.get("name")
        ],
    }
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    print(f"✅ Отчёт сохранён: {output_path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description="Анализ корпуса РПД")
    parser.add_argument(
        "--corpus-dir", default="rpd_corpus",
        help="Папка с .docx файлами корпуса (по умолчанию: rpd_corpus/)",
    )
    parser.add_argument(
        "--jsonl", default=None,
        help="Путь к data_clean.jsonl (быстрый режим, без чтения DOCX)",
    )
    parser.add_argument(
        "--out", default="corpus_analysis.json",
        help="Путь для сохранения JSON-отчёта (по умолчанию: corpus_analysis.json)",
    )
    parser.add_argument(
        "--min-coverage", type=int, default=None,
        help="Минимум РПД на домен для «покрыт» (по умолчанию: 3)",
    )
    # [FIX-§9]
    parser.add_argument(
        "--json", action="store_true",
        help="Вывести рекомендации по пробелам как JSON на stdout (для автоматизации)",
    )
    args = parser.parse_args()

    # Переопределяем глобальный порог, если задан явно
    global MIN_COVERAGE
    if args.min_coverage is not None:
        MIN_COVERAGE = args.min_coverage

    # --- Загрузка данных ---
    if args.jsonl and os.path.exists(args.jsonl):
        print(f"📦 Режим: data_clean.jsonl  ({args.jsonl})")
        items = _extract_from_jsonl(args.jsonl)
    elif os.path.isdir(args.corpus_dir):
        docx_files = sorted(
            os.path.join(args.corpus_dir, f)
            for f in os.listdir(args.corpus_dir)
            if f.endswith(".docx") and not f.startswith("~$")
               and not f.startswith("Шаблон")
        )
        if not docx_files:
            print(f"❌ .docx файлы не найдены в {args.corpus_dir!r}")
            sys.exit(1)
        print(f"📦 Режим: DOCX  ({len(docx_files)} файлов из {args.corpus_dir}/)")
        items = []
        for path in docx_files:
            result = _extract_from_docx(path)
            if "error" not in result:
                items.append(result)
            else:
                print(f"  ⚠️  {os.path.basename(path)}: {result['error']}")
    else:
        print(f"❌ Ни jsonl, ни corpus_dir не найдены. "
              f"Укажи --corpus-dir или --jsonl")
        sys.exit(1)

    # Фильтруем записи без названия
    named = [i for i in items if i.get("name")]
    unnamed = len(items) - len(named)
    if unnamed:
        print(f"  ⚠️  {unnamed} документов без распознанного названия пропущены")

    print(f"  ✅ Распознано: {len(named)} дисциплин\n")

    # --- Анализ ---
    by_domain = analyze(named)

    # --- Вывод ---
    print_report(by_domain, named)
    print_direction_report(by_domain, named)

    # [FIX-§9]
    if args.json:
        import json as _json
        gaps = {
            domain: [{"name": r["name"], "focus": r["focus"]} for r in recs]
            for domain, recs in RECOMMENDATIONS.items()
            if domain not in by_domain or len(by_domain.get(domain, [])) < MIN_COVERAGE
        }
        print(_json.dumps(gaps, ensure_ascii=False, indent=2))

    # --- Сохранение ---
    save_json(by_domain, args.out, named)


if __name__ == "__main__":
    main()
