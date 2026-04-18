"""
evaluate.py — оценка качества сгенерированного РПД (BLEU + ROUGE).

Алгоритм:
  1. Читает output_rpd.docx, извлекает текст по секциям.
  2. Загружает корпус из data_clean.jsonl.
  3. Ищет ближайший эталон в корпусе по cosine similarity
     эмбеддинга названия дисциплины (Ollama / bge-m3).
  4. Считает BLEU-4 + ROUGE-1/2/L на уровне документа и по секциям.
  5. Сохраняет eval_report.json + печатает таблицу.

Запуск:
  python evaluate.py                              # defaults
  python evaluate.py output_rpd.docx --config config.json
  python evaluate.py --reference rpd_5.json       # принудительный эталон
  python evaluate.py --ref-docx-dir rpd_docx/     # сравнивать docx-с-docx
"""

import argparse
import json
import os
import re
import sys
import time
from typing import Optional

import numpy as np
import requests
from docx import Document                        # [FIX-PEP8] перенесён из середины файла
from nltk.stem.snowball import SnowballStemmer   # [FIX-ROUGE] для стемминга русских токенов
# [FIX-#18] Единая embed-функция из utils.py
from utils import get_embedding as _embed_raw

# [FIX-#9] Кэш эмбеддингов: evaluate.py вычислял 49×N эмбеддингов при каждом
# запуске. Простой dict-кэш + персистенция в eval_cache.json — аналогично
# rpd_cache.json в rpd_generate.py.
_EMBED_CACHE: dict = {}
_EVAL_CACHE_FILE = "eval_cache.json"


def _load_eval_cache() -> None:
    global _EMBED_CACHE
    from pathlib import Path as _Path
    p = _Path(_EVAL_CACHE_FILE)
    if p.exists():
        try:
            _EMBED_CACHE = json.loads(p.read_text(encoding="utf-8"))
            print(f"📦 eval_cache: {len(_EMBED_CACHE)} эмбеддингов загружено")
        except Exception as e:
            print(f"⚠️  Ошибка загрузки eval_cache: {e}")


def _save_eval_cache() -> None:
    from pathlib import Path as _Path
    try:
        _Path(_EVAL_CACHE_FILE).write_text(
            json.dumps(_EMBED_CACHE, ensure_ascii=False), encoding="utf-8"
        )
    except Exception as e:
        print(f"⚠️  Ошибка сохранения eval_cache: {e}")


from nltk.translate.bleu_score import SmoothingFunction, corpus_bleu
from rouge_score import rouge_scorer as rs_lib

# [FIX-ROUGE] Стеммер для предобработки русских токенов перед ROUGE-подсчётом.
# use_stemmer=True в RougeScorer использует PorterStemmer (английский) — неприменим
# для русского. Предобработка через SnowballStemmer('russian') снижает занижение
# ROUGE на 10–20% (нейронных ≠ нейронные без стемминга).
_ru_stemmer = SnowballStemmer("russian")


def _stem_ru(text: str) -> str:
    """Стемминг русского текста для передачи в RougeScorer."""
    tokens = re.sub(r"[^\w\s]", " ", text.lower()).split()
    return " ".join(_ru_stemmer.stem(t) for t in tokens)


# ---------------------------------------------------------------------------
# Локальный Ollama — тот же стек что в rpd_generate.py
# ---------------------------------------------------------------------------
OLLAMA = {
    "embed_url":   "http://localhost:11434/api/embed",
    "embed_model": "bge-m3",
}

DEFAULT_DOCX   = "output_rpd.docx"
DEFAULT_JSONL  = "data_clean.jsonl"
DEFAULT_CONFIG = "config.json"
DEFAULT_OUT    = "eval_report.json"

# ---------------------------------------------------------------------------
# Секции: предикаты совпадают с _TABLE_PREDICATES в rpd_generate.py.
# Принимают frozenset уникальных текстов из первых 3 строк таблицы.
# ---------------------------------------------------------------------------
_SECTION_PREDICATES: dict = {
    "competencies": lambda h: any("Формируемые компетенции" in s for s in h),
    "outcomes":     lambda h: any("Индикаторы достижения компетенций" in s for s in h),
    # [FIX-T8] Добавлен признак таблицы лекций (T8): в заголовке ячейки стоит
    # «Лекция N. <тема>» вместо «Название темы (раздела)». Без этого T8 не
    # захватывалась → content = 464 симв. вместо ~1500+ → ROUGE-1 занижен.
    "content":      lambda h: (
        any("Название темы (раздела)" in s for s in h) or
        any(re.search(r"Лекция\s+\d+\.", s) for s in h)
    ),
    "lab_works":    lambda h: any("Название лабораторной работы" in s for s in h),
    "practice":     lambda h: any("Тема практического занятия" in s for s in h),
    "bibliography": lambda h: (
        any("Тип" in s for s in h) and
        any("Библиографическое описание" in s for s in h)
    ),
}

# Маппинг section_type из data_clean.jsonl → ключи _SECTION_PREDICATES.
# [FIX-MAP] assessment → None: после исправления баг-А в converter.py темы ЛР
# хранятся с section_type="content", а не "assessment". В "assessment" остался
# только ФОС-бойлерплейт (критерии оценки, шкалы) — использовать его как эталон
# для lab_works методологически некорректно: ROUGE сравнивал бы темы лабораторных
# работ с текстами типа «не зачтено / зачтено», давая заниженный score.
# Эталон для lab_works и practice теперь берётся из content-чанков (корректно).
_ST_MAP: dict = {
    "competencies":      "competencies",
    "learning_outcomes": "outcomes",
    "content":           "content",
    "lecture_content":   "content",      # [З-01] подтип лекционного контента
    "lab_content":       "lab_works",    # [З-01] подтип лабораторных работ
    "practice_content":  "practice",     # [З-01] подтип практических занятий
    "assessment":        None,            # [FIX-MAP] ФОС → не является эталоном
    "bibliography":      "bibliography",
    "place":             "bibliography",
}


# ---------------------------------------------------------------------------
# Embedding + cosine similarity (локальный Ollama / bge-m3)
# ---------------------------------------------------------------------------
def embed(text: str) -> list[float]:
    """Получает эмбеддинг через utils.get_embedding с локальным кэшем. [FIX-#9, #18]"""
    if text in _EMBED_CACHE:
        return _EMBED_CACHE[text]
    vec = _embed_raw(text, prefix="query", retry=3)
    if not vec:
        raise RuntimeError(f"embed failed for text[:50]={text[:50]!r}")
    _EMBED_CACHE[text] = vec
    return vec


def cosine_sim(a: list[float], b: list[float]) -> float:
    va, vb = np.array(a, dtype=float), np.array(b, dtype=float)
    denom = np.linalg.norm(va) * np.linalg.norm(vb)
    return float(np.dot(va, vb) / denom) if denom > 1e-10 else 0.0


# ---------------------------------------------------------------------------
# Извлечение текста из docx
# ---------------------------------------------------------------------------
def _table_header_set(table, max_rows: int = 5) -> frozenset:
    """
    Frozenset уникальных текстов ячеек из первых max_rows строк.
    [FIX-#10] max_rows поднят с 3 до 5 — синхронизировано с rpd_generate.py
    (_table_header_set там тоже max_rows=5). Расхождение давало разные
    результаты детекции таблиц при оценке vs генерации.
    """
    texts = set()
    for row in table.rows[:max_rows]:
        seen_tc: set = set()
        for cell in row.cells:
            if id(cell._tc) not in seen_tc:
                seen_tc.add(id(cell._tc))
                t = cell.text.strip()
                if t:
                    texts.add(t)
    return frozenset(texts)


def _table_all_text(table) -> str:
    """
    Весь текст таблицы без дублей merged-ячеек ВНУТРИ строки.

    [FIX-TC] Глобальная дедупликация по id(cell._tc) была некорректна:
    в шаблонах УГНТУ некоторые _tc-элементы повторяются ЧЕРЕЗ строки
    (крестообразный мердж в T4 компетенций). При глобальном seen строки 1–5
    давали 0 новых ячеек → 70 симв. вместо ~400.
    Теперь seen сбрасывается на каждой строке: дедупликация только горизонтальная.
    """
    parts = []
    for row in table.rows:
        row_seen: set = set()   # сброс на каждой строке
        for cell in row.cells:
            key = id(cell._tc)
            if key not in row_seen:
                row_seen.add(key)
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return " ".join(parts)


def extract_doc_sections(path: str, debug: bool = False) -> dict:
    """
    Читает docx, возвращает словарь секций:
    {"full": str, "competencies": str, "outcomes": str, ...}

    Таблицы определяются предикатами _SECTION_PREDICATES — теми же
    строками-маркерами, что и _TABLE_PREDICATES в rpd_generate.py.
    Это гарантирует совпадение при любом порядке таблиц в шаблоне.
    """
    doc = Document(path)
    sec_keys = list(_SECTION_PREDICATES.keys())
    sections: dict = {k: "" for k in ("full", *sec_keys)}

    para_parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    table_parts = []

    for table in doc.tables:
        hset     = _table_header_set(table)
        cell_txt = _table_all_text(table)
        table_parts.append(cell_txt)

        for sec_key, predicate in _SECTION_PREDICATES.items():
            if predicate(hset):
                sections[sec_key] = (sections[sec_key] + " " + cell_txt).strip()
                if debug:
                    print(f"   [debug] '{sec_key}': {len(cell_txt)} симв. "
                          f"| заголовки: {list(hset)[:4]}")
                break
        else:
            if debug:
                print(f"   [debug] не распознана | заголовки: {list(hset)[:4]}")

    raw_full = " ".join(para_parts) + " " + " ".join(table_parts)
    sections["full"] = re.sub(r"\s+", " ", raw_full).strip()
    return sections


# ---------------------------------------------------------------------------
# Загрузка корпуса из data_clean.jsonl
# ---------------------------------------------------------------------------

# Ключевые слова, по которым текст однозначно является шапкой вуза, а не
# названием дисциплины. Используются в _extract_discipline_name и find_reference.
_INSTITUTION_KW = re.compile(
    r"государственн|федеральн|бюджетн|образовательн|университет|"
    r"институт|академия|кафедр|министерств|рабочая программа|"
    r"направление подготовки|квалификация|форма обучения|уфа|"
    r"утверждаю|одобрено|протокол",
    re.IGNORECASE,
)

# Паттерн кода дисциплины: (38050) или (Б1.О.14) предшествует названию
_CODE_PREFIX = re.compile(r"^\s*\([А-ЯA-Z0-9\.\-]{2,12}\)\s*(.+)", re.IGNORECASE)


def _normalize_title(title: str) -> str:
    """
    [FIX-TITLE-NORM] Нормализация reference_title — удаление табличного мусора.

    Критическое исправление (отчёт §2.4): заголовки корпуса содержат мусор:
    «2 | Извлечение знаний из нейронных сетей | 7 | 3 | 9 | 4 | 1».
    Алгоритм: если ' | ' присутствует — берём самый длинный нечисловой фрагмент.
    """
    if not title:
        return title
    if " | " in title:
        parts = [p.strip() for p in title.split("|")]
        text_parts = [p for p in parts if p and not re.match(r"^\d+$", p) and len(p) > 3]
        if text_parts:
            title = max(text_parts, key=len)
    title = re.sub(r"^\d+\s+", "", title.strip())
    title = re.sub(r"\s+\d+$", "", title.strip())
    return title.strip()


def _extract_discipline_name(chunks: list) -> str:
    """
    Извлекает название дисциплины из чанков.

    Приоритет:
      1. document_meta.discipline
      2. document_meta.title (если не filename и не вузовская шапка)
      3. Текст вида «(КОД)Название дисциплины» в любом чанке
      4. Первый короткий (< 100 симв.) чанк без институциональных ключевых слов
    """
    # 1. document_meta.discipline
    for chunk in chunks[:5]:
        dm = chunk.get("document_meta") or {}
        d = dm.get("discipline", "") or dm.get("Discipline", "")
        if d and len(d) > 4 and not _INSTITUTION_KW.search(d):
            return d.strip()

    # 2. document_meta.title — не filename и не вузовская строка
    for chunk in chunks[:5]:
        dm = chunk.get("document_meta") or {}
        t = dm.get("title", "")
        if (t and len(t) > 4
                and not re.search(r"\.(docx?|json)$", t, re.I)
                and not _INSTITUTION_KW.search(t)):
            return t.strip()

    # 3. Паттерн (КОД)НазваниеДисциплины — характерен для шаблонов УГНТУ
    for chunk in chunks[:30]:
        text = chunk.get("text", "").strip()
        m = _CODE_PREFIX.match(text)
        if m:
            name = m.group(1).strip()
            if 4 < len(name) < 120 and not _INSTITUTION_KW.search(name):
                return name

    # 4. Первый короткий текст без институциональных ключевых слов
    for chunk in chunks[:30]:
        text = chunk.get("text", "").strip()
        if (5 < len(text) < 100
                and not _INSTITUTION_KW.search(text)
                and not re.search(r"\.(docx?|json)$", text, re.I)):
            return text

    return ""


# ---------------------------------------------------------------------------
# Загрузка корпуса из data_clean.jsonl
# ---------------------------------------------------------------------------
def load_corpus(jsonl_path: str) -> dict:
    """
    Возвращает:
    {
      "rpd_5.json": {
          "title":    "Нейронные сети",      ← реальное название дисциплины
          "full":     "<весь текст>",
          "sections": {"competencies": str, ...},
      },
      ...
    }
    """
    raw: dict = {}

    with open(jsonl_path, encoding="utf-8") as f:
        for line in f:
            try:
                r = json.loads(line)
            except Exception:
                continue
            src = r.get("source", "")
            if not src:
                continue
            if src not in raw:
                raw[src] = {"title": "", "chunks": []}
            raw[src]["chunks"].append(r)

    corpus: dict = {}
    for src, data in raw.items():
        # [FIX-TITLE] Извлекаем реальное название дисциплины из содержимого,
        # а не из поля title (которое часто содержит имя файла: "rpd_11.docx").
        # Без этого embedding search сравнивает "Интеллектуальные системы"
        # vs "rpd_11.docx" → бессмысленное сходство ~0.52 для всех.
        title = _extract_discipline_name(data["chunks"])
        title = _normalize_title(title)  # [FIX-TITLE-NORM]

        sec_texts: dict = {k: [] for k in _SECTION_PREDICATES.keys()}
        all_texts: list = []

        for chunk in data["chunks"]:
            text = chunk.get("text", "").strip()
            if not text:
                continue
            all_texts.append(text)
            mapped = _ST_MAP.get(chunk.get("section_type", ""))
            if mapped:
                sec_texts[mapped].append(text)
                # [FIX-LAB] В корпусе темы ЛР и ПЗ хранятся под section_type="content"
                # (после исправления баг-А в converter.py). Отдельного типа нет,
                # поэтому sec_texts["lab_works"] и ["practice"] всегда были пустыми
                # → empty_reference. Пробрасываем content-чанки в оба ключа.
                if mapped == "content":
                    sec_texts["lab_works"].append(text)
                    sec_texts["practice"].append(text)

        corpus[src] = {
            "title":    title,
            "full":     " ".join(all_texts),
            "sections": {k: " ".join(v) for k, v in sec_texts.items()},
        }

    return corpus


# ---------------------------------------------------------------------------
# Поиск эталона по embedding-similarity
# ---------------------------------------------------------------------------
def find_reference(
    discipline: str,
    corpus: dict,
) -> tuple:
    """
    Возвращает (source_name, corpus_entry, similarity).

    Стратегия:
      1. Embed discipline → cosine sim против title каждого источника.
      2. Если лучший title содержит институциональные слова (_INSTITUTION_KW)
         или similarity < 0.55 — fallback: embed первых 300 симв. текста
         секции competencies каждого источника (более содержательно, чем шапка).
      3. Берём top-1 по итоговому score.
    """
    print(f"  🔍 Embedding дисциплины: «{discipline}»")
    query_vec = embed(discipline)

    sources = [(src, data) for src, data in corpus.items()]
    print(f"  📚 Сравнение с {len(sources)} источниками (по title)...")

    best_src, best_score, best_entry = "", -1.0, {}
    for src, data in sources:
        title = data.get("title", "")
        if not title:
            continue
        try:
            title_vec = embed(title)
            score = cosine_sim(query_vec, title_vec)
            if score > best_score:
                best_score = score
                best_src   = src
                best_entry = data
        except Exception as e:
            print(f"    ⚠️  {src}: {e}")

    # [FIX-REF] Fallback когда title — название вуза, а не дисциплины.
    # Признаки плохого title: содержит институциональные слова ИЛИ similarity < 0.55.
    # Fallback: embed первых 300 симв. из секции competencies каждого источника.
    # Компетенции содержат коды (УК-1, ОПК-1) и формулировки, специфичные
    # для направления → дают осмысленное попарное сравнение.
    best_title = best_entry.get("title", "")
    if best_score < 0.55 or _INSTITUTION_KW.search(best_title):
        print(f"  ⚠️  title-поиск ненадёжен (score={best_score:.3f}, "
              f"title=«{best_title[:60]}»)\n"
              f"  🔄 Fallback: сравнение по тексту секции competencies...")
        fb_best_src, fb_best_score, fb_best_entry = "", -1.0, {}
        for src, data in sources:
            comp_text = data.get("sections", {}).get("competencies", "")
            if not comp_text:
                continue
            snippet = comp_text[:300]
            try:
                comp_vec = embed(snippet)
                score = cosine_sim(query_vec, comp_vec)
                if score > fb_best_score:
                    fb_best_score = score   # [FIX-DEAD] удалена мёртвая строка
                    fb_best_src   = src
                    fb_best_entry = data
            except Exception as e:
                print(f"    ⚠️  {src}: {e}")

        if fb_best_src and fb_best_score > best_score:
            best_src, best_score, best_entry = fb_best_src, fb_best_score, fb_best_entry
            print(f"  ✅ Fallback-эталон: {best_src} | score={best_score:.4f}")

    return best_src, best_entry, best_score


def find_section_reference(
    section_text: str,
    section_key: str,
    corpus: dict,
) -> tuple:
    """
    Ищет ближайший эталон в корпусе по embedding сгенерированного текста
    конкретной секции (а не по названию дисциплины).

    Применяется для lab_works и practice: глобальный эталон выбирается по
    заголовку дисциплины (cosine по названию), но тематика ПЗ/ЛР может
    лучше совпадать с другим документом корпуса.

    Возвращает (source_name, section_text_from_corpus, similarity).
    При отсутствии подходящих секций в корпусе возвращает ("", "", 0.0).
    """
    # Берём первые 400 символов — достаточно для характеристики тематики
    query_vec = embed(section_text[:400])
    best_src, best_score, best_text = "", -1.0, ""
    for src, data in corpus.items():
        ref_text = data.get("sections", {}).get(section_key, "").strip()
        if not ref_text:
            continue
        try:
            ref_vec = embed(ref_text[:400])
            score = cosine_sim(query_vec, ref_vec)
            if score > best_score:
                best_score = score
                best_src   = src
                best_text  = ref_text
        except Exception:
            continue
    return best_src, best_text, best_score


# ---------------------------------------------------------------------------
# Метрики
# ---------------------------------------------------------------------------
def _tokenize(text: str) -> list[str]:
    """Lowercase + split по пробелам/пунктуации (подходит для русского)."""
    return re.sub(r"[^\w\s]", " ", text.lower()).split()


def compute_bleu(hypothesis: str, reference: str) -> float:
    """
    BLEU-4 c smoothing method4.
    corpus_bleu принимает list-of-documents:
      [[ref1_tokens]], [hyp_tokens]
    """
    hyp_tok = _tokenize(hypothesis)
    ref_tok  = _tokenize(reference)
    if not hyp_tok or not ref_tok:
        return 0.0
    smooth = SmoothingFunction().method4
    return round(
        float(corpus_bleu([[ref_tok]], [hyp_tok], smoothing_function=smooth)),
        4,
    )


def compute_rouge(hypothesis: str, reference: str) -> dict[str, float]:
    """ROUGE-1, ROUGE-2, ROUGE-L (F1). Тексты предварительно стеммируются
    через SnowballStemmer('russian') — устраняет занижение 10–20% из-за
    морфологии (нейронных ≠ нейронные). [FIX-ROUGE]"""
    if not hypothesis.strip() or not reference.strip():
        return {"rouge1": 0.0, "rouge2": 0.0, "rougeL": 0.0}
    scorer = rs_lib.RougeScorer(["rouge1", "rouge2", "rougeL"], use_stemmer=False)
    scores = scorer.score(_stem_ru(reference), _stem_ru(hypothesis))
    return {
        "rouge1": round(scores["rouge1"].fmeasure, 4),
        "rouge2": round(scores["rouge2"].fmeasure, 4),
        "rougeL": round(scores["rougeL"].fmeasure, 4),
    }


def compute_all_metrics(
    gen_sections: dict[str, str],
    ref_sections: dict[str, str],
    per_section_refs: dict = None,
) -> dict[str, dict]:
    """
    Считает BLEU + ROUGE для каждой именованной секции + overall.

    per_section_refs: {section_key: (src, ref_text, score)} — если задан,
    для указанных секций используется ref_text вместо значения из ref_sections.
    Применяется для lab_works/practice, где глобальный эталон может быть
    тематически далёк от сгенерированного содержимого.

    Если секция пустая у гипотезы или эталона — ставит 0 с пометкой.
    """
    results: dict[str, dict] = {}

    for key in _SECTION_PREDICATES.keys():
        hyp = gen_sections.get(key, "").strip()

        # Секционный эталон приоритетнее глобального
        if per_section_refs and key in per_section_refs:
            ref = per_section_refs[key][1].strip()
        else:
            ref = ref_sections.get(key, "").strip()

        if not hyp and not ref:
            results[key] = {"note": "empty_both"}
            continue
        if not hyp:
            results[key] = {"bleu": 0.0, "rouge1": 0.0, "rouge2": 0.0,
                            "rougeL": 0.0, "note": "empty_hypothesis"}
            continue
        if not ref:
            results[key] = {"bleu": 0.0, "rouge1": 0.0, "rouge2": 0.0,
                            "rougeL": 0.0, "note": "empty_reference"}
            continue

        results[key] = {
            "bleu": compute_bleu(hyp, ref),
            **compute_rouge(hyp, ref),
        }

    # Overall — полный текст документа
    hyp_full = gen_sections.get("full", "").strip()
    ref_full = ref_sections.get("full", "").strip()
    results["overall"] = {
        "bleu":   compute_bleu(hyp_full, ref_full),
        **compute_rouge(hyp_full, ref_full),
    }

    return results


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
def main() -> None:
    _load_eval_cache()  # [FIX-#9] загружаем кэш эмбеддингов
    parser = argparse.ArgumentParser(
        description="Оценка качества РПД (BLEU + ROUGE)"
    )
    parser.add_argument(
        "docx", nargs="?", default=DEFAULT_DOCX,
        help=f"Сгенерированный DOCX (по умолчанию: {DEFAULT_DOCX})",
    )
    parser.add_argument(
        "--jsonl", default=DEFAULT_JSONL,
        help=f"Путь к data_clean.jsonl (по умолчанию: {DEFAULT_JSONL})",
    )
    parser.add_argument(
        "--config", default=DEFAULT_CONFIG,
        help="config.json для чтения названия дисциплины",
    )
    parser.add_argument(
        "--reference", default=None,
        help="Принудительно задать source-эталон, например: rpd_5.json",
    )
    parser.add_argument(
        "--ref-docx-dir", default=None,
        dest="ref_docx_dir",
        help="Папка с оригинальными .docx корпуса (rpd_1.docx … rpd_40.docx). "
             "Если указана — эталон извлекается из .docx напрямую, а не из jsonl. "
             "По умолчанию: автоопределение (rpd_corpus/ рядом со скриптом).",
    )
    parser.add_argument(
        "--out", default=DEFAULT_OUT,
        help=f"Путь к eval_report.json (по умолчанию: {DEFAULT_OUT})",
    )
    parser.add_argument(
        "--debug", action="store_true",
        help="Печатать заголовки каждой таблицы (диагностика пустых секций)",
    )
    args = parser.parse_args()

    # [FIX-DOCX-AUTO] Авто-детект папки с оригинальными DOCX-файлами корпуса.
    # Сравнение docx↔docx значительно точнее jsonl↔docx: при извлечении из jsonl
    # таблицы (ЛР, ПЗ, библиография) теряют структуру и попадают в виде
    # плоского чанк-текста, тогда как extract_doc_sections() читает ячейки
    # напрямую → ROUGE для lab_works/practice/bibliography систематически
    # занижен в jsonl-режиме на 5–15 пунктов.
    # Если --ref-docx-dir не задан явно, проверяем rpd_corpus/ рядом со скриптом.
    if args.ref_docx_dir is None:
        _auto_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "rpd_corpus")
        if os.path.isdir(_auto_dir):
            args.ref_docx_dir = _auto_dir
            print(f"  ℹ️  ref-docx-dir авто-определён: {_auto_dir}")

    # --- Дисциплина из config ---
    discipline = "Интеллектуальные системы"  # fallback
    if os.path.exists(args.config):
        with open(args.config, encoding="utf-8") as f:
            cfg = json.load(f)
        discipline = cfg.get("discipline", discipline)

    print(f"\n{'='*55}")
    print(f"  evaluate.py — оценка РПД")
    print(f"  Дисциплина : «{discipline}»")
    print(f"  Документ   : {args.docx}")
    print(f"  Корпус     : {args.jsonl}")
    print(f"{'='*55}\n")

    # --- Сгенерированный документ ---
    if not os.path.exists(args.docx):
        print(f"❌ Файл {args.docx!r} не найден")
        sys.exit(1)

    print("📄 Извлечение секций из сгенерированного DOCX...")
    gen_sections = extract_doc_sections(args.docx, debug=args.debug)
    for k, v in gen_sections.items():
        if k != "full":
            status = f"{len(v):5d} симв." if v else "  пусто"
            print(f"   {k:<20s}: {status}")

    # --- Корпус ---
    if not os.path.exists(args.jsonl):
        print(f"❌ Корпус {args.jsonl!r} не найден — запусти prepare_texts.py")
        sys.exit(1)

    print("\n📦 Загрузка корпуса...")
    corpus = load_corpus(args.jsonl)
    print(f"   Источников найдено: {len(corpus)}")

    # --- Поиск / выбор эталона ---
    if args.reference:
        ref_src = args.reference
        if ref_src not in corpus:
            print(f"❌ Источник {ref_src!r} не найден в корпусе")
            sys.exit(1)
        ref_entry = corpus[ref_src]
        sim_score = 1.0
        print(f"\n📌 Эталон задан вручную: {ref_src} | «{ref_entry['title']}»")
    else:
        print("\n🔎 Поиск ближайшего эталона по embedding...")
        ref_src, ref_entry, sim_score = find_reference(discipline, corpus)
        if not ref_src:
            print("❌ Не удалось найти эталон")
            sys.exit(1)
        print(
            f"   ✅ Найден : {ref_src}\n"
            f"      Заголовок : «{ref_entry['title']}»\n"
            f"      Similarity: {sim_score:.4f}"
        )

    # --- Секции эталона ---
    # Если задана папка с оригинальными docx — извлекаем из файла напрямую.
    # Это даёт более честное сравнение (docx↔docx вместо docx↔jsonl-текст).
    ref_sections_raw = ref_entry["sections"].copy()
    ref_sections_raw["full"] = ref_entry["full"]

    if args.ref_docx_dir:
        ref_docx_name = re.sub(r"\.json$", ".docx", ref_src)
        ref_docx_path = os.path.join(args.ref_docx_dir, ref_docx_name)
        if os.path.exists(ref_docx_path):
            print(f"\n📄 Извлечение секций из эталонного DOCX: {ref_docx_path}")
            ref_sections_raw = extract_doc_sections(ref_docx_path)
            print("   ✅ Использован docx-эталон (наиболее точное сравнение)")
        else:
            print(
                f"   ⚠️  {ref_docx_path} не найден — "
                "используются тексты из data_clean.jsonl"
            )

    # --- Подсчёт метрик ---
    # [PER-SEC] Для lab_works и practice ищем секционный эталон отдельно:
    # глобальный эталон выбирается по cosine заголовка дисциплины, но тематика
    # ПЗ/ЛР может лучше совпасть с другим документом корпуса.
    # Секционный эталон применяется только если score заметно выше глобального.
    _PER_SEC_KEYS = ("lab_works", "practice")
    per_section_refs: dict = {}

    print("\n🔎 Поиск секционных эталонов для ЛР/ПЗ...")
    for sec_key in _PER_SEC_KEYS:
        sec_text = gen_sections.get(sec_key, "").strip()
        if not sec_text:
            print(f"   {sec_key}: секция пуста — пропуск")
            continue
        src_s, text_s, score_s = find_section_reference(sec_text, sec_key, corpus)
        if not src_s:
            print(f"   {sec_key}: эталон не найден в корпусе")
            continue
        if score_s > sim_score + 0.02:
            # Секционный эталон лучше → подгружаем из docx, если доступен
            if args.ref_docx_dir:
                sec_docx_name = re.sub(r"\.json$", ".docx", src_s)
                sec_docx_path = os.path.join(args.ref_docx_dir, sec_docx_name)
                if os.path.exists(sec_docx_path):
                    sec_ref_sections = extract_doc_sections(sec_docx_path)
                    text_s = sec_ref_sections.get(sec_key, text_s)
            per_section_refs[sec_key] = (src_s, text_s, score_s)
            title_s = corpus.get(src_s, {}).get("title", src_s)
            print(f"   {sec_key}: «{title_s}» ({src_s}, {score_s:.4f})"
                  f" > глобальный ({sim_score:.4f}) ✅")
        else:
            print(f"   {sec_key}: глобальный эталон достаточен"
                  f" ({sim_score:.4f} vs {score_s:.4f})")

    print("\n📊 Подсчёт метрик...")
    metrics = compute_all_metrics(gen_sections, ref_sections_raw, per_section_refs)

    # --- Вывод таблицы ---
    col = "{:<20s} {:>7} {:>8} {:>8} {:>8}"
    print(f"\n{'='*57}")
    print(col.format("Секция", "BLEU", "ROUGE-1", "ROUGE-2", "ROUGE-L"))
    print("-"*57)
    for key in (*_SECTION_PREDICATES.keys(), "overall"):
        vals = metrics.get(key, {})
        note = vals.get("note", "")
        if note and "empty" in note:
            print(f"  {key:<20s} {'—':>7} {'—':>8} {'—':>8}   ({note})")
        else:
            print(col.format(
                f"  {key}",
                f"{vals.get('bleu',  0):.4f}",
                f"{vals.get('rouge1',0):.4f}",
                f"{vals.get('rouge2',0):.4f}",
                f"{vals.get('rougeL',0):.4f}",
            ))
    print("="*57)

    # --- Интерпретация overall ---
    # Пороги скорректированы под RAG-генерацию:
    # LLM перефразирует корпус → BLEU всегда низкий (~0.01–0.15).
    # ROUGE-1 более показателен: 0.25+ = структурное совпадение;
    # 0.40+ = хорошее доменное попадание.
    ov = metrics.get("overall", {})
    bleu_val = ov.get("bleu", 0)
    r1_val   = ov.get("rouge1", 0)
    out_r1   = metrics.get("outcomes", {}).get("rouge1", 0)
    if r1_val >= 0.35 or out_r1 >= 0.45:
        verdict = "✅ Хорошее совпадение — структура и outcomes совпадают с эталоном"
    elif r1_val >= 0.20 or out_r1 >= 0.30:
        verdict = "🟡 Умеренное совпадение — для RAG-генерации норма; проверь domain drift"
    else:
        verdict = "🔴 Низкое совпадение — возможен domain drift или слабый retrieval"
    print(f"\n  {verdict}")

    # --- Сохранение отчёта ---
    report = {
        "discipline":        discipline,
        "generated_docx":    args.docx,
        "reference_source":  ref_src,
        "reference_title":   ref_entry.get("title", ""),
        "reference_mode":    "docx" if (args.ref_docx_dir and ref_sections_raw.get("full") != ref_entry["full"]) else "jsonl",
        "embedding_similarity": round(sim_score, 4),
        # [PER-SEC] Секционные эталоны для lab_works/practice (если отличаются от глобального)
        "section_references": {
            k: {"source": v[0], "title": corpus.get(v[0], {}).get("title", ""), "similarity": round(v[2], 4)}
            for k, v in per_section_refs.items()
        },
        "metrics":           metrics,
        "timestamp":         time.strftime("%Y-%m-%dT%H:%M:%S"),
    }

    with open(args.out, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    print(f"\n✅ Отчёт сохранён: {args.out}\n")
    _save_eval_cache()  # [FIX-#9] сохраняем кэш эмбеддингов


if __name__ == "__main__":
    main()
